"""
ExportService — Phase 8B.

Provides export functions for Lab Notebook records and Protocols.

PDF / DOCX:  reportlab ≥ 4.x  /  python-docx ≥ 1.x  (optional but pre-installed)
JSON / Markdown:  pure Python (always available)

Error handling
--------------
Heavy-dependency functions raise ExportDependencyError if the library
is missing.  Callers should catch that and show a toast instead of crashing.
"""
from __future__ import annotations

import json
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any


# ── Helpers ───────────────────────────────────────────────────────────────────

class ExportDependencyError(Exception):
    """Raised when an optional export dependency is not installed."""
    def __init__(self, dep: str, install_cmd: str) -> None:
        self.dep         = dep
        self.install_cmd = install_cmd
        super().__init__(f"Missing: {dep}. Install with: {install_cmd}")


def _safe_name(text: str, max_len: int = 60) -> str:
    """Strip characters illegal in file names and truncate."""
    s = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", str(text))
    s = s.strip(". ")
    return s[:max_len] or "export"


def _fmt_secs(secs: float) -> str:
    s = int(secs)
    if s <= 0:
        return "—"
    h, rem = divmod(s, 3600)
    m, sc  = divmod(rem, 60)
    if h:
        return f"{h}h {m}m" if m else f"{h}h"
    return f"{m}m {sc}s" if sc else f"{m}m"


def _fmt_mins(mins: float) -> str:
    m = int(round(mins))
    if m <= 0:
        return "—"
    h, r = divmod(m, 60)
    if h:
        return f"{h}h {r}m" if r else f"{h}h"
    return f"{m}m"


def _fmt_ts(ts_ms: int | float | None, fmt: str = "%B %d, %Y  %H:%M") -> str:
    if not ts_ms:
        return "—"
    try:
        return datetime.fromtimestamp(float(ts_ms) / 1000).strftime(fmt)
    except Exception:
        return "—"


def notebook_default_name(record: dict, fmt: str) -> str:
    title    = record.get("title", "Session")
    ts       = record.get("startedAt", 0)
    date_str = _fmt_ts(ts, "%Y-%m-%d") if ts else "unknown"
    ext      = {"pdf": ".pdf", "docx": ".docx", "json": ".json"}.get(fmt, ".txt")
    return f"BenchFlow_{_safe_name(title)}_{date_str}{ext}"


def protocol_default_name(protocol: dict, fmt: str) -> str:
    name = protocol.get("name", "Protocol")
    ext  = {"pdf": ".pdf", "json": ".json", "md": ".md"}.get(fmt, ".txt")
    return f"BenchFlow_Protocol_{_safe_name(name)}{ext}"


# ── JSON exports (no dependencies) ───────────────────────────────────────────

def export_notebook_json(record: dict, path: str) -> None:
    """Write the full record dict as pretty JSON."""
    payload = {
        "_exported_by": "BenchFlow Qt",
        "_exported_at": datetime.now().isoformat(),
        "record": record,
    }
    Path(path).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def export_protocol_json(protocol: dict, path: str) -> None:
    """Write the protocol dict as pretty JSON."""
    payload = {
        "_exported_by": "BenchFlow Qt",
        "_exported_at": datetime.now().isoformat(),
        "protocol": protocol,
    }
    Path(path).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── Markdown export (no dependencies) ────────────────────────────────────────

def export_protocol_markdown(protocol: dict, path: str) -> None:
    """Write protocol as a structured Markdown document."""
    lines: list[str] = []
    a = lines.append

    name = protocol.get("name", "Protocol")
    a(f"# {name}")
    a("")

    cat  = protocol.get("category", "")
    tags = protocol.get("tags", [])
    desc = protocol.get("description", "")

    if cat:
        a(f"**Category:** {cat}  ")
    if tags:
        a(f"**Tags:** {', '.join(tags)}  ")
    updated = protocol.get("updatedAt", protocol.get("createdAt", 0))
    if updated:
        a(f"**Updated:** {_fmt_ts(updated, '%Y-%m-%d')}  ")
    a("")
    if desc:
        a(f"{desc}")
        a("")

    steps = protocol.get("steps", [])
    total_m = sum(
        float(s.get("handsOnMinutes", 0)) +
        float(s.get("waitMinutes",    0)) +
        float(s.get("bufferMinutes",  0))
        for s in steps
    )
    a(f"**{len(steps)} steps · Total: {_fmt_mins(total_m)}**")
    a("")
    a("---")
    a("")

    for i, step in enumerate(steps):
        title    = step.get("title", f"Step {i+1}")
        stype    = step.get("type", "other").replace("_", " ").title()
        hands_on = float(step.get("handsOnMinutes", 0))
        wait     = float(step.get("waitMinutes",    0))
        buf      = float(step.get("bufferMinutes",  0))
        total_s  = hands_on + wait + buf

        a(f"## Step {i+1}: {title}")
        a("")

        timing_parts = [f"**Type:** {stype}"]
        if total_s > 0:
            timing_parts.append(f"**Total:** {_fmt_mins(total_s)}")
        if hands_on > 0:
            timing_parts.append(f"Hands-on: {_fmt_mins(hands_on)}")
        if wait > 0:
            timing_parts.append(f"Wait: {_fmt_mins(wait)}")
        a("  ·  ".join(timing_parts))
        a("")

        for field, label in [
            ("temperature",         "🌡 Temperature"),
            ("centrifugeCondition", "🔄 Centrifuge"),
            ("shakingRotation",     "↺ Shaking/RPM"),
        ]:
            val = step.get(field, "").strip()
            if val:
                a(f"**{label}:** {val}  ")

        description = step.get("description", "").strip()
        if description:
            a("")
            a(description)

        reagents = step.get("reagents", [])
        if reagents:
            a("")
            a("**Reagents:**")
            for r in reagents:
                n = r.get("name", "")
                amt = r.get("amount", r.get("amountUnit", "")).strip()
                unit = r.get("unit",  r.get("amountUnit", "")).strip()
                # avoid duplicating if amt and unit both came from 'amountUnit'
                parts = [x for x in [amt, unit] if x and x != amt]
                suffix = f": {amt} {' '.join(parts)}".strip() if amt else ""
                a(f"- {n}{suffix}")

        equipment = step.get("equipment", [])
        if equipment:
            a("")
            a("**Equipment:**")
            for e in equipment:
                a(f"- {str(e).strip()}")

        notes    = step.get("notes", "").strip()
        warnings = step.get("warnings", "").strip()
        if notes:
            a("")
            a(f"**Notes:** {notes}")
        if warnings:
            a("")
            a(f"**⚠ Warnings:** {warnings}")

        checklist = step.get("checklist", [])
        if checklist:
            a("")
            a("**Checklist:**")
            for item in checklist:
                txt = str(item.get("text", item) if isinstance(item, dict) else item).strip()
                if txt:
                    a(f"- [ ] {txt}")

        substeps = step.get("substeps", [])
        if substeps:
            a("")
            a("**Substeps:**")
            for j, s in enumerate(substeps):
                txt = str(s.get("text", s) if isinstance(s, dict) else s).strip()
                if txt:
                    a(f"  {j+1}. {txt}")

        a("")
        a("---")
        a("")

    Path(path).write_text("\n".join(lines), encoding="utf-8")


# ── PDF helpers (reportlab) ───────────────────────────────────────────────────

_PDF_STYLES: Any = None
_PDF_COLORS: dict[str, Any] = {}

def _init_pdf():
    global _PDF_STYLES, _PDF_COLORS
    try:
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
        _PDF_STYLES = getSampleStyleSheet()
        _PDF_COLORS = {
            "dark":    colors.HexColor("#1e293b"),
            "mid":     colors.HexColor("#475569"),
            "muted":   colors.HexColor("#64748b"),
            "accent":  colors.HexColor("#3b82f6"),
            "success": colors.HexColor("#22c55e"),
            "warning": colors.HexColor("#f97316"),
            "danger":  colors.HexColor("#ef4444"),
            "white":   colors.white,
            "light":   colors.HexColor("#f1f5f9"),
            "hdr_bg":  colors.HexColor("#1e3a5f"),
            "alt_row": colors.HexColor("#f8fafc"),
        }
        return True
    except ImportError:
        return False


def _para(text: str, style_name: str = "Normal",
          font_size: int = 10, bold: bool = False,
          color_hex: str = "#1e293b",
          space_before: float = 4, space_after: float = 4) -> Any:
    """Build a reportlab Paragraph with given style."""
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors

    style = ParagraphStyle(
        name=f"_custom_{id(text)}",
        fontSize=font_size,
        fontName="Helvetica-Bold" if bold else "Helvetica",
        textColor=colors.HexColor(color_hex),
        spaceAfter=space_after,
        spaceBefore=space_before,
        leading=font_size * 1.35,
    )
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def _h1(text: str) -> Any:
    return _para(text, font_size=20, bold=True, color_hex="#1e293b",
                 space_before=0, space_after=8)

def _h2(text: str) -> Any:
    return _para(text, font_size=14, bold=True, color_hex="#1e3a5f",
                 space_before=12, space_after=4)

def _h3(text: str) -> Any:
    return _para(text, font_size=11, bold=True, color_hex="#334155",
                 space_before=8, space_after=3)

def _body(text: str, color: str = "#475569") -> Any:
    return _para(text, font_size=10, color_hex=color)

def _muted(text: str) -> Any:
    return _para(text, font_size=9, color_hex="#64748b")


# ── Notebook PDF export ───────────────────────────────────────────────────────

def export_notebook_pdf(record: dict, path: str) -> None:
    """Export a Lab Notebook session record to PDF using reportlab."""
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable,
        )
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
    except ImportError:
        raise ExportDependencyError("reportlab", "pip install reportlab")

    if not _init_pdf():
        raise ExportDependencyError("reportlab", "pip install reportlab")

    C = _PDF_COLORS
    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=20 * mm, leftMargin=20 * mm,
        topMargin=20 * mm,   bottomMargin=20 * mm,
    )
    story: list[Any] = []
    sp = lambda h=6: Spacer(1, h)
    hr = lambda: HRFlowable(width="100%", thickness=0.5, color=C["muted"])

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(_h1(record.get("title", "Lab Session")))

    started  = record.get("startedAt", 0)
    ended    = record.get("endedAt",   0)
    dur_secs = record.get("actualDuration", 0)
    if not dur_secs and started and ended:
        dur_secs = max(0, (ended - started) / 1000)

    meta = [
        ["Date",     _fmt_ts(started, "%B %d, %Y")],
        ["Time",     f"{_fmt_ts(started, '%H:%M')}  –  {_fmt_ts(ended, '%H:%M')}"],
        ["Duration", _fmt_secs(dur_secs)],
        ["Protocol", record.get("protocolName", "—")],
    ]
    meta_tbl = Table(meta, colWidths=[35 * mm, None])
    meta_tbl.setStyle(TableStyle([
        ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1,-1), 9),
        ("TEXTCOLOR",   (0, 0), (0, -1), C["mid"]),
        ("TEXTCOLOR",   (1, 0), (1, -1), C["dark"]),
        ("TOPPADDING",  (0, 0), (-1,-1), 2),
        ("BOTTOMPADDING",(0,0), (-1,-1), 2),
    ]))
    story += [meta_tbl, sp(8), hr(), sp(6)]

    # ── Stats ─────────────────────────────────────────────────────────────────
    step_records = record.get("stepRecords", [])
    total    = len(step_records)
    done     = sum(1 for s in step_records if s.get("status") == "completed")
    skipped  = sum(1 for s in step_records if s.get("status") == "skipped")
    pct      = int(done / total * 100) if total else 0
    story.append(_muted(
        f"Steps: {total}  ·  Completed: {done}  ·  Skipped: {skipped}  ·  "
        f"Completion: {pct}%"
    ))
    story.append(sp(4))

    # ── Observations ─────────────────────────────────────────────────────────
    obs = (record.get("observations") or record.get("notes") or "").strip()
    if obs:
        story.append(_h3("Observations / Summary"))
        story.append(_body(obs))
        story.append(sp(6))

    # ── Step Records table ────────────────────────────────────────────────────
    if step_records:
        story += [hr(), sp(6), _h2("Step Records")]
        hdr = ["#", "Step", "Planned", "Actual", "Status"]
        rows = [hdr]
        for i, sr in enumerate(step_records):
            status = sr.get("status", "idle")
            status_color_map = {
                "completed": "#22c55e",
                "skipped":   "#f97316",
                "pending":   "#64748b",
                "idle":      "#64748b",
                "incomplete":"#ef4444",
            }
            rows.append([
                str(i + 1),
                sr.get("stepTitle", f"Step {i+1}"),
                _fmt_secs(float(sr.get("plannedSecs", 0))),
                _fmt_secs(float(sr.get("usedSecs",    0))) if sr.get("usedSecs") else "—",
                status,
            ])

        col_widths = [10 * mm, None, 22 * mm, 22 * mm, 26 * mm]
        tbl = Table(rows, colWidths=col_widths, repeatRows=1)
        style = TableStyle([
            # Header
            ("BACKGROUND",    (0, 0), (-1, 0), C["hdr_bg"]),
            ("TEXTCOLOR",     (0, 0), (-1, 0), C["white"]),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1,-1), 8),
            ("ALIGN",         (2, 0), (-1,-1), "CENTER"),
            ("ALIGN",         (0, 0), (0, -1), "CENTER"),
            ("ROWBACKGROUNDS",(0, 1), (-1,-1), [colors.white, C["alt_row"]]),
            ("TOPPADDING",    (0, 0), (-1,-1), 3),
            ("BOTTOMPADDING", (0, 0), (-1,-1), 3),
            ("LINEBELOW",     (0, 0), (-1, 0), 0.5, C["mid"]),
            ("GRID",          (0, 0), (-1,-1), 0.25, C["muted"]),
        ])
        # Color status cells
        for row_idx, sr in enumerate(step_records, start=1):
            status = sr.get("status", "idle")
            c_hex  = {"completed": "#22c55e", "skipped": "#f97316",
                      "incomplete": "#ef4444"}.get(status, "#64748b")
            style.add("TEXTCOLOR", (4, row_idx), (4, row_idx), colors.HexColor(c_hex))
        tbl.setStyle(style)
        story.append(tbl)
        story.append(sp(8))

    # ── Timeline ──────────────────────────────────────────────────────────────
    timeline = record.get("timeline", [])
    if timeline:
        story += [hr(), sp(6), _h2("Timeline")]
        tl_rows = [["Time", "Event"]]
        for entry in timeline:
            tl_rows.append([
                entry.get("time", ""),
                entry.get("text", ""),
            ])
        tl_tbl = Table(tl_rows, colWidths=[22 * mm, None], repeatRows=1)
        tl_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0), C["hdr_bg"]),
            ("TEXTCOLOR",     (0, 0), (-1, 0), C["white"]),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1,-1), 8),
            ("ROWBACKGROUNDS",(0, 1), (-1,-1), [colors.white, C["alt_row"]]),
            ("TOPPADDING",    (0, 0), (-1,-1), 2),
            ("BOTTOMPADDING", (0, 0), (-1,-1), 2),
            ("LINEBELOW",     (0, 0), (-1, 0), 0.5, C["mid"]),
        ]))
        story.append(tl_tbl)
        story.append(sp(8))

    # ── Additional notes ─────────────────────────────────────────────────────
    notes = record.get("summary", "").strip()
    if notes:
        story += [hr(), sp(4), _h3("Additional Notes"), _body(notes)]

    # Footer
    story += [sp(12), hr(), _muted(f"Exported by BenchFlow Qt  ·  {datetime.now().strftime('%Y-%m-%d %H:%M')}")]

    doc.build(story)


# ── Notebook DOCX export ──────────────────────────────────────────────────────

def export_notebook_docx(record: dict, path: str) -> None:
    """Export a Lab Notebook session record to .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise ExportDependencyError("python-docx", "pip install python-docx")

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def _h(text: str, level: int = 1) -> None:
        doc.add_heading(text, level=level)

    def _p(text: str, bold: bool = False, italic: bool = False,
           color: tuple = (71, 85, 105)) -> None:
        p = doc.add_paragraph()
        r = p.add_run(str(text))
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(*color)

    def _meta(key: str, val: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(val)
        p.paragraph_format.space_after = Pt(2)

    started  = record.get("startedAt", 0)
    ended    = record.get("endedAt",   0)
    dur_secs = record.get("actualDuration", 0)
    if not dur_secs and started and ended:
        dur_secs = max(0, (ended - started) / 1000)

    # ── Title ─────────────────────────────────────────────────────────────────
    _h(record.get("title", "Lab Session"), level=1)

    _meta("Date",     _fmt_ts(started, "%B %d, %Y"))
    _meta("Time",     f"{_fmt_ts(started, '%H:%M')} – {_fmt_ts(ended, '%H:%M')}")
    _meta("Duration", _fmt_secs(dur_secs))
    _meta("Protocol", record.get("protocolName", "—"))
    doc.add_paragraph()

    # ── Stats ─────────────────────────────────────────────────────────────────
    step_records = record.get("stepRecords", [])
    total   = len(step_records)
    done    = sum(1 for s in step_records if s.get("status") == "completed")
    skipped = sum(1 for s in step_records if s.get("status") == "skipped")
    pct     = int(done / total * 100) if total else 0
    _p(f"Steps: {total}  ·  Completed: {done}  ·  Skipped: {skipped}  ·  {pct}% complete",
       color=(100, 116, 139))

    # ── Observations ─────────────────────────────────────────────────────────
    obs = (record.get("observations") or record.get("notes") or "").strip()
    if obs:
        doc.add_paragraph()
        _h("Observations / Summary", level=2)
        _p(obs)

    # ── Step Records table ────────────────────────────────────────────────────
    if step_records:
        doc.add_paragraph()
        _h("Step Records", level=2)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for j, h in enumerate(["#", "Step", "Planned", "Actual", "Status"]):
            hdr[j].text = h
            for para in hdr[j].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for i, sr in enumerate(step_records):
            row   = tbl.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = sr.get("stepTitle", f"Step {i+1}")
            row[2].text = _fmt_secs(float(sr.get("plannedSecs", 0)))
            row[3].text = _fmt_secs(float(sr.get("usedSecs", 0))) if sr.get("usedSecs") else "—"
            row[4].text = sr.get("status", "—")
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    # ── Timeline ──────────────────────────────────────────────────────────────
    timeline = record.get("timeline", [])
    if timeline:
        doc.add_paragraph()
        _h("Timeline", level=2)
        for entry in timeline:
            t = entry.get("time", "")
            txt = entry.get("text", "")
            if txt:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{t}  ").bold = True
                p.add_run(txt).font.size = Pt(9)

    # ── Additional notes ──────────────────────────────────────────────────────
    notes = record.get("summary", "").strip()
    if notes:
        doc.add_paragraph()
        _h("Additional Notes", level=2)
        _p(notes)

    doc.add_paragraph()
    _p(f"Exported by BenchFlow Qt  ·  {datetime.now().strftime('%Y-%m-%d %H:%M')}",
       italic=True, color=(148, 163, 184))

    doc.save(path)


# ── Protocol PDF export ───────────────────────────────────────────────────────

def export_protocol_pdf(protocol: dict, path: str) -> None:
    """Export a protocol to PDF using reportlab."""
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Spacer, Table, TableStyle, HRFlowable,
        )
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
    except ImportError:
        raise ExportDependencyError("reportlab", "pip install reportlab")

    if not _init_pdf():
        raise ExportDependencyError("reportlab", "pip install reportlab")

    C = _PDF_COLORS
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        rightMargin=20 * mm, leftMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    story: list[Any] = []
    sp = lambda h=6: Spacer(1, h)
    hr = lambda: HRFlowable(width="100%", thickness=0.5, color=C["muted"])

    # Header
    story.append(_h1(protocol.get("name", "Protocol")))

    cat  = protocol.get("category", "")
    tags = protocol.get("tags", [])
    desc = protocol.get("description", "")

    if cat or tags:
        story.append(_muted(
            "  ·  ".join(filter(None, [cat, ", ".join(tags)]))
        ))
    if desc:
        story.append(_body(desc))

    steps = protocol.get("steps", [])
    total_m = sum(
        float(s.get("handsOnMinutes", 0)) +
        float(s.get("waitMinutes",    0)) +
        float(s.get("bufferMinutes",  0))
        for s in steps
    )
    story.append(_muted(f"{len(steps)} steps  ·  Total: {_fmt_mins(total_m)}"))
    story += [sp(6), hr(), sp(6)]

    for i, step in enumerate(steps):
        title   = step.get("title", f"Step {i+1}")
        stype   = step.get("type", "other").replace("_", " ").title()
        hands_on = float(step.get("handsOnMinutes", 0))
        wait     = float(step.get("waitMinutes",    0))
        buf      = float(step.get("bufferMinutes",  0))
        total_s  = hands_on + wait + buf

        story.append(_h2(f"{i+1}. {title}"))

        meta_row: list[str] = [stype]
        if total_s > 0:
            meta_row.append(f"Total: {_fmt_mins(total_s)}")
        if hands_on > 0:
            meta_row.append(f"Hands-on: {_fmt_mins(hands_on)}")
        if wait > 0:
            meta_row.append(f"Wait: {_fmt_mins(wait)}")
        story.append(_muted("  ·  ".join(meta_row)))

        for field, label in [
            ("temperature", "Temperature"),
            ("centrifugeCondition", "Centrifuge"),
            ("shakingRotation", "Shaking/RPM"),
        ]:
            val = step.get(field, "").strip()
            if val:
                story.append(_muted(f"{label}: {val}"))

        description = step.get("description", "").strip()
        if description:
            story += [sp(3), _body(description)]

        reagents  = step.get("reagents", [])
        if reagents:
            story.append(_h3("Reagents"))
            rows = [["Name", "Amount", "Unit"]]
            for r in reagents:
                rows.append([
                    r.get("name", ""),
                    r.get("amount", ""),
                    r.get("unit", r.get("amountUnit", "")),
                ])
            r_tbl = Table(rows, colWidths=[None, 25 * mm, 25 * mm])
            r_tbl.setStyle(TableStyle([
                ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",    (0, 0), (-1,-1), 8),
                ("BACKGROUND",  (0, 0), (-1, 0), C["hdr_bg"]),
                ("TEXTCOLOR",   (0, 0), (-1, 0), C["white"]),
                ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.white, C["alt_row"]]),
                ("TOPPADDING",  (0, 0), (-1,-1), 2),
                ("BOTTOMPADDING",(0,0), (-1,-1), 2),
                ("GRID",        (0, 0), (-1,-1), 0.25, C["muted"]),
            ]))
            story.append(r_tbl)

        equipment = step.get("equipment", [])
        if equipment:
            story += [_h3("Equipment"),
                      _body(", ".join(str(e).strip() for e in equipment))]

        for field, label, color in [
            ("notes",    "Notes",      "#475569"),
            ("warnings", "⚠ Warnings", "#f97316"),
        ]:
            val = step.get(field, "").strip()
            if val:
                story += [sp(2), _muted(label), _body(val, color)]

        story.append(sp(4))

    story += [hr(), _muted(
        f"Exported by BenchFlow Qt  ·  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )]
    doc.build(story)
