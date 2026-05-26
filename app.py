#!/usr/bin/env python3
"""BenchFlow — Wet Lab Protocol Manager"""

import json, os, sys, time, threading, uuid, calendar as _calendar
from datetime import datetime, timedelta
from pathlib import Path
import tkinter as tk
from tkinter import messagebox, filedialog
import customtkinter as ctk

# Optional import libs
try:
    import fitz as pymupdf  # PyMuPDF
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False

try:
    import docx as _docx_lib  # python-docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("blue")

# ─── Data directory ───────────────────────────────────────────────────────────
APP_DIR = Path.home() / "Library" / "Application Support" / "BenchFlow"
APP_DIR.mkdir(parents=True, exist_ok=True)
PROTOCOLS_FILE  = APP_DIR / "protocols.json"
RUNS_FILE       = APP_DIR / "runs.json"
CATEGORIES_FILE = APP_DIR / "categories.json"
TAGS_FILE       = APP_DIR / "tags.json"
SCHEDULE_FILE   = APP_DIR / "schedule.json"
RUNTIME_FILE    = APP_DIR / "runtime_session.json"  # crash-recovery checkpoint
# Built-in protocol templates directory.
# When frozen by PyInstaller, data files live in sys._MEIPASS.
# When running from source, they sit next to app.py.
_APP_BASE = (Path(sys._MEIPASS)      # type: ignore[attr-defined]
             if getattr(sys, "frozen", False)
             else Path(__file__).parent)
TEMPLATES_DIR   = _APP_BASE / "templates"

# ─── Palette ──────────────────────────────────────────────────────────────────
# (light, dark)
SB       = ("#1e293b", "#0f172a")       # sidebar bg
SB_ACT   = ("#3b82f6", "#3b82f6")       # active nav item
SB_HOV   = ("#334155", "#334155")       # hover
SB_TXT   = ("#94a3b8", "#94a3b8")       # sidebar text
SB_ATXT  = ("#ffffff", "#ffffff")       # active text
BG       = ("#f1f5f9", "#1e293b")       # main bg
CARD     = ("#ffffff", "#1e293b")       # card bg — note: CTk uses (light, dark) tuples
CARD_B   = ("#e2e8f0", "#334155")       # card border
T1       = ("#0f172a", "#f1f5f9")       # primary text
T2       = ("#64748b", "#94a3b8")       # secondary text
T3       = ("#94a3b8", "#475569")       # muted
ACC      = ("#3b82f6", "#60a5fa")       # accent blue
ACCH     = ("#2563eb", "#3b82f6")       # accent hover
DANGER   = ("#ef4444", "#f87171")       # red
GREEN    = ("#22c55e", "#4ade80")       # green
ORANGE   = ("#f97316", "#fb923c")       # orange

STEP_COLORS = {
    "preparation":      ("#dbeafe", "#1e3a5f"),
    "reagent_addition": ("#ccfbf1", "#042f2e"),
    "mixing":           ("#ede9fe", "#2e1065"),
    "incubation":       ("#ffedd5", "#431407"),
    "waiting":          ("#f1f5f9", "#1e293b"),
    "centrifuge":       ("#f5f3ff", "#1e1b4b"),
    "wash":             ("#cffafe", "#164e63"),
    "transfer":         ("#e0e7ff", "#1e1b4b"),
    "pipetting":        ("#e0f2fe", "#0c4a6e"),
    "resuspension":     ("#d1fae5", "#052e16"),
    "staining":         ("#fce7f3", "#500724"),
    "blocking":         ("#ffe4e6", "#4c0519"),
    "electrophoresis":  ("#f3e8ff", "#3b0764"),
    "gel_running":      ("#ede9fe", "#2e1065"),
    "membrane_transfer":("#dbeafe", "#1e3a5f"),
    "imaging":          ("#ecfdf5", "#064e3b"),
    "measurement":      ("#fefce8", "#422006"),
    "sample_collection":("#fef3c7", "#451a03"),
    "harvest":          ("#f7fee7", "#1a2e05"),
    "lysis":            ("#fee2e2", "#450a0a"),
    "heating":          ("#fff7ed", "#431407"),
    "cooling":          ("#f0f9ff", "#0c4a6e"),
    "storage":          ("#f8fafc", "#1e293b"),
    "note":             ("#fefce8", "#422006"),
    "checklist_block":  ("#dcfce7", "#052e16"),
    "decision":         ("#fef3c7", "#451a03"),
    "custom":           ("#fef9c3", "#422006"),
    "analysis":         ("#dcfce7", "#052e16"),  # backward compat
}
STEP_BADGES = {
    "preparation":      ("#93c5fd",  "#1d4ed8"),
    "reagent_addition": ("#5eead4",  "#0d9488"),
    "mixing":           ("#c4b5fd",  "#6d28d9"),
    "incubation":       ("#fdba74",  "#c2410c"),
    "waiting":          ("#cbd5e1",  "#475569"),
    "centrifuge":       ("#a5b4fc",  "#4338ca"),
    "wash":             ("#67e8f9",  "#0e7490"),
    "transfer":         ("#a5b4fc",  "#3730a3"),
    "pipetting":        ("#7dd3fc",  "#0284c7"),
    "resuspension":     ("#6ee7b7",  "#059669"),
    "staining":         ("#f9a8d4",  "#9d174d"),
    "blocking":         ("#fda4af",  "#be123c"),
    "electrophoresis":  ("#d8b4fe",  "#7c3aed"),
    "gel_running":      ("#c4b5fd",  "#6d28d9"),
    "membrane_transfer":("#93c5fd",  "#1d4ed8"),
    "imaging":          ("#6ee7b7",  "#059669"),
    "measurement":      ("#fef08a",  "#ca8a04"),
    "sample_collection":("#fcd34d",  "#d97706"),
    "harvest":          ("#bef264",  "#65a30d"),
    "lysis":            ("#fca5a5",  "#dc2626"),
    "heating":          ("#fed7aa",  "#ea580c"),
    "cooling":          ("#7dd3fc",  "#0369a1"),
    "storage":          ("#94a3b8",  "#334155"),
    "note":             ("#fef08a",  "#854d0e"),
    "checklist_block":  ("#86efac",  "#15803d"),
    "decision":         ("#fde68a",  "#b45309"),
    "custom":           ("#fde047",  "#854d0e"),
    "analysis":         ("#86efac",  "#15803d"),  # backward compat
}
STEP_LABELS = {
    "preparation":      "Preparation",
    "reagent_addition": "Reagent Addition",
    "mixing":           "Mixing",
    "incubation":       "Incubation",
    "waiting":          "Waiting",
    "centrifuge":       "Centrifuge",
    "wash":             "Wash",
    "transfer":         "Transfer",
    "pipetting":        "Pipetting",
    "resuspension":     "Resuspension",
    "staining":         "Staining",
    "blocking":         "Blocking",
    "electrophoresis":  "Electrophoresis",
    "gel_running":      "Gel Running",
    "membrane_transfer":"Membrane Transfer",
    "imaging":          "Imaging",
    "measurement":      "Measurement",
    "sample_collection":"Sample Collection",
    "harvest":          "Harvest",
    "lysis":            "Lysis",
    "heating":          "Heating",
    "cooling":          "Cooling",
    "storage":          "Storage",
    "note":             "Note",
    "checklist_block":  "Checklist",
    "decision":         "Decision",
    "custom":           "Custom",
    "analysis":         "Analysis",   # backward compat
}

# ─── Timeline block types (Schedule right-panel editor) ──────────────────────
BLOCK_TYPE_LABELS = {
    "protocol_step": "Protocol Step",
    "break":         "Break",
    "task":          "Task",
    "note":          "Note",
    "decision":      "Decision",
    "custom":        "Custom",
}
# (light_bg, accent_color, badge_bg)
BLOCK_TYPE_COLORS = {
    "protocol_step": ("#dbeafe", "#3b82f6", "#93c5fd"),
    "break":         ("#f3f4f6", "#6b7280", "#d1d5db"),
    "task":          ("#d1fae5", "#10b981", "#6ee7b7"),
    "note":          ("#fefce8", "#ca8a04", "#fde68a"),
    "decision":      ("#fee2e2", "#ef4444", "#fca5a5"),
    "custom":        ("#f5f3ff", "#8b5cf6", "#c4b5fd"),
}
BLOCK_STATUS_LABELS = {
    "planned":  "Planned",
    "done":     "Done",
    "skipped":  "Skipped",
    "canceled": "Canceled",
    "modified": "Modified",
}

STEP_TYPES = [
    "preparation","reagent_addition","mixing","incubation","waiting",
    "centrifuge","wash","transfer","pipetting","resuspension",
    "staining","blocking","electrophoresis","gel_running","membrane_transfer",
    "imaging","measurement","sample_collection","harvest","lysis",
    "heating","cooling","storage","note","checklist_block","decision","custom",
]
STEP_TIMER_MODE = {
    "incubation": "countdown", "waiting": "countdown", "centrifuge": "countdown",
    "electrophoresis": "countdown", "gel_running": "countdown",
    "membrane_transfer": "countdown", "imaging": "countdown",
    "heating": "countdown", "cooling": "countdown", "storage": "countdown",
    "note": "none", "checklist_block": "none", "decision": "none",
}

# Human-readable label shown above the large countdown display
TIMER_LABEL = {
    "incubation":        "Incubation timer",
    "waiting":           "Waiting timer",
    "centrifuge":        "Centrifuge timer",
    "electrophoresis":   "Electrophoresis timer",
    "gel_running":       "Gel running timer",
    "membrane_transfer": "Transfer timer",
    "imaging":           "Imaging timer",
    "heating":           "Heating timer",
    "cooling":           "Cooling timer",
    "storage":           "Storage timer",
    "transfer":          "Transfer timer",
}

def _step_timer_type(step):
    """Return 'countdown', 'hands_on_only', or 'none' for a step dict."""
    stype  = step.get("type", "custom")
    wait_m = step.get("waitMinutes", 0)
    ho_m   = step.get("handsOnMinutes", 0)
    if STEP_TIMER_MODE.get(stype) == "none":
        return "none"
    if wait_m > 0 or STEP_TIMER_MODE.get(stype) == "countdown":
        return "countdown"
    if ho_m > 0:
        return "hands_on_only"
    return "none"

FONT = "SF Pro Display"

R_XS = 8
R_SM = 12
R_MD = 16
R_LG = 20
R_XL = 24
R_2XL = 32

AMOUNT_UNITS = [
    "", "µL", "mL", "L", "µg", "mg", "g",
    "cells", "K cells", "M cells",
    "%", "drops", "tablets", "tubes", "wells", "slides",
    "custom...",
]
CONCENTRATION_UNITS = [
    "", "mg/mL", "µg/mL", "ng/mL", "M", "mM", "µM", "nM",
    "%", "X", "OD", "CFU/mL", "custom...",
]

# ─── Helpers ──────────────────────────────────────────────────────────────────
def new_id():
    return str(uuid.uuid4())

def now_ts():
    return int(time.time() * 1000)

def fmt_mins(m):
    m = int(m)
    if m <= 0: return "—"
    if m < 60: return f"{m}m"
    h, r = divmod(m, 60)
    return f"{h}h {r}m" if r else f"{h}h"

def fmt_secs(s):
    m, sec = divmod(int(s), 60)
    return f"{m:02d}:{sec:02d}"

def fmt_date(ts):
    return datetime.fromtimestamp(ts / 1000).strftime("%b %d, %Y")

def fmt_time(ts):
    return datetime.fromtimestamp(ts / 1000).strftime("%I:%M %p")

def total_mins(steps):
    return sum(s.get("handsOnMinutes", 0) + s.get("waitMinutes", 0) + s.get("bufferMinutes", 0) for s in steps)

def is_missing(value):
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, tuple, dict, set)):
        return len(value) == 0
    return False

def display_na(value):
    return "N/A" if is_missing(value) else str(value)

def display_list(values, formatter=None):
    if not values:
        return ["N/A"]
    rows = []
    for value in values:
        if formatter:
            text = formatter(value)
        else:
            text = str(value)
        rows.append(display_na(text))
    return rows or ["N/A"]

def safe_text(value):
    return "" if value is None else str(value)

def fmt_optional_mins(value):
    try:
        mins = float(value or 0)
    except (TypeError, ValueError):
        mins = 0
    return "N/A" if mins <= 0 else fmt_mins(mins)

def reagent_value_with_unit(item, value_key, unit_key):
    value = safe_text(item.get(value_key)).strip()
    unit = safe_text(item.get(unit_key)).strip()
    if value and unit:
        return f"{value} {unit}"
    return value or unit

def reagent_display_text(rg):
    name = safe_text(rg.get("name")).strip()
    amount = reagent_value_with_unit(rg, "amount", "amountUnit")
    concentration = reagent_value_with_unit(rg, "concentration", "concentrationUnit")
    notes = safe_text(rg.get("notes")).strip()
    parts = []
    if name:
        parts.append(name)
    if amount:
        parts.append(amount)
    if concentration:
        parts.append(concentration)
    if notes:
        parts.append(notes)
    return " | ".join(parts)

# ─── Storage ──────────────────────────────────────────────────────────────────
def load_protocols():
    if PROTOCOLS_FILE.exists():
        try: return json.loads(PROTOCOLS_FILE.read_text())
        except: pass
    return []

def save_protocols(protocols):
    PROTOCOLS_FILE.write_text(json.dumps(protocols, ensure_ascii=False, indent=2))

def load_runs():
    if RUNS_FILE.exists():
        try: return json.loads(RUNS_FILE.read_text())
        except: pass
    return []

def save_runs(runs):
    RUNS_FILE.write_text(json.dumps(runs, ensure_ascii=False, indent=2))

def load_terms(path):
    if path.exists():
        try:
            data = json.loads(path.read_text())
            if isinstance(data, list):
                return sorted({str(x).strip() for x in data if str(x).strip()})
        except Exception:
            pass
    return []

def save_terms(path, terms):
    path.write_text(json.dumps(sorted({str(x).strip() for x in terms if str(x).strip()}),
                               ensure_ascii=False, indent=2))

def load_categories():
    return load_terms(CATEGORIES_FILE)

def save_categories(categories):
    save_terms(CATEGORIES_FILE, categories)

def load_tags():
    return load_terms(TAGS_FILE)

def save_tags(tags):
    save_terms(TAGS_FILE, tags)

def _discard_runtime_session():
    """Delete the crash-recovery checkpoint file (session complete or user discarded)."""
    try:
        RUNTIME_FILE.unlink(missing_ok=True)
    except Exception:
        pass

def load_schedule():
    if SCHEDULE_FILE.exists():
        try:
            data = json.loads(SCHEDULE_FILE.read_text())
            return [_migrate_exp_blocks(e) for e in data]
        except:
            pass
    return []

def save_schedule(experiments):
    SCHEDULE_FILE.write_text(json.dumps(experiments, ensure_ascii=False, indent=2))

def load_templates():
    """Return all built-in templates from TEMPLATES_DIR, sorted by category then name.
    Each template is a dict loaded straight from a .json file in that folder.
    Returns an empty list if the folder doesn't exist (safe for packaged builds)."""
    if not TEMPLATES_DIR.exists():
        return []
    templates = []
    for f in sorted(TEMPLATES_DIR.glob("*.json")):
        try:
            t = json.loads(f.read_text(encoding="utf-8"))
            if isinstance(t, dict) and t.get("name"):
                templates.append(t)
        except Exception:
            pass
    templates.sort(key=lambda t: (t.get("category", ""), t.get("name", "")))
    return templates

def _step_sched_cat(stype):
    """Map step type → schedule display category (hands_on/waiting/machine/note)."""
    _HO = {"preparation","reagent_addition","mixing","pipetting","resuspension",
           "transfer","wash","staining","harvest","lysis","sample_collection","blocking"}
    _WT = {"incubation","waiting","heating","cooling","storage"}
    _MC = {"centrifuge","electrophoresis","gel_running","membrane_transfer",
           "imaging","measurement"}
    if stype in _HO: return "hands_on"
    if stype in _WT: return "waiting"
    if stype in _MC: return "machine"
    return "note"


def _migrate_exp_blocks(exp):
    """Ensure exp has timelineBlocks. Migrates from legacy scheduledSteps if needed."""
    if "timelineBlocks" not in exp:
        blocks = []
        for ss in exp.get("scheduledSteps", []):
            blocks.append({
                "id":                   ss.get("id", new_id()),
                "blockType":            "protocol_step",
                "title":                ss.get("title", "Step"),
                "type":                 ss.get("type", "custom"),
                "startTime":            ss.get("plannedStart", 0),
                "endTime":              ss.get("plannedEnd",   0),
                "durationMinutes":      ss.get("durationMinutes", 5),
                "handsOnMinutes":       ss.get("handsOnMinutes", 0),
                "waitMinutes":          ss.get("waitMinutes",    0),
                "notes":                ss.get("notes", ""),
                "status":               "planned",
                "sourceProtocolStepId": ss.get("protocolStepId", ""),
                "isParallelTask":       False,
                "parallelWithBlockId":  None,
                "keepTime":             True,
            })
        exp["timelineBlocks"] = blocks
    # Ensure every block has all required keys (forward-compat)
    for b in exp.get("timelineBlocks", []):
        b.setdefault("id",                  new_id())
        b.setdefault("blockType",           "protocol_step")
        b.setdefault("status",              "planned")
        b.setdefault("isParallelTask",      False)
        b.setdefault("parallelWithBlockId", None)
        b.setdefault("keepTime",            True)
        b.setdefault("durationMinutes",     5)
        b.setdefault("startTime",           exp.get("plannedStart", 0))
        b.setdefault("endTime",             exp.get("plannedStart", 0))
    return exp


def _recalc_timeline(exp):
    """Recalculate startTime/endTime for all blocks from exp.plannedStart.
    Skipped/canceled with keepTime=False take 0 ms (later blocks shift earlier)."""
    blocks = exp.get("timelineBlocks", [])
    cur_ms = exp.get("plannedStart", 0)
    for blk in blocks:
        if blk.get("isParallelTask"):
            continue
        dur_ms = blk.get("durationMinutes", 0) * 60_000
        if blk.get("status") in ("skipped", "canceled") and not blk.get("keepTime", True):
            blk["startTime"] = cur_ms
            blk["endTime"]   = cur_ms
        else:
            blk["startTime"] = cur_ms
            blk["endTime"]   = cur_ms + dur_ms
            cur_ms += dur_ms
    # Update session end from last sequential block
    seq = [b for b in blocks if not b.get("isParallelTask")]
    if seq:
        exp["plannedEnd"] = max(seq[-1]["endTime"], exp.get("plannedStart", 0))
    return exp


# ─── Default step ─────────────────────────────────────────────────────────────
def new_step(order=0):
    return {
        "id": new_id(), "order": order, "title": "", "type": "preparation",
        "description": "", "reagents": [], "equipment": [],
        "handsOnMinutes": 5, "waitMinutes": 0, "bufferMinutes": 0,
        "temperature": "", "centrifugeCondition": "", "shakingRotation": "",
        "checklist": [], "notes": "", "warnings": "", "substeps": [],
    }

def new_source(source_type="manual"):
    return {
        "sourceType": source_type,   # manual / text / pdf / docx / web
        "sourceName": "",
        "sourceUrl": "",
        "importedAt": now_ts(),
        "rawText": "",
        "originalFileName": "",
        "notes": "",
    }

def new_protocol():
    return {
        "id": new_id(), "name": "New Protocol", "category": "",
        "description": "", "createdAt": now_ts(), "updatedAt": now_ts(),
        "tags": [],
        "steps": [new_step(0)],
        "source": new_source("manual"),
    }

# ══════════════════════════════════════════════════════════════════════════════
# Reusable widgets
# ══════════════════════════════════════════════════════════════════════════════
def label(parent, text, size=13, weight="normal", color=None, **kw):
    return ctk.CTkLabel(parent, text=text, font=(FONT, size, weight),
                        text_color=color or T1, **kw)

def btn(parent, text, cmd, color=None, text_color=("#fff","#fff"), width=100, height=34, size=13, **kw):
    return ctk.CTkButton(parent, text=text, command=cmd,
                         fg_color=color or ACC, hover_color=ACCH,
                         text_color=text_color, font=(FONT, size),
                         width=width, height=height,
                         corner_radius=kw.pop("corner_radius", R_SM), **kw)

def entry(parent, placeholder="", width=200, **kw):
    return ctk.CTkEntry(parent, placeholder_text=placeholder,
                        font=(FONT, 13), width=width,
                        height=kw.pop("height", 34),
                        corner_radius=kw.pop("corner_radius", R_SM),
                        border_color=CARD_B, **kw)

def textbox(parent, width=200, height=60, **kw):
    return ctk.CTkTextbox(parent, font=(FONT, 13), width=width, height=height,
                          corner_radius=kw.pop("corner_radius", R_SM),
                          border_color=CARD_B, border_width=1, **kw)

def separator(parent, orient="horizontal"):
    color = CARD_B
    if orient == "horizontal":
        return ctk.CTkFrame(parent, height=1, fg_color=color, corner_radius=0)
    return ctk.CTkFrame(parent, width=1, fg_color=color, corner_radius=0)

def card_frame(parent, **kw):
    return ctk.CTkFrame(parent, fg_color=CARD, corner_radius=kw.pop("corner_radius", R_MD),
                        border_width=1, border_color=CARD_B, **kw)

def canvas_round_rect(canvas, x1, y1, x2, y2, radius=R_MD, **kw):
    r = min(radius, abs(x2 - x1) / 2, abs(y2 - y1) / 2)
    points = [
        x1+r, y1, x2-r, y1, x2, y1, x2, y1+r,
        x2, y2-r, x2, y2, x2-r, y2, x1+r, y2,
        x1, y2, x1, y2-r, x1, y1+r, x1, y1,
    ]
    return canvas.create_polygon(points, smooth=True, splinesteps=12, **kw)

# ══════════════════════════════════════════════════════════════════════════════
# Dialog centering helper
# ══════════════════════════════════════════════════════════════════════════════
def _center_on_parent(dialog, parent):
    """Compute centered position and apply it.
    Call after geometry('WxH') while the window is still withdrawn."""
    dialog.update_idletasks()
    try:
        g = dialog.geometry()          # "WxH" or "WxH+x+y"
        w_str, rest = g.split("x", 1)
        h_str = rest.split("+")[0].split("-")[0]
        dw, dh = int(w_str), int(h_str)
    except Exception:
        dw, dh = 500, 400
    try:
        px  = parent.winfo_rootx()
        py  = parent.winfo_rooty()
        pw  = parent.winfo_width()
        ph  = parent.winfo_height()
        if pw <= 1:
            raise ValueError("parent not yet drawn")
        cx = px + (pw - dw) // 2
        cy = py + (ph - dh) // 2
    except Exception:
        sw = dialog.winfo_screenwidth()
        sh = dialog.winfo_screenheight()
        cx = (sw - dw) // 2
        cy = (sh - dh) // 2
    sw = dialog.winfo_screenwidth()
    sh = dialog.winfo_screenheight()
    cx = max(20, min(cx, sw - dw - 20))
    cy = max(20, min(cy, sh - dh - 20))
    dialog.geometry(f"{dw}x{dh}+{cx}+{cy}")

def _show_dialog(dialog, parent, w, h):
    """Standard open sequence: position BEFORE revealing so there is no jump.
    1. dialog.withdraw() must have been called right after super().__init__()
    2. call this after the UI is fully built."""
    dialog.geometry(f"{w}x{h}")
    dialog.transient(parent)
    _center_on_parent(dialog, parent)
    dialog.deiconify()
    dialog.lift()
    dialog.grab_set()
    dialog.focus_set()

# ══════════════════════════════════════════════════════════════════════════════
# Scrollable frame helper
# ══════════════════════════════════════════════════════════════════════════════
class ScrollFrame(ctk.CTkScrollableFrame):
    def __init__(self, parent, **kw):
        super().__init__(
            parent, fg_color="transparent",
            scrollbar_button_color=("#cbd5e1", "#475569"),
            scrollbar_button_hover_color=("#94a3b8", "#64748b"),
            **kw)
        # ── Unified canvas scroll setup ────────────────────────────────────────
        # CTkScrollableFrame already binds <MouseWheel> to _parent_canvas via
        # _mouse_wheel_all. Our _refresh_scroll_bindings recursively visits all
        # winfo_children(), which includes _parent_canvas, and adds another
        # binding with add="+". This causes TWO handlers to fire on every wheel
        # event over the canvas (empty space between widgets): the first scrolls
        # N units, the second scrolls M more — the combined jump is the bug.
        #
        # Fix: replace CTkScrollableFrame's canvas binding with ours (single
        # consistent handler, same math everywhere), then pre-mark the canvas
        # in _scroll_bound_widgets so the recursive traversal skips it.
        try:
            canvas = self._parent_canvas
            canvas.configure(yscrollincrement=20)
            # Replace, not add-to: no add="+" here.
            canvas.bind("<MouseWheel>", self._on_scroll_input)
            canvas.bind("<Button-4>",   self._on_scroll_input)
            canvas.bind("<Button-5>",   self._on_scroll_input)
            self._scroll_bound_widgets = {str(canvas)}   # pre-mark canvas
        except AttributeError:
            self._scroll_bound_widgets = set()
        self._bind_scroll_inputs(self)
        self.after(250, self._refresh_scroll_bindings)

    def _bind_scroll_inputs(self, widget):
        # Never rebind the canvas — it is set up in __init__ with a single
        # handler that replaces CTkScrollableFrame's built-in one.
        if isinstance(widget, tk.Canvas):
            return
        wid = str(widget)
        if wid in self._scroll_bound_widgets:
            return
        self._scroll_bound_widgets.add(wid)
        for sequence in ("<MouseWheel>", "<Button-4>", "<Button-5>"):
            try:
                widget.bind(sequence, self._on_scroll_input, add="+")
            except (NotImplementedError, tk.TclError):
                pass
        if widget is self:
            for sequence in ("<Up>", "<Down>", "<Prior>", "<Next>", "<Home>", "<End>"):
                try:
                    widget.bind(sequence, self._on_scroll_input, add="+")
                except (NotImplementedError, tk.TclError):
                    pass

    def _refresh_scroll_bindings(self):
        def bind_tree(widget):
            self._bind_scroll_inputs(widget)
            for child in widget.winfo_children():
                bind_tree(child)
        try:
            bind_tree(self)
        except tk.TclError:
            return
        self.after(5000, self._refresh_scroll_bindings)

    def _on_scroll_input(self, event):
        try:
            canvas = self._parent_canvas
        except AttributeError:
            return
        sequence = getattr(event, "keysym", "")
        if sequence == "Up":
            canvas.yview_scroll(-3, "units")
        elif sequence == "Down":
            canvas.yview_scroll(3, "units")
        elif sequence == "Prior":
            canvas.yview_scroll(-1, "pages")
        elif sequence == "Next":
            canvas.yview_scroll(1, "pages")
        elif sequence == "Home":
            canvas.yview_moveto(0)
        elif sequence == "End":
            canvas.yview_moveto(1)
        elif getattr(event, "num", None) == 4:
            canvas.yview_scroll(-5, "units")   # Linux scroll up  = 100 px
        elif getattr(event, "num", None) == 5:
            canvas.yview_scroll(5, "units")    # Linux scroll down = 100 px
        elif getattr(event, "delta", 0):
            delta = event.delta
            if abs(delta) >= 120:
                # Physical mouse wheel: ±120 per notch on macOS/Windows.
                # 5 units × 20 px = 100 px per notch — fast, comfortable.
                units = int(-delta / 120) * 5
            else:
                # Trackpad: delta is typically ±1–30 on macOS.
                # Divide by 3 to stay proportional; floor to ±1 minimum so
                # tiny swipes still register. This matches approximately the
                # speed of the native CTkScrollableFrame handler while keeping
                # a single code path for canvas and child widgets alike.
                units = int(-delta / 3) or (-1 if delta > 0 else 1)
            canvas.yview_scroll(units, "units")
        return "break"

# ══════════════════════════════════════════════════════════════════════════════
# Step editor dialog
# ══════════════════════════════════════════════════════════════════════════════
class StepEditorDialog(ctk.CTkToplevel):
    def __init__(self, parent, step: dict, on_save):
        super().__init__(parent)
        self.step = dict(step)
        self.step["reagents"] = [dict(r) for r in step.get("reagents", [])]
        self.step["checklist"] = [dict(c) for c in step.get("checklist", [])]
        self.step["substeps"] = [dict(s) for s in step.get("substeps", [])]
        for reagent in self.step["reagents"]:
            reagent.setdefault("amountUnit", "")
            reagent.setdefault("concentrationUnit", "")
            if "unit" in reagent and not reagent.get("amountUnit"):
                reagent["amountUnit"] = reagent.get("unit", "")
        self._original_step = self._comparison_step(self.step)
        if not self.step["reagents"]:
            self.step["reagents"].append({
                "id": new_id(), "name": "", "amount": "", "amountUnit": "",
                "concentration": "", "concentrationUnit": "", "notes": ""
            })
        if not self.step["checklist"]:
            self.step["checklist"].append({"id": new_id(), "text": "", "checked": False})
        if not self.step["substeps"]:
            self.step["substeps"].append({"id": new_id(), "text": ""})
        self.on_save = on_save
        self.withdraw()
        self.title("Edit Step")
        self.resizable(True, True)
        self.protocol("WM_DELETE_WINDOW", self._back)
        self._build()
        max_w = max(900, int(self.winfo_screenwidth() * 0.86))
        max_h = max(640, int(self.winfo_screenheight() * 0.9))
        _show_dialog(self, parent, min(1040, max_w), min(840, max_h))

    def _section_card(self, parent, title, row):
        card = card_frame(parent)
        card.grid(row=row, column=0, sticky="ew", pady=(0, 12))
        card.grid_columnconfigure(0, weight=1)
        hdr = ctk.CTkFrame(card, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=16, pady=(14, 8))
        label(hdr, title, size=14, weight="bold").pack(side="left")
        body = ctk.CTkFrame(card, fg_color="transparent")
        body.grid(row=1, column=0, sticky="ew", padx=16, pady=(0, 16))
        return body

    def _field_entry(self, parent, row, col, title, attr, placeholder, width=180):
        box = ctk.CTkFrame(parent, fg_color="transparent")
        box.grid(row=row, column=col, sticky="ew", padx=(0, 10), pady=(0, 10))
        box.grid_columnconfigure(0, weight=1)
        label(box, title, size=11, color=T2, weight="bold").grid(row=0, column=0, sticky="w", pady=(0, 3))
        e = entry(box, placeholder, width=width)
        e.grid(row=1, column=0, sticky="ew")
        setattr(self, attr, e)
        return e

    def _field_textbox(self, parent, row, col, title, attr, placeholder="", width=320, height=72, colspan=1):
        box = ctk.CTkFrame(parent, fg_color="transparent")
        box.grid(row=row, column=col, columnspan=colspan, sticky="ew", padx=(0, 10), pady=(0, 10))
        box.grid_columnconfigure(0, weight=1)
        label(box, title, size=11, color=T2, weight="bold").grid(row=0, column=0, sticky="w", pady=(0, 3))
        tb = textbox(box, width=width, height=height)
        if placeholder:
            tb.configure(border_color=CARD_B)
        tb.grid(row=1, column=0, sticky="ew")
        setattr(self, attr, tb)
        return tb

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        header = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG, border_width=0)
        header.grid(row=0, column=0, sticky="ew")
        header.grid_columnconfigure(1, weight=1)
        btn(header, "← Back", self._back, color=("#e2e8f0","#334155"),
            text_color=T1, width=90, height=34).grid(row=0, column=0, padx=(16, 10), pady=12)
        label(header, "Edit Step", size=18, weight="bold").grid(row=0, column=1, sticky="w")
        btn(header, "Save Step", self._save, width=120, height=34).grid(
            row=0, column=2, padx=(10, 16), pady=12)

        scroll = ScrollFrame(self)
        scroll.grid(row=1, column=0, sticky="nsew", padx=20, pady=(16, 0))
        scroll.grid_columnconfigure(0, weight=1)
        body_row = 0

        basic = self._section_card(scroll, "Basic Information", body_row); body_row += 1
        basic.grid_columnconfigure((0, 1), weight=1)
        self.e_title = self._field_entry(basic, 0, 0, "Step Title", "e_title", "Step title...", width=360)
        self.e_title.insert(0, self.step.get("title", ""))
        type_box = ctk.CTkFrame(basic, fg_color="transparent")
        type_box.grid(row=0, column=1, sticky="ew", padx=(0, 10), pady=(0, 10))
        label(type_box, "Step Type", size=11, color=T2, weight="bold").grid(row=0, column=0, sticky="w", pady=(0, 3))
        self.type_var = ctk.StringVar(value=self.step.get("type", "preparation"))
        type_menu = ctk.CTkOptionMenu(type_box, variable=self.type_var, values=STEP_TYPES,
                                      font=(FONT, 13), width=220, height=34,
                                      corner_radius=R_SM)
        type_menu.grid(row=1, column=0, sticky="w")
        self.e_desc = self._field_textbox(basic, 1, 0, "Description", "e_desc", width=760, height=86, colspan=2)
        self.e_desc.insert("0.0", self.step.get("description", ""))

        timing = self._section_card(scroll, "Timing & Conditions", body_row); body_row += 1
        timing.grid_columnconfigure((0, 1, 2), weight=1)
        for col, (lbl, key) in enumerate([("Hands-on time (min)", "handsOnMinutes"),
                                          ("Wait / Incubation time (min)", "waitMinutes"),
                                          ("Buffer time (min)", "bufferMinutes")]):
            e = self._field_entry(timing, 0, col, lbl, f"e_{key}", "0", width=120)
            e.insert(0, str(self.step.get(key, 0)))
        self.e_temperature = self._field_entry(timing, 1, 0, "Temperature", "e_temperature", "e.g. 4°C, RT, 37°C, on ice", width=220)
        self.e_temperature.insert(0, self.step.get("temperature", ""))
        self.e_centrifugeCondition = self._field_entry(timing, 1, 1, "Centrifuge", "e_centrifugeCondition", "e.g. 300g × 5 min", width=220)
        self.e_centrifugeCondition.insert(0, self.step.get("centrifugeCondition", ""))
        self.e_shakingRotation = self._field_entry(timing, 1, 2, "Rotation / Shaking", "e_shakingRotation", "e.g. 250 rpm", width=180)
        self.e_shakingRotation.insert(0, self.step.get("shakingRotation", ""))

        materials = self._section_card(scroll, "Equipment & Reagents", body_row); body_row += 1
        materials.grid_columnconfigure(0, weight=1)
        self.e_equip = self._field_entry(materials, 0, 0, "Equipment", "e_equip", "comma separated, e.g. Centrifuge, pipette, ice bucket", width=760)
        self.e_equip.insert(0, ", ".join(self.step.get("equipment", [])))
        self.reagent_frame = ctk.CTkFrame(materials, fg_color=("#f8fafc","#1e293b"),
                                          corner_radius=R_SM, border_width=1, border_color=CARD_B)
        self.reagent_frame.grid(row=1, column=0, sticky="ew", pady=(2, 8))
        self._render_reagents()
        btn(materials, "+ Add Reagent", self._add_reagent, color=("#e2e8f0","#334155"),
            text_color=T1, width=130, height=30).grid(row=2, column=0, sticky="w")

        workflow = self._section_card(scroll, "Workflow Details", body_row); body_row += 1
        workflow.grid_columnconfigure((0, 1), weight=1)
        checklist_col = ctk.CTkFrame(workflow, fg_color="transparent")
        checklist_col.grid(row=0, column=0, sticky="nsew", padx=(0, 12))
        label(checklist_col, "Checklist", size=11, color=T2, weight="bold").pack(anchor="w", pady=(0, 4))
        self.checklist_frame = ctk.CTkFrame(checklist_col, fg_color="transparent")
        self.checklist_frame.pack(fill="x")
        self._render_checklist()
        btn(checklist_col, "+ Add Item", self._add_checklist, color=("#e2e8f0","#334155"),
            text_color=T1, width=110, height=28).pack(anchor="w", pady=(4, 0))

        substep_col = ctk.CTkFrame(workflow, fg_color="transparent")
        substep_col.grid(row=0, column=1, sticky="nsew")
        label(substep_col, "Sub-steps", size=11, color=T2, weight="bold").pack(anchor="w", pady=(0, 4))
        self.substep_frame = ctk.CTkFrame(substep_col, fg_color="transparent")
        self.substep_frame.pack(fill="x")
        self._render_substeps()
        btn(substep_col, "+ Add Sub-step", self._add_substep, color=("#e2e8f0","#334155"),
            text_color=T1, width=125, height=28).pack(anchor="w", pady=(4, 0))

        notes = self._section_card(scroll, "Notes & Warnings", body_row); body_row += 1
        notes.grid_columnconfigure((0, 1), weight=1)
        self.tb_notes = self._field_textbox(notes, 0, 0, "Notes", "tb_notes", width=360, height=82)
        self.tb_notes.insert("0.0", self.step.get("notes", ""))
        self.tb_warnings = self._field_textbox(notes, 0, 1, "Warnings", "tb_warnings", width=360, height=82)
        self.tb_warnings.insert("0.0", self.step.get("warnings", ""))

        ctk.CTkFrame(scroll, height=10, fg_color="transparent").grid(row=body_row, column=0, sticky="ew")

        footer = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG, border_width=0)
        footer.grid(row=2, column=0, sticky="ew")
        footer.grid_columnconfigure(0, weight=1)
        btn(footer, "Cancel", self._back, color=("#e2e8f0","#334155"),
            text_color=T1, width=100).grid(row=0, column=0, sticky="w", padx=16, pady=12)
        btn(footer, "Save Step", self._save, width=120).grid(row=0, column=1, sticky="e", padx=16, pady=12)

    def _render_reagents(self):
        for w in self.reagent_frame.winfo_children(): w.destroy()
        widths = [160, 88, 96, 110, 96, 190, 36]
        headers = ["Name", "Amount", "Unit", "Concentration", "Unit", "Notes", ""]
        for col, width in enumerate(widths):
            self.reagent_frame.grid_columnconfigure(col, minsize=width, weight=1 if col in (0, 5) else 0)
        for col, (txt, width) in enumerate(zip(headers, widths)):
            label(self.reagent_frame, txt, size=10, color=T2, weight="bold").grid(
                row=0, column=col, sticky="w", padx=(10 if col == 0 else 4, 4), pady=(8, 4))
        for i, r in enumerate(self.step["reagents"]):
            grid_row = i + 1
            e_name = entry(self.reagent_frame, "Competent cells", width=widths[0])
            e_name.insert(0, r.get("name", ""))
            e_name.grid(row=grid_row, column=0, sticky="ew", padx=(10, 4), pady=(0, 6))
            e_name.bind("<KeyRelease>", lambda ev, idx=i, widget=e_name: self._update_reagent(idx, "name", widget.get()))

            self._reagent_value_cell(self.reagent_frame, grid_row, 1, i, "amount", "50", widths[1])
            self._reagent_unit_cell(self.reagent_frame, grid_row, 2, i, "amountUnit", AMOUNT_UNITS, widths[2])
            self._reagent_value_cell(self.reagent_frame, grid_row, 3, i, "concentration", "1:1000", widths[3])
            self._reagent_unit_cell(self.reagent_frame, grid_row, 4, i, "concentrationUnit", CONCENTRATION_UNITS, widths[4])

            e_notes = entry(self.reagent_frame, "keep on ice", width=widths[5])
            e_notes.insert(0, r.get("notes", ""))
            e_notes.grid(row=grid_row, column=5, sticky="ew", padx=4, pady=(0, 6))
            e_notes.bind("<KeyRelease>", lambda ev, idx=i, widget=e_notes: self._update_reagent(idx, "notes", widget.get()))
            def del_r(idx=i):
                self.step["reagents"].pop(idx)
                if not self.step["reagents"]:
                    self.step["reagents"].append({
                        "id": new_id(), "name": "", "amount": "", "amountUnit": "",
                        "concentration": "", "concentrationUnit": "", "notes": ""
                    })
                self._render_reagents()
            ctk.CTkButton(self.reagent_frame, text="✕", width=30, height=30,
                          fg_color=DANGER, hover_color=("#dc2626","#dc2626"),
                          font=(FONT,12), command=del_r, corner_radius=R_XS).grid(
                              row=grid_row, column=6, sticky="e", padx=(4, 10), pady=(0, 6))
        ctk.CTkFrame(self.reagent_frame, height=4, fg_color="transparent").grid(
            row=len(self.step["reagents"]) + 1, column=0, columnspan=7, sticky="ew")

    def _reagent_value_cell(self, parent, row, col, idx, key, placeholder, width):
        e_val = entry(parent, placeholder, width=width)
        e_val.insert(0, self.step["reagents"][idx].get(key, ""))
        e_val.grid(row=row, column=col, sticky="ew", padx=4, pady=(0, 6))
        e_val.bind("<KeyRelease>", lambda ev, i=idx, k=key, widget=e_val: self._update_reagent(i, k, widget.get()))

    def _reagent_unit_cell(self, parent, row, col, idx, unit_key, units, width):
        current_unit = self.step["reagents"][idx].get(unit_key, "")
        if current_unit and current_unit not in units:
            e_unit = entry(parent, "unit", width=width)
            e_unit.insert(0, current_unit)
            e_unit.grid(row=row, column=col, sticky="ew", padx=4, pady=(0, 6))
            e_unit.bind("<KeyRelease>", lambda ev, i=idx, k=unit_key, widget=e_unit: self._update_reagent(i, k, widget.get()))
            return
        unit_var = ctk.StringVar(value=current_unit if current_unit in units else "")
        unit_menu = ctk.CTkOptionMenu(
            parent, variable=unit_var, values=units, width=width, height=32,
            font=(FONT, 11), corner_radius=R_SM,
            command=lambda choice, i=idx, k=unit_key: self._set_reagent_unit(i, k, choice))
        unit_menu.grid(row=row, column=col, sticky="ew", padx=4, pady=(0, 6))

    def _set_reagent_unit(self, idx, key, choice):
        if idx >= len(self.step["reagents"]):
            return
        if choice == "custom...":
            self.step["reagents"][idx][key] = ""
        else:
            self.step["reagents"][idx][key] = choice
        self._render_reagents()

    def _update_reagent(self, idx, key, val):
        if idx < len(self.step["reagents"]):
            self.step["reagents"][idx][key] = val

    def _add_reagent(self):
        self.step["reagents"].append({
            "id": new_id(), "name": "", "amount": "", "amountUnit": "",
            "concentration": "", "concentrationUnit": "", "notes": ""
        })
        self._render_reagents()

    def _render_checklist(self):
        for w in self.checklist_frame.winfo_children(): w.destroy()
        for i, c in enumerate(self.step["checklist"]):
            row_f = ctk.CTkFrame(self.checklist_frame, fg_color="transparent")
            row_f.pack(fill="x", pady=2)
            e = entry(row_f, "Checklist item...", width=560)
            e.insert(0, c.get("text", ""))
            e.pack(side="left", padx=(0,4))
            e.bind("<KeyRelease>", lambda ev, idx=i, widget=e: self._update_cl(idx, widget.get()))
            def del_c(idx=i):
                self.step["checklist"].pop(idx)
                if not self.step["checklist"]:
                    self.step["checklist"].append({"id": new_id(), "text": "", "checked": False})
                self._render_checklist()
            ctk.CTkButton(row_f, text="✕", width=28, height=28, fg_color=DANGER, font=(FONT,12),
                          command=del_c, corner_radius=R_XS).pack(side="left")

    def _update_cl(self, idx, val):
        if idx < len(self.step["checklist"]):
            self.step["checklist"][idx]["text"] = val

    def _add_checklist(self):
        self.step["checklist"].append({"id": new_id(), "text": "", "checked": False})
        self._render_checklist()

    def _render_substeps(self):
        for w in self.substep_frame.winfo_children(): w.destroy()
        for i, s in enumerate(self.step["substeps"]):
            row_f = ctk.CTkFrame(self.substep_frame, fg_color="transparent")
            row_f.pack(fill="x", pady=2)
            label(row_f, f"{i+1}.", size=12, color=T2).pack(side="left", padx=(0,4))
            e = entry(row_f, "Sub-step description...", width=556)
            e.insert(0, s.get("text", ""))
            e.pack(side="left", padx=(0,4))
            e.bind("<KeyRelease>", lambda ev, idx=i, widget=e: self._update_ss(idx, widget.get()))
            def del_s(idx=i):
                self.step["substeps"].pop(idx)
                if not self.step["substeps"]:
                    self.step["substeps"].append({"id": new_id(), "text": ""})
                self._render_substeps()
            ctk.CTkButton(row_f, text="✕", width=28, height=28, fg_color=DANGER, font=(FONT,12),
                          command=del_s, corner_radius=R_XS).pack(side="left")

    def _update_ss(self, idx, val):
        if idx < len(self.step["substeps"]):
            self.step["substeps"][idx]["text"] = val

    def _add_substep(self):
        self.step["substeps"].append({"id": new_id(), "text": ""})
        self._render_substeps()

    def _comparison_step(self, step):
        clean = dict(step)
        for key in ["description", "temperature", "centrifugeCondition", "shakingRotation",
                    "warnings", "notes", "category"]:
            clean.setdefault(key, "")
        for key in ["handsOnMinutes", "waitMinutes", "bufferMinutes"]:
            clean.setdefault(key, 0)
        clean.setdefault("equipment", [])
        clean["reagents"] = [
            r for r in clean.get("reagents", [])
            if any(str(r.get(k, "")).strip() for k in (
                "name", "amount", "amountUnit", "concentration", "concentrationUnit", "notes"))
        ]
        clean["checklist"] = [
            c for c in clean.get("checklist", [])
            if str(c.get("text", "")).strip()
        ]
        clean["substeps"] = [
            s for s in clean.get("substeps", [])
            if str(s.get("text", "")).strip()
        ]
        return clean

    def _draft_step(self):
        draft = dict(self.step)
        draft["title"] = self.e_title.get().strip()
        draft["type"] = self.type_var.get()
        draft["description"] = self.e_desc.get("0.0", "end").strip()
        for key in ["handsOnMinutes", "waitMinutes", "bufferMinutes"]:
            try:
                draft[key] = int(getattr(self, f"e_{key}").get())
            except Exception:
                draft[key] = 0
        draft["temperature"] = self.e_temperature.get().strip()
        draft["centrifugeCondition"] = self.e_centrifugeCondition.get().strip()
        draft["shakingRotation"] = self.e_shakingRotation.get().strip()
        eq = self.e_equip.get().strip()
        draft["equipment"] = [x.strip() for x in eq.split(",") if x.strip()] if eq else []
        draft["warnings"] = self.tb_warnings.get("0.0", "end").strip()
        draft["notes"] = self.tb_notes.get("0.0", "end").strip()
        draft["reagents"] = [
            r for r in self.step["reagents"]
            if any(str(r.get(k, "")).strip() for k in (
                "name", "amount", "amountUnit", "concentration", "concentrationUnit", "notes"))
        ]
        draft["checklist"] = [
            c for c in self.step["checklist"]
            if str(c.get("text", "")).strip()
        ]
        draft["substeps"] = [
            s for s in self.step["substeps"]
            if str(s.get("text", "")).strip()
        ]
        return draft

    def _has_unsaved_changes(self):
        return json.dumps(self._original_step, sort_keys=True, ensure_ascii=False) != json.dumps(
            self._comparison_step(self._draft_step()), sort_keys=True, ensure_ascii=False)

    def _back(self):
        if self._has_unsaved_changes():
            if not messagebox.askyesno("Discard unsaved changes?", "Discard unsaved changes?"):
                return
        self.destroy()

    def _save(self):
        self.step = self._draft_step()
        self.on_save(self.step)
        self.destroy()

# ══════════════════════════════════════════════════════════════════════════════
# Pages
# ══════════════════════════════════════════════════════════════════════════════
class PageBase(ctk.CTkFrame):
    def __init__(self, parent, app):
        super().__init__(parent, fg_color=BG, corner_radius=R_XL)
        self.app = app

    def refresh(self): pass


class ProtocolCreateDialog(ctk.CTkToplevel):
    # (icon, title, description, action_key)
    _ALL_OPTIONS = [
        ("✦", "Blank Protocol",           "Create a protocol manually from scratch.",            "blank"),
        ("⌨", "Import from Text / Web",   "Paste text or a URL and extract steps manually.",     "text"),
        ("⎘", "Import PDF",               "Extract protocol text from a local PDF file.",         "pdf"),
        ("◫", "Import Word (.docx)",      "Extract protocol text from a Word document.",          "docx"),
        ("⧉", "Duplicate Existing",       "Copy an existing protocol as a new editable template.", "dup"),
    ]

    def __init__(self, parent, app, mode="full"):
        super().__init__(parent)
        self.withdraw()
        self.app  = app
        self.mode = mode
        self.title("Create Protocol")
        self.resizable(False, False)
        self.configure(fg_color=BG)
        self._build()
        _show_dialog(self, parent, 760, 580)

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # ── header (fixed) ────────────────────────────────────────────────────
        hdr = ctk.CTkFrame(self, fg_color=CARD, corner_radius=0, border_width=0)
        hdr.grid(row=0, column=0, sticky="ew")
        hdr.grid_columnconfigure(0, weight=1)
        inner_hdr = ctk.CTkFrame(hdr, fg_color="transparent")
        inner_hdr.pack(fill="x", padx=28, pady=(22, 18))
        label(inner_hdr, "Create Protocol", size=20, weight="bold").pack(anchor="w")
        label(inner_hdr, "Choose how you want to create a reusable wet-lab protocol.",
              size=13, color=T2).pack(anchor="w", pady=(3, 0))
        separator(self).grid(row=0, column=0, sticky="ews")

        # ── scrollable options ────────────────────────────────────────────────
        body = ScrollFrame(self)
        body.grid(row=1, column=0, sticky="nsew", padx=28, pady=(16, 0))
        body.grid_columnconfigure(0, weight=1)

        options = []
        if self.mode == "full":
            options.append(self._ALL_OPTIONS[0])
        options += self._ALL_OPTIONS[1:4]
        if self.mode == "full":
            options.append(self._ALL_OPTIONS[4])

        for row_i, (icon, title, desc, key) in enumerate(options):
            self._option_card(body, row_i, icon, title, desc, key)

        # ── footer (fixed) ────────────────────────────────────────────────────
        footer = ctk.CTkFrame(self, fg_color="transparent")
        footer.grid(row=2, column=0, sticky="ew", padx=28, pady=(8, 16))
        btn(footer, "Cancel", self.destroy,
            color=("#e2e8f0", "#334155"), text_color=T1,
            width=90, height=34).pack(anchor="w")

    def _option_card(self, parent, row_i, icon, title, desc, key):
        card = ctk.CTkFrame(parent, fg_color=CARD, corner_radius=R_MD,
                            border_width=1, border_color=CARD_B, cursor="hand2")
        card.grid(row=row_i, column=0, sticky="ew", pady=(0, 8))
        card.grid_columnconfigure(1, weight=1)

        # icon badge
        icon_f = ctk.CTkFrame(card, fg_color=("#eff6ff", "#1e3a5f"),
                               corner_radius=R_SM, width=44, height=44)
        icon_f.grid(row=0, column=0, padx=(16, 12), pady=14)
        icon_f.grid_propagate(False)
        label(icon_f, icon, size=18, color=ACC).place(relx=.5, rely=.5, anchor="center")

        # text
        txt_f = ctk.CTkFrame(card, fg_color="transparent")
        txt_f.grid(row=0, column=1, sticky="ew", pady=14)
        label(txt_f, title, size=14, weight="bold").pack(anchor="w")
        label(txt_f, desc, size=12, color=T2).pack(anchor="w", pady=(2, 0))

        # chevron
        label(card, "›", size=18, color=T3).grid(row=0, column=2, padx=(8, 16))

        # hover effect — bind to all children
        def _enter(e, c=card):
            c.configure(fg_color=("#f0f7ff", "#1e3a5f"), border_color=ACC)
        def _leave(e, c=card):
            c.configure(fg_color=CARD, border_color=CARD_B)
        def _click(e=None, k=key):
            self._choose(k)

        for w in (card, icon_f, txt_f) + tuple(card.winfo_children()) + tuple(txt_f.winfo_children()):
            try:
                w.bind("<Enter>", _enter, add="+")
                w.bind("<Leave>", _leave, add="+")
                w.bind("<Button-1>", _click, add="+")
            except Exception:
                pass

    def _choose(self, key):
        self.destroy()
        if key == "blank":
            self.app.open_editor(None)
        elif key == "text":
            self.app.open_import("text")
        elif key == "pdf":
            self.app.open_import("pdf")
        elif key == "docx":
            self.app.open_import("docx")
        elif key == "dup":
            self._duplicate_existing()

    def _duplicate_existing(self):
        protocols = self.app.protocols
        if not protocols:
            messagebox.showinfo("No Protocols", "There are no protocols to duplicate yet.")
            return
        DuplicateProtocolDialog(self.app, self.app)


class DuplicateProtocolDialog(ctk.CTkToplevel):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app
        self.withdraw()
        self.title("Duplicate Protocol")
        self.resizable(False, False)
        self._name_to_id = {p.get("name", "Untitled"): p["id"] for p in app.protocols}
        self._var = ctk.StringVar(value=next(iter(self._name_to_id.keys())))
        self._build()
        _show_dialog(self, parent, 460, 220)

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=18)
        label(f, "Duplicate Existing Protocol", size=16, weight="bold").pack(anchor="w")
        label(f, "Select a protocol to copy into a new editable template.", size=12, color=T2).pack(anchor="w", pady=(2, 12))
        ctk.CTkOptionMenu(f, variable=self._var, values=list(self._name_to_id.keys()),
                          width=390, height=36, font=(FONT, 13),
                          corner_radius=R_SM).pack(anchor="w")
        row = ctk.CTkFrame(f, fg_color="transparent")
        row.pack(fill="x", pady=(18, 0))
        btn(row, "Cancel", self.destroy, color=("#e2e8f0","#334155"), text_color=T1, width=90).pack(side="left")
        btn(row, "Duplicate", self._duplicate, width=120).pack(side="right")

    def _duplicate(self):
        import copy
        pid = self._name_to_id.get(self._var.get())
        orig = next((p for p in self.app.protocols if p["id"] == pid), None)
        if not orig:
            return
        dup = copy.deepcopy(orig)
        dup["id"] = new_id()
        dup["name"] = orig.get("name", "Untitled") + " (copy)"
        dup["createdAt"] = dup["updatedAt"] = now_ts()
        for s in dup.get("steps", []):
            s["id"] = new_id()
            for c in s.get("checklist", []):
                c["id"] = new_id()
        self.app.protocols.insert(0, dup)
        save_protocols(self.app.protocols)
        self.destroy()
        self.app.open_editor(dup["id"])


class TermManagerDialog(ctk.CTkToplevel):
    def __init__(self, parent, app, kind):
        super().__init__(parent)
        self.app = app
        self.kind = kind
        self.withdraw()
        self.title(f"Manage {kind.title()}")
        self.resizable(False, False)
        self._build()
        _show_dialog(self, parent, 440, 420)

    @property
    def _terms(self):
        return self.app.categories if self.kind == "categories" else self.app.tags

    @_terms.setter
    def _terms(self, values):
        if self.kind == "categories":
            self.app.categories = sorted(values)
            save_categories(self.app.categories)
        else:
            self.app.tags = sorted(values)
            save_tags(self.app.tags)

    def _build(self):
        f = ctk.CTkFrame(self, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=18)
        label(f, f"Manage {self.kind.title()}", size=18, weight="bold").pack(anchor="w")
        self.list_frame = ctk.CTkScrollableFrame(f, fg_color=CARD, corner_radius=R_SM,
                                                 border_width=1, border_color=CARD_B, height=230)
        self.list_frame.pack(fill="both", expand=True, pady=(12, 10))
        add_row = ctk.CTkFrame(f, fg_color="transparent")
        add_row.pack(fill="x")
        self.e_new = entry(add_row, f"New {self.kind[:-1]}...", width=280)
        self.e_new.pack(side="left", fill="x", expand=True, padx=(0,8))
        btn(add_row, "Add", self._add, width=80).pack(side="left")
        btn(f, "Done", self._done, width=90).pack(anchor="e", pady=(14, 0))
        self._render()

    def _render(self):
        for w in self.list_frame.winfo_children():
            w.destroy()
        if not self._terms:
            label(self.list_frame, f"No {self.kind} yet.", size=12, color=T3).pack(pady=30)
            return
        for term in self._terms:
            row = ctk.CTkFrame(self.list_frame, fg_color="transparent")
            row.pack(fill="x", padx=10, pady=4)
            e = entry(row, width=250)
            e.insert(0, term)
            e.pack(side="left", fill="x", expand=True, padx=(0, 8))
            btn(row, "Save", lambda old=term, widget=e: self._rename(old, widget.get()),
                width=60, height=28, size=11).pack(side="left", padx=(0,4))
            ctk.CTkButton(row, text="✕", width=28, height=28, fg_color=DANGER,
                          command=lambda t=term: self._delete(t), corner_radius=R_XS).pack(side="left")

    def _add(self):
        term = self.e_new.get().strip()
        if not term:
            return
        self._terms = set(self._terms) | {term}
        self.e_new.delete(0, "end")
        self._render()

    def _rename(self, old, new):
        new = new.strip()
        if not new:
            return
        terms = {new if t == old else t for t in self._terms}
        self._terms = terms
        if self.kind == "categories":
            for p in self.app.protocols:
                if p.get("category") == old:
                    p["category"] = new
            save_protocols(self.app.protocols)
        else:
            for p in self.app.protocols:
                p["tags"] = [new if t == old else t for t in p.get("tags", [])]
            save_protocols(self.app.protocols)
        self._render()

    def _delete(self, term):
        if not messagebox.askyesno("Delete", f"Delete '{term}'? Existing protocols will keep their text values."):
            return
        self._terms = {t for t in self._terms if t != term}
        self._render()

    def _done(self):
        try:
            self.app.pages["library"].refresh()
            self.app.pages["editor"].refresh_terms()
        except Exception:
            pass
        self.destroy()

# ─── Dashboard ────────────────────────────────────────────────────────────────
class DashboardPage(PageBase):
    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._build()

    def _build(self):
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(28,16))
        label(hdr, "Dashboard", size=24, weight="bold").pack(anchor="w")
        label(hdr, "Welcome to BenchFlow", size=13, color=T2).pack(anchor="w")

        self.scroll = ScrollFrame(self)
        self.scroll.grid(row=1, column=0, sticky="nsew", padx=28, pady=(0,28))
        self.scroll.grid_columnconfigure((0,1,2,3), weight=1)

        self.stat_frames = []
        for i in range(4):
            f = card_frame(self.scroll)
            f.grid(row=0, column=i, sticky="ew", padx=(0 if i==0 else 8, 0), pady=(0,16))
            self.stat_frames.append(f)

        self.recent_proto_frame = card_frame(self.scroll)
        self.recent_proto_frame.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=(0,8), pady=(0,16))
        self.recent_runs_frame = card_frame(self.scroll)
        self.recent_runs_frame.grid(row=1, column=2, columnspan=2, sticky="nsew", pady=(0,16))

    def refresh(self):
        protocols = self.app.protocols
        runs = self.app.runs

        stats = [
            ("Protocols", len(protocols), ACC),
            ("Total Steps", sum(len(p.get("steps",[])) for p in protocols), ("#7c3aed","#7c3aed")),
            ("Runs Completed", sum(1 for r in runs if r.get("endedAt")), GREEN),
            ("Hands-on Total", fmt_mins(sum(s.get("handsOnMinutes",0) for p in protocols for s in p.get("steps",[]))), ORANGE),
        ]
        for f, (lbl_txt, val, col) in zip(self.stat_frames, stats):
            for w in f.winfo_children(): w.destroy()
            f.configure(width=160)
            inner = ctk.CTkFrame(f, fg_color="transparent")
            inner.pack(padx=16, pady=14, fill="x")
            ctk.CTkFrame(inner, width=8, height=8, fg_color=col, corner_radius=R_XS).pack(anchor="w")
            label(inner, str(val), size=26, weight="bold").pack(anchor="w", pady=(4,0))
            label(inner, lbl_txt, size=12, color=T2).pack(anchor="w")

        # Recent protocols
        f = self.recent_proto_frame
        for w in f.winfo_children(): w.destroy()
        hdr = ctk.CTkFrame(f, fg_color="transparent")
        hdr.pack(fill="x", padx=16, pady=(14,8))
        label(hdr, "Recent Protocols", size=14, weight="bold").pack(side="left")
        ctk.CTkButton(hdr, text="View all →", width=80, height=26, fg_color="transparent",
                      text_color=ACC, font=(FONT,12), hover_color=CARD_B,
                      command=lambda: self.app.navigate("library")).pack(side="right")
        recent = protocols[:5]
        if not recent:
            label(f, "No protocols yet", size=13, color=T3).pack(pady=20)
        for p in recent:
            pf = ctk.CTkFrame(f, fg_color="transparent", cursor="hand2")
            pf.pack(fill="x", padx=12, pady=2)
            pf.bind("<Button-1>", lambda e, pid=p["id"]: self.app.open_editor(pid))
            label(pf, p["name"] or "Untitled", size=13, weight="bold").pack(side="left", padx=4)
            steps = p.get("steps", [])
            label(pf, f"{len(steps)} steps · {fmt_mins(total_mins(steps))}", size=11, color=T2).pack(side="left", padx=4)
            label(pf, fmt_date(p["updatedAt"]), size=11, color=T3).pack(side="right", padx=4)
        btn(f, "+ New Protocol", lambda: self.app.open_create_protocol("full"),
            width=130, height=30).pack(padx=16, pady=10, anchor="w")

        # Recent runs
        f = self.recent_runs_frame
        for w in f.winfo_children(): w.destroy()
        hdr2 = ctk.CTkFrame(f, fg_color="transparent")
        hdr2.pack(fill="x", padx=16, pady=(14,8))
        label(hdr2, "Recent Sessions", size=14, weight="bold").pack(side="left")
        ctk.CTkButton(hdr2, text="Lab Notebook →", width=110, height=26, fg_color="transparent",
                      text_color=ACC, font=(FONT,12), hover_color=CARD_B,
                      command=lambda: self.app.navigate("history")).pack(side="right")
        recent_r = [r for r in runs if r.get("endedAt")][:5]
        if not recent_r:
            label(f, "No runs yet", size=13, color=T3).pack(pady=20)
        for r in recent_r:
            rf = ctk.CTkFrame(f, fg_color=("#f8fafc","#1e293b"), corner_radius=R_SM)
            rf.pack(fill="x", padx=12, pady=3)
            label(rf, r.get("protocolName","?"), size=13, weight="bold").pack(anchor="w", padx=10, pady=(8,2))
            row_ = ctk.CTkFrame(rf, fg_color="transparent")
            row_.pack(fill="x", padx=10, pady=(0,8))
            label(row_, fmt_date(r["startedAt"]), size=11, color=T2).pack(side="left")
            if r.get("actualDuration") is not None:
                label(row_, f"Actual: {fmt_mins(r['actualDuration'])}", size=11, color=GREEN).pack(side="left", padx=8)

# ─── Library ──────────────────────────────────────────────────────────────────
class LibraryPage(PageBase):
    # ── Colours used for the two different view-mode toggle buttons ──────────
    _TOGGLE_ON  = ACC                           # active tab colour
    _TOGGLE_OFF = ("#e2e8f0", "#334155")        # inactive tab colour
    _TOGGLE_OFF_TXT = T1

    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._search_var   = ctk.StringVar()
        self._search_var.trace_add("write", lambda *_: self._render_list())
        self._filter_var   = ctk.StringVar(value="All Categories")
        self._tmpl_cat_var = ctk.StringVar(value="All Categories")
        self._view_mode    = "protocols"   # "protocols" | "templates"
        self._build()

    # ── Build ─────────────────────────────────────────────────────────────────
    def _build(self):
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(22, 0))

        # ── Line 1: Title · Toggle · Search (all via pack) ───────────────────
        line1 = ctk.CTkFrame(hdr, fg_color="transparent")
        line1.pack(fill="x", pady=(0, 2))

        label(line1, "Protocol Library", size=24, weight="bold").pack(side="left")

        tog = ctk.CTkFrame(line1, fg_color="transparent")
        tog.pack(side="left", padx=(20, 0))
        self._btn_protos = ctk.CTkButton(
            tog, text="My Protocols", font=(FONT, 12, "bold"),
            width=118, height=32, corner_radius=R_SM,
            command=lambda: self._set_view("protocols"))
        self._btn_protos.pack(side="left", padx=(0, 2))
        self._btn_templates = ctk.CTkButton(
            tog, text="Templates", font=(FONT, 12, "bold"),
            width=95, height=32, corner_radius=R_SM,
            command=lambda: self._set_view("templates"))
        self._btn_templates.pack(side="left")

        # Search on the far right of line 1
        self._search_entry = ctk.CTkEntry(
            line1, placeholder_text="🔍  Search...",
            textvariable=self._search_var, font=(FONT, 13),
            height=34, width=220, corner_radius=R_XL, border_color=CARD_B)
        self._search_entry.pack(side="right")

        # ── Line 2a: Protocol controls (shown when view = "protocols") ───────
        self._proto_ctrl = ctk.CTkFrame(hdr, fg_color="transparent")
        # packed/forgotten by _set_view — NOT packed here
        btn(self._proto_ctrl, "+ New Protocol", self._open_create_menu,
            width=145).pack(side="left", padx=(0, 6))
        btn(self._proto_ctrl, "Import Protocol", self._open_import_menu,
            color=("#e2e8f0","#334155"), text_color=T1, width=140).pack(side="left", padx=6)
        self._filter_menu = ctk.CTkOptionMenu(
            self._proto_ctrl, variable=self._filter_var,
            values=["All Categories"],
            command=lambda _: self._render_list(),
            width=148, height=32, font=(FONT, 12), corner_radius=R_SM)
        self._filter_menu.pack(side="left", padx=(12, 4))
        btn(self._proto_ctrl, "Manage",
            lambda: TermManagerDialog(self.app, self.app, "categories"),
            color=("#e2e8f0","#334155"), text_color=T1,
            width=82, height=30, size=11).pack(side="left", padx=3)
        btn(self._proto_ctrl, "Tags",
            lambda: TermManagerDialog(self.app, self.app, "tags"),
            color=("#e2e8f0","#334155"), text_color=T1,
            width=58, height=30, size=11).pack(side="left", padx=3)

        # ── Line 2b: Template controls (shown when view = "templates") ───────
        self._tmpl_ctrl = ctk.CTkFrame(hdr, fg_color="transparent")
        # packed/forgotten by _set_view — NOT packed here
        self._tmpl_filter_menu = ctk.CTkOptionMenu(
            self._tmpl_ctrl, variable=self._tmpl_cat_var,
            values=["All Categories"],
            command=lambda _: self._render_list(),
            width=175, height=32, font=(FONT, 12), corner_radius=R_SM)
        self._tmpl_filter_menu.pack(side="left", padx=(0, 6))

        # ── Scrollable content area ──────────────────────────────────────────
        self.scroll = ScrollFrame(self)
        self.scroll.grid(row=1, column=0, sticky="nsew", padx=28, pady=(10, 28))
        self.scroll.grid_columnconfigure(0, weight=1)

        # Apply initial toggle state
        self._set_view("protocols")

    # ── View-mode switch ──────────────────────────────────────────────────────
    def _set_view(self, mode):
        self._view_mode = mode
        if mode == "protocols":
            self._btn_protos.configure(
                fg_color=self._TOGGLE_ON, text_color=("#fff", "#fff"))
            self._btn_templates.configure(
                fg_color=self._TOGGLE_OFF, text_color=self._TOGGLE_OFF_TXT)
            self._tmpl_ctrl.pack_forget()
            self._proto_ctrl.pack(fill="x", pady=(6, 0))
            self._search_entry.configure(placeholder_text="🔍  Search protocols...")
        else:
            self._btn_templates.configure(
                fg_color=self._TOGGLE_ON, text_color=("#fff", "#fff"))
            self._btn_protos.configure(
                fg_color=self._TOGGLE_OFF, text_color=self._TOGGLE_OFF_TXT)
            self._proto_ctrl.pack_forget()
            self._tmpl_ctrl.pack(fill="x", pady=(6, 0))
            self._search_entry.configure(placeholder_text="🔍  Search templates...")
        self._render_list()

    # ── Refresh ───────────────────────────────────────────────────────────────
    def refresh(self):
        self._refresh_filters()
        self._render_list()

    def _refresh_filters(self):
        # Protocol category filter
        cats = sorted({
            *self.app.categories,
            *(p.get("category","").strip()
              for p in self.app.protocols if p.get("category","").strip())
        })
        values = ["All Categories"] + cats
        self._filter_menu.configure(values=values)
        if self._filter_var.get() not in values:
            self._filter_var.set("All Categories")

        # Template category filter
        tcats = sorted({t.get("category","") for t in self.app.templates
                        if t.get("category","")})
        tvalues = ["All Categories"] + tcats
        self._tmpl_filter_menu.configure(values=tvalues)
        if self._tmpl_cat_var.get() not in tvalues:
            self._tmpl_cat_var.set("All Categories")

    def _open_create_menu(self):
        self.app.open_create_protocol("full")

    def _open_import_menu(self):
        self.app.open_create_protocol("import")

    # ── Render dispatcher ─────────────────────────────────────────────────────
    def _render_list(self):
        for w in self.scroll.winfo_children():
            w.destroy()
        try:
            self.scroll._parent_canvas.yview_moveto(0)
        except Exception:
            pass
        if self._view_mode == "templates":
            self._render_templates()
        else:
            self._render_protocols()

    # ── My Protocols section ──────────────────────────────────────────────────
    def _render_protocols(self):
        q   = self._search_var.get().lower()
        cat = self._filter_var.get()
        protocols = [p for p in self.app.protocols
                     if q in p.get("name","").lower()
                     or q in p.get("category","").lower()
                     or q in " ".join(p.get("tags",[])).lower()]
        if cat != "All Categories":
            protocols = [p for p in protocols
                         if p.get("category","").strip() == cat]
        if not protocols:
            label(self.scroll, "No protocols found", size=14,
                  color=T3).grid(row=0, column=0, pady=60)
            return
        for i, p in enumerate(protocols):
            self._proto_card(p, i)

    def _proto_card(self, p, row_idx):
        card = card_frame(self.scroll)
        card.grid(row=row_idx, column=0, sticky="ew", pady=(0, 8))
        card.grid_columnconfigure(0, weight=1)

        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="x", padx=16, pady=12)
        inner.grid_columnconfigure(0, weight=1)

        left = ctk.CTkFrame(inner, fg_color="transparent")
        left.grid(row=0, column=0, sticky="ew")

        name_row = ctk.CTkFrame(left, fg_color="transparent")
        name_row.pack(fill="x", anchor="w")
        label(name_row, p.get("name") or "Untitled",
              size=15, weight="bold").pack(side="left")
        if p.get("category"):
            badge = ctk.CTkFrame(name_row,
                                 fg_color=("#e2e8f0","#334155"), corner_radius=R_SM)
            badge.pack(side="left", padx=8)
            label(badge, p["category"], size=11, color=T2).pack(padx=8, pady=2)
        for tag in p.get("tags", [])[:3]:
            tb = ctk.CTkFrame(name_row,
                              fg_color=("#dbeafe","#1e3a5f"), corner_radius=R_SM)
            tb.pack(side="left", padx=(0, 6))
            label(tb, tag, size=10, color=ACC).pack(padx=7, pady=2)

        if p.get("description"):
            _desc = p["description"]
            _desc_short = (_desc[:80] + "…") if len(_desc) > 80 else _desc
            label(left, _desc_short, size=12,
                  color=T2).pack(anchor="w", pady=(2, 0))

        steps = p.get("steps", [])
        ho = sum(s.get("handsOnMinutes", 0) for s in steps)
        wt = sum(s.get("waitMinutes", 0) for s in steps)
        meta = (f"{len(steps)} steps  ·  Total: {fmt_mins(total_mins(steps))}"
                f"  ·  Hands-on: {fmt_mins(ho)}  ·  Wait: {fmt_mins(wt)}"
                f"  ·  {fmt_date(p['updatedAt'])}")
        label(left, meta, size=11, color=T3).pack(anchor="w", pady=(4, 0))

        actions = ctk.CTkFrame(inner, fg_color="transparent")
        actions.grid(row=0, column=1, sticky="e")

        def start(pid=p["id"]):  self.app.start_run(pid)
        def edit(pid=p["id"]):   self.app.open_editor(pid)
        def flow(pid=p["id"]):   self.app.open_flowchart(pid)
        def dup(pid=p["id"]):    self._dup(pid)
        def delete(pid=p["id"]): self._delete(pid)

        for txt, cmd, col in [
            ("▶  Run",  start,  GREEN),
            ("⎇  Flow", flow,   ("#7c3aed","#7c3aed")),
            ("✎  Edit", edit,   ACC),
            ("⧉  Copy", dup,    ("#64748b","#475569")),
            ("✕  Del",  delete, DANGER),
        ]:
            btn(actions, txt, cmd, color=col,
                width=80, height=30, size=12).pack(side="left", padx=3)

    def _dup(self, pid):
        protocols = self.app.protocols
        orig = next((p for p in protocols if p["id"] == pid), None)
        if not orig: return
        import copy
        dup = copy.deepcopy(orig)
        dup["id"] = new_id()
        dup["name"] = orig["name"] + " (copy)"
        dup["createdAt"] = dup["updatedAt"] = now_ts()
        for s in dup["steps"]:
            s["id"] = new_id()
            for c in s.get("checklist", []): c["id"] = new_id()
        protocols.append(dup)
        save_protocols(protocols)
        self.refresh()

    def _delete(self, pid):
        if not messagebox.askyesno(
                "Delete Protocol",
                "Delete this protocol? This cannot be undone."): return
        self.app.protocols = [p for p in self.app.protocols if p["id"] != pid]
        save_protocols(self.app.protocols)
        self.refresh()

    # ── Built-in Templates section ────────────────────────────────────────────
    def _render_templates(self):
        # Disclaimer banner
        disc = ctk.CTkFrame(
            self.scroll,
            fg_color=("#fef9c3", "#1c1917"),
            corner_radius=R_LG,
            border_width=1,
            border_color=("#fde68a", "#78350f"))
        disc.grid(row=0, column=0, sticky="ew", pady=(0, 14))
        label(disc,
              "ℹ  These templates are general starting points. "
              "Please adjust conditions according to your lab, reagent, instrument, and protocol requirements.",
              size=12, color=("#92400e", "#fbbf24")).pack(padx=14, pady=9)

        q          = self._search_var.get().lower()
        cat_filter = self._tmpl_cat_var.get()

        templates = self.app.templates
        if q:
            templates = [t for t in templates
                         if q in t.get("name","").lower()
                         or q in t.get("category","").lower()
                         or q in " ".join(t.get("tags",[])).lower()
                         or q in t.get("description","").lower()]
        if cat_filter != "All Categories":
            templates = [t for t in templates
                         if t.get("category","") == cat_filter]

        if not templates:
            label(self.scroll, "No templates found.", size=14,
                  color=T3).grid(row=1, column=0, pady=60)
            return

        # Group by category, preserve sorted order
        from collections import defaultdict as _dd
        by_cat = _dd(list)
        for t in templates:
            by_cat[t.get("category","Other")].append(t)

        row = 1
        for cat_name in sorted(by_cat):
            # Category section header
            cat_hdr = ctk.CTkFrame(self.scroll, fg_color="transparent")
            cat_hdr.grid(row=row, column=0, sticky="ew", pady=(14, 4))
            label(cat_hdr, cat_name, size=15, weight="bold").pack(side="left")
            sep_line = ctk.CTkFrame(cat_hdr, fg_color=CARD_B, height=1)
            sep_line.pack(side="left", fill="x", expand=True, padx=(12, 0))
            row += 1
            for tmpl in by_cat[cat_name]:
                self._template_card(tmpl, row)
                row += 1

    def _template_card(self, tmpl, row_idx):
        card = ctk.CTkFrame(
            self.scroll,
            fg_color=("#f0fdf4", "#052e16"),
            corner_radius=R_LG,
            border_width=2,
            border_color=("#6ee7b7", "#059669"))
        card.grid(row=row_idx, column=0, sticky="ew", pady=(0, 8))
        card.grid_columnconfigure(0, weight=1)

        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="x", padx=16, pady=12)
        inner.grid_columnconfigure(0, weight=1)

        left = ctk.CTkFrame(inner, fg_color="transparent")
        left.grid(row=0, column=0, sticky="ew")

        # ── Name row ─────────────────────────────────────────────────────────
        name_row = ctk.CTkFrame(left, fg_color="transparent")
        name_row.pack(fill="x", anchor="w")

        # "TEMPLATE" pill badge
        tmpl_pill = ctk.CTkFrame(
            name_row, fg_color=("#6ee7b7","#065f46"), corner_radius=R_SM)
        tmpl_pill.pack(side="left", padx=(0, 8))
        label(tmpl_pill, "TEMPLATE", size=10,
              color=("#064e3b","#d1fae5")).pack(padx=7, pady=2)

        label(name_row, tmpl.get("name",""), size=15,
              weight="bold").pack(side="left")

        for tag in tmpl.get("tags", [])[:4]:
            tb = ctk.CTkFrame(name_row,
                              fg_color=("#d1fae5","#064e3b"), corner_radius=R_SM)
            tb.pack(side="left", padx=(6, 0))
            label(tb, tag, size=10, color=("#065f46","#6ee7b7")).pack(padx=7, pady=2)

        # ── Description ──────────────────────────────────────────────────────
        desc = tmpl.get("description","")
        if desc:
            label(left, (desc[:95] + "…") if len(desc) > 95 else desc,
                  size=12, color=T2).pack(anchor="w", pady=(4, 0))

        # ── Stats ─────────────────────────────────────────────────────────────
        steps  = tmpl.get("steps", [])
        total_m = sum(s.get("handsOnMinutes",0) + s.get("waitMinutes",0) for s in steps)
        ho_m    = sum(s.get("handsOnMinutes",0) for s in steps)
        meta = (f"{len(steps)} steps  ·  Total: {fmt_mins(total_m)}"
                f"  ·  Hands-on: {fmt_mins(ho_m)}")
        label(left, meta, size=11, color=T3).pack(anchor="w", pady=(4, 0))

        # ── Use Template button ───────────────────────────────────────────────
        acts = ctk.CTkFrame(inner, fg_color="transparent")
        acts.grid(row=0, column=1, sticky="e", padx=(12, 0))
        btn(acts, "＋ Use Template",
            lambda t=tmpl: self._use_template(t),
            color=GREEN, width=138, height=36, size=13).pack()

    def _use_template(self, tmpl):
        """Deep-copy a built-in template into a new user protocol and open the editor."""
        import copy

        p = new_protocol()
        p["name"]        = tmpl.get("name", "Protocol")
        p["category"]    = tmpl.get("category", "")
        p["description"] = tmpl.get("description", "")
        p["tags"]        = list(tmpl.get("tags", []))
        p["source"] = {
            "sourceType":       "template",
            "sourceName":       f"Built-in Template: {tmpl.get('name','')}",
            "sourceUrl":        tmpl.get("source_reference_url", ""),
            "importedAt":       now_ts(),
            "rawText":          "",
            "originalFileName": tmpl.get("templateId", ""),
            "notes":            tmpl.get("source_note", ""),
        }

        steps = []
        for i, ts in enumerate(tmpl.get("steps", [])):
            s = copy.deepcopy(ts)
            s["id"]    = new_id()
            s["order"] = i
            # Ensure all fields from new_step() schema are present
            s.setdefault("type",               "preparation")
            s.setdefault("description",        "")
            s.setdefault("reagents",           [])
            s.setdefault("equipment",          [])
            s.setdefault("handsOnMinutes",     5)
            s.setdefault("waitMinutes",        0)
            s.setdefault("bufferMinutes",      0)
            s.setdefault("temperature",        "")
            s.setdefault("centrifugeCondition","")
            s.setdefault("shakingRotation",    "")
            s.setdefault("checklist",          [])
            s.setdefault("notes",              "")
            s.setdefault("warnings",           "")
            s.setdefault("substeps",           [])
            steps.append(s)

        p["steps"] = steps
        self.app.protocols.append(p)
        save_protocols(self.app.protocols)
        messagebox.showinfo(
            "Template Copied",
            f'"{p["name"]}" has been added to your protocols.\n\n'
            "You can now edit, rename, and customise it freely.\n"
            "The original template is never modified.")
        self.app.open_editor(p["id"])

# ─── Editor ───────────────────────────────────────────────────────────────────
class EditorPage(PageBase):
    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.protocol = None
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._build()

    def _build(self):
        # Toolbar — two rows
        self.toolbar = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG,
                                    border_width=0)
        self.toolbar.grid(row=0, column=0, sticky="ew")
        self.toolbar.grid_columnconfigure(1, weight=1)

        # Row 0: Back | name (expands) | Save | Run
        btn(self.toolbar, "← Library", self._back_to_library,
            color=("#e2e8f0","#334155"), text_color=T1, width=95, height=34, size=12
            ).grid(row=0, column=0, padx=(12,8), pady=(10,4), sticky="w")

        self.e_name = entry(self.toolbar, "Protocol name...", width=300)
        self.e_name.configure(font=(FONT, 15, "bold"), height=36)
        self.e_name.grid(row=0, column=1, padx=4, pady=(10,4), sticky="ew")

        self.btn_save = btn(self.toolbar, "💾  Save", self._save, width=100, height=34)
        self.btn_save.grid(row=0, column=2, padx=4, pady=(10,4))
        btn(self.toolbar, "▶  Run", self._run, color=GREEN, width=90, height=34
            ).grid(row=0, column=3, padx=(4,12), pady=(10,4))

        # Row 1: category | Categories | tags | Tags | meta
        meta_row = ctk.CTkFrame(self.toolbar, fg_color="transparent")
        meta_row.grid(row=1, column=0, columnspan=4, sticky="ew", padx=12, pady=(0,8))

        self.category_var = ctk.StringVar(value="")
        self.e_cat = ctk.CTkComboBox(meta_row, variable=self.category_var,
                                     values=self._category_values(),
                                     font=(FONT, 12), width=150, height=30,
                                     corner_radius=R_SM, border_color=CARD_B)
        self.e_cat.pack(side="left", padx=(0,4))
        btn(meta_row, "Categories", lambda: TermManagerDialog(self.app, self.app, "categories"),
            color=("#e2e8f0","#334155"), text_color=T1, width=90, height=28, size=11
            ).pack(side="left", padx=(0,10))
        self.e_tags = entry(meta_row, "Tags, comma separated", width=200)
        self.e_tags.pack(side="left", padx=(0,4))
        btn(meta_row, "Tags", lambda: TermManagerDialog(self.app, self.app, "tags"),
            color=("#e2e8f0","#334155"), text_color=T1, width=55, height=28, size=11
            ).pack(side="left", padx=(0,12))
        self.lbl_meta = label(meta_row, "", size=12, color=T2)
        self.lbl_meta.pack(side="left")

        # Scroll
        self.scroll = ScrollFrame(self)
        self.scroll.grid(row=1, column=0, sticky="nsew", padx=28, pady=(12,0))
        self.scroll.grid_columnconfigure(0, weight=1)

        # Description
        desc_f = card_frame(self.scroll)
        desc_f.grid(row=0, column=0, sticky="ew", pady=(0,12))
        label(desc_f, "Description", size=12, color=T2).pack(anchor="w", padx=14, pady=(10,2))
        self.e_desc = entry(desc_f, "Protocol description...", width=900)
        self.e_desc.pack(fill="x", padx=14, pady=(0,10))

        # Steps header
        steps_hdr = ctk.CTkFrame(self.scroll, fg_color="transparent")
        steps_hdr.grid(row=1, column=0, sticky="ew", pady=(0,8))
        label(steps_hdr, "Steps", size=14, weight="bold").pack(side="left")
        btn(steps_hdr, "+ Add Step", self._add_step, width=110, height=30).pack(side="right")

        # Step list
        self.steps_frame = ctk.CTkFrame(self.scroll, fg_color="transparent")
        self.steps_frame.grid(row=2, column=0, sticky="ew")
        self.steps_frame.grid_columnconfigure(0, weight=1)

        # Footer add button
        self.footer = ctk.CTkFrame(self.scroll, fg_color="transparent")
        self.footer.grid(row=3, column=0, sticky="ew", pady=16)

    def load(self, protocol):
        self.protocol = protocol
        self.refresh_terms()
        self.e_name.delete(0, "end")
        self.e_name.insert(0, protocol.get("name", ""))
        self.category_var.set(protocol.get("category", ""))
        self.e_tags.delete(0, "end")
        self.e_tags.insert(0, ", ".join(protocol.get("tags", [])))
        self.e_desc.delete(0, "end")
        self.e_desc.insert(0, protocol.get("description", ""))
        self._render_steps()
        self._update_meta()

    def _category_values(self):
        vals = [""] + sorted({*self.app.categories, *(p.get("category", "").strip() for p in self.app.protocols if p.get("category", "").strip())})
        return vals

    def refresh_terms(self):
        try:
            self.e_cat.configure(values=self._category_values())
        except Exception:
            pass

    def _update_meta(self):
        if not self.protocol: return
        steps = self.protocol.get("steps", [])
        self.lbl_meta.configure(text=f"{len(steps)} steps  ·  {fmt_mins(total_mins(steps))}")

    def _render_steps(self):
        for w in self.steps_frame.winfo_children(): w.destroy()
        if not self.protocol: return
        for i, step in enumerate(self.protocol.get("steps", [])):
            self._step_row(step, i)

    def _step_row(self, step, idx):
        stype = step.get("type", "custom")
        bg_col = STEP_COLORS.get(stype, STEP_COLORS["custom"])
        bd_col = STEP_BADGES.get(stype, STEP_BADGES["custom"])

        card = ctk.CTkFrame(self.steps_frame, fg_color=bg_col, corner_radius=R_SM,
                            border_width=2, border_color=bd_col)
        card.grid(row=idx, column=0, sticky="ew", pady=(0,6))
        card.grid_columnconfigure(1, weight=1)

        # Index + drag hint
        label(card, f"{idx+1}", size=12, color=T2).grid(row=0, column=0, padx=(12,4), pady=12)

        # Badge
        badge = ctk.CTkFrame(card, fg_color=bd_col, corner_radius=R_SM)
        badge.grid(row=0, column=1, sticky="w", padx=4)
        label(badge, STEP_LABELS.get(stype,"Custom"), size=11,
              color=("#ffffff","#ffffff")).pack(padx=8, pady=3)

        # Title
        title = step.get("title") or "Untitled step"
        label(card, title, size=13, weight="bold").grid(row=0, column=2, sticky="w", padx=8)

        # Time badges
        total_m = step.get("handsOnMinutes",0) + step.get("waitMinutes",0) + step.get("bufferMinutes",0)
        if total_m:
            label(card, fmt_mins(total_m), size=12, color=T2).grid(row=0, column=3, padx=8)

        if step.get("temperature"):
            label(card, f"🌡 {step['temperature']}", size=11, color=T2).grid(row=0, column=4, padx=4)
        if step.get("centrifugeCondition"):
            label(card, f"🔄 {step['centrifugeCondition']}", size=11, color=T2).grid(row=0, column=5, padx=4)
        if step.get("shakingRotation"):
            label(card, f"↻ {step['shakingRotation']}", size=11, color=T2).grid(row=0, column=6, padx=4)

        # Actions
        acts = ctk.CTkFrame(card, fg_color="transparent")
        acts.grid(row=0, column=7, sticky="e", padx=(8,12))

        def edit_step(s=step): self._edit_step(s)
        def dup_step(i=idx): self._dup_step(i)
        def del_step(i=idx): self._del_step(i)
        def move_up(i=idx): self._move_step(i, -1)
        def move_dn(i=idx): self._move_step(i, 1)

        for txt, cmd, col, en in [
            ("↑", move_up, ("#94a3b8","#475569"), idx > 0),
            ("↓", move_dn, ("#94a3b8","#475569"), idx < len(self.protocol["steps"])-1),
            ("Edit", edit_step, ACC, True),
            ("⧉", dup_step, ("#64748b","#475569"), True),
            ("✕", del_step, DANGER, True),
        ]:
            b = ctk.CTkButton(acts, text=txt, width=36, height=28, fg_color=col,
                              font=(FONT,12), command=cmd, corner_radius=R_XS,
                              state="normal" if en else "disabled")
            b.pack(side="left", padx=2)

    def _edit_step(self, step):
        def on_save(updated):
            steps = self.protocol["steps"]
            for i, s in enumerate(steps):
                if s["id"] == updated["id"]:
                    steps[i] = updated
                    break
            self._render_steps()
            self._update_meta()
        StepEditorDialog(self.app, step, on_save)

    def _add_step(self):
        s = new_step(len(self.protocol.get("steps", [])))
        self.protocol.setdefault("steps", []).append(s)
        self._render_steps()
        self._update_meta()
        StepEditorDialog(self.app, s, lambda upd: (
            self.protocol["steps"].__setitem__(
                next(i for i, x in enumerate(self.protocol["steps"]) if x["id"]==upd["id"]), upd
            ),
            self._render_steps(),
            self._update_meta()
        ))

    def _dup_step(self, idx):
        import copy
        orig = self.protocol["steps"][idx]
        dup = copy.deepcopy(orig)
        dup["id"] = new_id()
        self.protocol["steps"].insert(idx+1, dup)
        for i, s in enumerate(self.protocol["steps"]): s["order"] = i
        self._render_steps()
        self._update_meta()

    def _del_step(self, idx):
        self.protocol["steps"].pop(idx)
        for i, s in enumerate(self.protocol["steps"]): s["order"] = i
        self._render_steps()
        self._update_meta()

    def _move_step(self, idx, delta):
        steps = self.protocol["steps"]
        new_idx = idx + delta
        if 0 <= new_idx < len(steps):
            steps[idx], steps[new_idx] = steps[new_idx], steps[idx]
            for i, s in enumerate(steps): s["order"] = i
            self._render_steps()

    def _has_unsaved_changes(self):
        if not self.protocol:
            return False
        return (
            self.e_name.get().strip() != self.protocol.get("name", "") or
            self.category_var.get().strip() != self.protocol.get("category", "") or
            self.e_tags.get().strip() != ", ".join(self.protocol.get("tags", [])) or
            self.e_desc.get().strip() != self.protocol.get("description", "")
        )

    def _back_to_library(self):
        if self._has_unsaved_changes():
            if not messagebox.askyesno(
                "Discard Changes?",
                "You have unsaved changes. Discard and return to Library?",
                icon="warning",
            ):
                return
        self.app.navigate("library")

    def _save(self):
        if not self.protocol: return
        self.protocol["name"] = self.e_name.get().strip() or "New Protocol"
        self.protocol["category"] = self.category_var.get().strip()
        self.protocol["tags"] = [x.strip() for x in self.e_tags.get().split(",") if x.strip()]
        self.protocol["description"] = self.e_desc.get().strip()
        self.protocol["updatedAt"] = now_ts()
        if self.protocol["category"] and self.protocol["category"] not in self.app.categories:
            self.app.categories.append(self.protocol["category"])
            save_categories(self.app.categories)
        new_tags = [t for t in self.protocol["tags"] if t not in self.app.tags]
        if new_tags:
            self.app.tags.extend(new_tags)
            save_tags(self.app.tags)
        protocols = self.app.protocols
        existing = next((i for i, p in enumerate(protocols) if p["id"] == self.protocol["id"]), None)
        if existing is not None:
            protocols[existing] = self.protocol
        else:
            protocols.append(self.protocol)
        self.app.protocols = protocols
        save_protocols(protocols)
        self.btn_save.configure(text="✓ Saved!", fg_color=GREEN)
        self.after(1800, lambda: self.btn_save.configure(text="💾  Save", fg_color=ACC))

    def _run(self):
        self._save()
        self.app.start_run(self.protocol["id"])

# ─── Flowchart ────────────────────────────────────────────────────────────────
class FlowchartPage(PageBase):
    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.protocol = None
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._build()

    def _build(self):
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(20,8))
        self.lbl_title = label(hdr, "Flowchart", size=20, weight="bold")
        self.lbl_title.pack(side="left")
        label(hdr, "Click a step to view details", size=12, color=T2).pack(side="left", padx=16)

        # Protocol selector
        sel_row = ctk.CTkFrame(hdr, fg_color="transparent")
        sel_row.pack(side="left", padx=16)
        label(sel_row, "Protocol:", size=12, color=T2).pack(side="left", padx=(0,6))
        self._proto_var = ctk.StringVar(value="— Select protocol —")
        self._proto_menu = ctk.CTkOptionMenu(
            sel_row, variable=self._proto_var,
            values=["— Select protocol —"],
            font=(FONT, 12), width=260, corner_radius=R_SM,
            command=self._on_proto_select)
        self._proto_menu.pack(side="left")
        # Fit view button
        btn(hdr, "⊞ Fit View", self._fit_view,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=90, height=30, size=12).pack(side="right")

        main = ctk.CTkFrame(self, fg_color="transparent")
        main.grid(row=1, column=0, sticky="nsew", padx=28, pady=(0,28))
        main.grid_columnconfigure(0, weight=1)
        main.grid_rowconfigure(0, weight=1)

        # Canvas for flowchart
        self.canvas_frame = card_frame(main)
        self.canvas_frame.grid(row=0, column=0, sticky="nsew")
        self.canvas_frame.grid_columnconfigure(0, weight=1)
        self.canvas_frame.grid_rowconfigure(0, weight=1)

        self.canvas = tk.Canvas(self.canvas_frame, bg="#f8fafc", highlightthickness=0)
        self.canvas.grid(row=0, column=0, sticky="nsew", padx=2, pady=2)
        vsb = ctk.CTkScrollbar(self.canvas_frame, command=self.canvas.yview, orientation="vertical",
                               corner_radius=R_SM,
                               width=12, fg_color="transparent",
                               button_color=("#cbd5e1","#475569"),
                               button_hover_color=("#94a3b8","#64748b"))
        vsb.grid(row=0, column=1, sticky="ns")
        hsb = ctk.CTkScrollbar(self.canvas_frame, command=self.canvas.xview, orientation="horizontal",
                               corner_radius=R_SM,
                               height=12, fg_color="transparent",
                               button_color=("#cbd5e1","#475569"),
                               button_hover_color=("#94a3b8","#64748b"))
        hsb.grid(row=1, column=0, sticky="ew")
        self.canvas.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.canvas.bind("<Button-1>", self._on_canvas_click)
        self.canvas.bind("<MouseWheel>", self._fc_wheel)
        self.canvas.bind("<Button-4>", lambda e: self.canvas.yview_scroll(-3, "units"))
        self.canvas.bind("<Button-5>", lambda e: self.canvas.yview_scroll(3, "units"))

        # Detail panel
        self.detail_frame = card_frame(main)
        self.detail_frame.grid(row=0, column=1, sticky="nsew", padx=(12,0))
        self.detail_frame.configure(width=260)
        self.detail_frame.grid_columnconfigure(0, weight=1)
        self.detail_frame.grid_rowconfigure(0, weight=1)
        main.grid_columnconfigure(1, minsize=260, weight=0)
        label(self.detail_frame, "Select a step", size=13, color=T3).grid(row=0, column=0, pady=40)
        self._node_ids = {}  # canvas item id → step

    def _fc_wheel(self, event):
        """Proper trackpad/mouse-wheel handler for the flowchart canvas."""
        delta = getattr(event, "delta", 0)
        if abs(delta) >= 120:
            units = int(-delta / 120) * 3
        else:
            units = int(-delta / 3) or (-1 if delta > 0 else 1)
        self.canvas.yview_scroll(units, "units")

    def load(self, protocol):
        self.protocol = protocol
        self.lbl_title.configure(text=protocol.get("name", "Flowchart"))
        self._draw_graph()

    def _draw_graph(self):
        self.canvas.delete("all")
        self._node_ids = {}
        if not self.protocol: return

        steps = self.protocol.get("steps", [])
        if not steps:
            w = self.canvas.winfo_width() or 600
            h = self.canvas.winfo_height() or 400
            self.canvas.create_text(w//2, h//2,
                text="This protocol has no steps yet.",
                font=(FONT, 14), fill="#94a3b8", anchor="center")
            return

        NODE_W, NODE_H = 220, 96   # extra height for potential 2-line titles
        GAP = 40
        X_CENTER = 300
        total_h = len(steps) * (NODE_H + GAP) + GAP

        self.canvas.configure(scrollregion=(0, 0, X_CENTER*2, total_h + 20))

        for i, step in enumerate(steps):
            x = X_CENTER - NODE_W // 2
            y = GAP + i * (NODE_H + GAP)

            stype = step.get("type", "custom")
            bg = STEP_COLORS.get(stype, STEP_COLORS["custom"])
            bd = STEP_BADGES.get(stype, STEP_BADGES["custom"])
            # use light colors only (index 0)
            bg_c = bg[0]
            bd_c = bd[0]

            # connector line (arrow)
            if i > 0:
                prev_y = GAP + (i-1) * (NODE_H + GAP) + NODE_H
                self.canvas.create_line(
                    X_CENTER, prev_y, X_CENTER, y,
                    arrow=tk.LAST, fill="#94a3b8", width=2, arrowshape=(10,12,4)
                )

            # Card
            rid = canvas_round_rect(self.canvas, x, y, x+NODE_W, y+NODE_H,
                                    radius=R_LG, fill=bg_c, outline=bd_c,
                                    width=2, tags=f"node_{step['id']}")
            # Badge
            canvas_round_rect(self.canvas, x+8, y+8, x+8+len(STEP_LABELS.get(stype,""))*7+16, y+26,
                              radius=R_XS, fill=bd_c, outline="", tags=f"node_{step['id']}")
            self.canvas.create_text(x+8+8, y+17, text=STEP_LABELS.get(stype,""), anchor="w",
                                     font=(FONT, 10, "bold"), fill="white", tags=f"node_{step['id']}")
            # Step number
            self.canvas.create_text(x+NODE_W-8, y+17, text=f"#{i+1}", anchor="e",
                                     font=(FONT, 10), fill="#94a3b8", tags=f"node_{step['id']}")
            # Title — use canvas width parameter so tk wraps long titles
            # Safety cap at 64 chars to avoid extreme overflow in very tall nodes
            title = step.get("title") or "Untitled"
            if len(title) > 64: title = title[:64] + "…"
            self.canvas.create_text(X_CENTER, y+46, text=title, anchor="center",
                                     font=(FONT, 12, "bold"), fill="#0f172a",
                                     width=NODE_W-24,
                                     tags=f"node_{step['id']}")
            # Time / conditions row (nudged down to y+78 to clear 2-line title)
            total_m = step.get("handsOnMinutes",0)+step.get("waitMinutes",0)+step.get("bufferMinutes",0)
            info = fmt_mins(total_m)
            if step.get("temperature"): info += f"  🌡 {step['temperature']}"
            if step.get("centrifugeCondition"): info += f"  🔄 {step['centrifugeCondition']}"
            if step.get("shakingRotation"): info += f"  ↻ {step['shakingRotation']}"
            self.canvas.create_text(X_CENTER, y+80, text=info, anchor="center",
                                     font=(FONT, 10), fill="#64748b",
                                     width=NODE_W-24,
                                     tags=f"node_{step['id']}")

            self._node_ids[rid] = step
            for tid in self.canvas.find_withtag(f"node_{step['id']}"):
                self._node_ids[tid] = step

    def refresh(self):
        protos = self.app.protocols
        names = [p.get("name","Untitled") for p in protos]
        if not names:
            self._proto_menu.configure(values=["— No protocols —"])
            self._proto_var.set("— No protocols —")
            return
        self._proto_menu.configure(values=["— Select protocol —"] + names)
        if self.protocol:
            cur_name = self.protocol.get("name","")
            if cur_name in names:
                self._proto_var.set(cur_name)
                return
        self._proto_var.set("— Select protocol —")
        self._show_empty_state()

    def _on_proto_select(self, name):
        if name.startswith("—"):
            self._show_empty_state(); return
        proto = next((p for p in self.app.protocols
                      if p.get("name","") == name), None)
        if proto:
            self.load(proto)

    def _fit_view(self):
        self.canvas.update_idletasks()
        bbox = self.canvas.bbox("all")
        if not bbox: return
        x1,y1,x2,y2 = bbox
        cw = self.canvas.winfo_width()
        ch = self.canvas.winfo_height()
        if cw > 0 and ch > 0:
            self.canvas.xview_moveto((x1 - 20) / max(x2 - x1 + 40, 1))
            self.canvas.yview_moveto(0)

    def _show_empty_state(self):
        self.canvas.delete("all")
        self._node_ids = {}
        self.lbl_title.configure(text="Flowchart")
        for w in self.detail_frame.winfo_children(): w.destroy()
        label(self.detail_frame, "Select a step", size=13, color=T3).grid(row=0, column=0, pady=40)
        w = self.canvas.winfo_width() or 600
        h = self.canvas.winfo_height() or 400
        self.canvas.create_text(
            w // 2, h // 2,
            text="Select a protocol above to view its workflow.",
            font=(FONT, 14), fill="#94a3b8", anchor="center"
        )

    def _on_canvas_click(self, event):
        cx = self.canvas.canvasx(event.x)
        cy = self.canvas.canvasy(event.y)
        items = self.canvas.find_overlapping(cx-1, cy-1, cx+1, cy+1)
        for item in items:
            if item in self._node_ids:
                self._show_detail(self._node_ids[item])
                return

    def _show_detail(self, step):
        for w in self.detail_frame.winfo_children(): w.destroy()
        f = ScrollFrame(self.detail_frame)
        f.grid(row=0, column=0, sticky="nsew", padx=8, pady=8)
        f.grid_columnconfigure(0, weight=1)

        stype = step.get("type","custom")
        bd = STEP_BADGES.get(stype, STEP_BADGES["custom"])
        badge = ctk.CTkFrame(f, fg_color=bd, corner_radius=R_SM)
        badge.grid(row=0, column=0, sticky="w", pady=(8,4))
        label(badge, STEP_LABELS.get(stype,""), size=11, color=("#fff","#fff")).pack(padx=8, pady=3)

        label(f, step.get("title") or "Untitled", size=14, weight="bold",
              wraplength=220, justify="left").grid(row=1, column=0, sticky="ew", pady=(0,8))

        def section(lbl_txt, val, row, color=None):
            label(f, lbl_txt, size=11, color=T2, weight="bold").grid(
                row=row, column=0, sticky="w", pady=(8,1))
            label(f, display_na(val), size=12, color=color or T1,
                  wraplength=220, justify="left").grid(row=row+1, column=0, sticky="ew")
            return row + 2

        def reagent_text(rg):
            return reagent_display_text(rg)

        r = 2
        r = section("Step Type", STEP_LABELS.get(stype, "Custom"), r)
        r = section("Description", step.get("description"), r)
        r = section("Hands-on Time", fmt_optional_mins(step.get("handsOnMinutes", 0)), r)
        r = section("Wait / Incubation Time", fmt_optional_mins(step.get("waitMinutes", 0)), r)
        r = section("Buffer Time", fmt_optional_mins(step.get("bufferMinutes", 0)), r)
        r = section("Temperature", step.get("temperature"), r)
        r = section("Centrifuge", step.get("centrifugeCondition"), r)
        r = section("Shaking / Rotation", step.get("shakingRotation"), r)
        r = section("Equipment", "\n".join(display_list(step.get("equipment", []))), r)
        r = section("Reagents / Materials",
                    "\n".join(display_list(step.get("reagents", []), reagent_text)), r)
        r = section("Checklist",
                    "\n".join(display_list(step.get("checklist", []), lambda c: c.get("text", ""))), r)
        r = section("Sub-steps",
                    "\n".join(display_list(step.get("substeps", []), lambda s: s.get("text", ""))), r)
        r = section("Notes", step.get("notes"), r)
        section("Warnings", step.get("warnings"), r, color=("#92400e","#fde68a") if step.get("warnings") else None)

# ─── Time-difference dialog ────────────────────────────────────────────────────
class TimeDiffDialog(ctk.CTkToplevel):
    """Shown after completing a step when actual ≠ planned."""
    def __init__(self, parent, step_title, planned_mins, actual_mins, on_choice):
        super().__init__(parent)
        self.on_choice = on_choice
        self.withdraw()
        self.title("Step Completed")
        self.resizable(False, False)

        f = ctk.CTkFrame(self, fg_color="transparent")
        f.pack(fill="both", expand=True, padx=24, pady=20)

        diff = actual_mins - planned_mins
        diff_txt = f"+{diff}m over" if diff > 0 else f"{abs(diff)}m under"
        diff_col = DANGER if diff > 0 else GREEN

        label(f, f"✓  {step_title}", size=15, weight="bold").pack(anchor="w")
        row = ctk.CTkFrame(f, fg_color="transparent")
        row.pack(anchor="w", pady=(6,16))
        label(row, f"Planned: {fmt_mins(planned_mins)}", size=13, color=T2).pack(side="left")
        label(row, f"  Actual: {fmt_mins(actual_mins)}", size=13, color=T1).pack(side="left")
        label(row, f"  ({diff_txt})", size=13, color=diff_col).pack(side="left")

        label(f, "What would you like to do?", size=13, color=T2).pack(anchor="w", pady=(0,10))

        for txt, val in [
            ("Keep original protocol unchanged",         "keep"),
            ("Update this step in the current protocol", "update"),
            ("Save as my default speed for this step",   "default"),
        ]:
            btn(f, txt, lambda v=val: self._choose(v),
                color=ACC, text_color=("#fff","#fff"), width=400, height=36).pack(anchor="w", pady=3)

        btn(f, "Dismiss", lambda: self._choose("keep"),
            color=("#e2e8f0","#334155"), text_color=T1, width=120, height=32).pack(anchor="w", pady=(10,0))
        _show_dialog(self, parent, 480, 320)

    def _choose(self, val):
        self.on_choice(val)
        self.destroy()


# ─── Add Block Dialog ─────────────────────────────────────────────────────────
class AddBlockDialog(ctk.CTkToplevel):
    """Mini dialog to add a temporary block during a run."""
    BLOCK_TYPES = [
        ("action",    "Action Block",    "preparation"),
        ("timer",     "Timer Block",     "incubation"),
        ("note",      "Note",            "note"),
        ("checklist", "Checklist",       "checklist_block"),
        ("custom",    "Custom",          "custom"),
    ]

    def __init__(self, parent, on_confirm):
        super().__init__(parent)
        self.on_confirm = on_confirm
        self.withdraw()
        self.title("Add Temporary Block")
        self.resizable(True, True)
        self._type_var = ctk.StringVar(value="action")   # must match BLOCK_TYPES[0][0]
        self._pos_var  = ctk.StringVar(value="end")
        self._build()
        _show_dialog(self, parent, 480, 500)

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Sticky header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 4))
        label(hdr, "Add Block During Run", size=14, weight="bold").pack(anchor="w")

        # Scrollable body — all form fields go here
        body = ScrollFrame(self)
        body.grid(row=1, column=0, sticky="nsew", padx=8, pady=0)

        f = ctk.CTkFrame(body, fg_color="transparent")
        f.pack(fill="x", padx=12, pady=(8, 12))

        label(f, "Block type:", size=12, color=T2).pack(anchor="w")
        type_menu = ctk.CTkOptionMenu(f, variable=self._type_var,
                                       values=[t[0] for t in self.BLOCK_TYPES],
                                       font=(FONT,12), width=240, corner_radius=R_SM)
        type_menu.pack(anchor="w", pady=(2,10))

        label(f, "Title:", size=12, color=T2).pack(anchor="w")
        self.e_title = entry(f, "Block title...", width=400)
        self.e_title.pack(anchor="w", pady=(2,8))

        label(f, "Description / note:", size=12, color=T2).pack(anchor="w")
        self.tb_desc = textbox(f, width=400, height=72)
        self.tb_desc.pack(anchor="w", pady=(2,8))

        row_time = ctk.CTkFrame(f, fg_color="transparent")
        row_time.pack(fill="x", pady=(0,4))
        label(row_time, "Planned time (min):", size=12, color=T2).pack(side="left", padx=(0,8))
        self.e_time = entry(row_time, "5", width=80)
        self.e_time.pack(side="left")

        label(f, "Insert position:", size=12, color=T2).pack(anchor="w", pady=(10,2))
        for val, txt in [("end","At end of protocol"),
                         ("after_current","After current step"),
                         ("before_current","Before current step")]:
            ctk.CTkRadioButton(f, text=txt, variable=self._pos_var, value=val,
                               font=(FONT,12)).pack(anchor="w", pady=2)
        # Spacer so last radio isn't flush against the footer
        ctk.CTkFrame(f, height=8, fg_color="transparent").pack()

        # Sticky footer — always visible at the bottom
        footer = ctk.CTkFrame(self, fg_color=CARD,
                               border_width=1, border_color=CARD_B,
                               corner_radius=0)
        footer.grid(row=2, column=0, sticky="ew")
        brow = ctk.CTkFrame(footer, fg_color="transparent")
        brow.pack(fill="x", padx=20, pady=12)
        btn(brow, "Cancel", self.destroy,
            color=("#e2e8f0","#334155"), text_color=T1, width=90).pack(side="left")
        btn(brow, "Add Block", self._confirm, color=ACC, width=120).pack(side="right")

    def _confirm(self):
        # Read ALL widget values BEFORE destroy() — accessing Tcl vars after
        # destroy() is unreliable on some CTk/Tk versions.
        title     = self.e_title.get().strip() or "Unnamed Block"
        desc      = self.tb_desc.get("0.0", "end").strip()
        position  = self._pos_var.get()
        btype_key = self._type_var.get()
        try:
            mins = float(self.e_time.get().strip() or "5")
        except Exception:
            mins = 5.0
        step_type = next(
            (t[2] for t in self.BLOCK_TYPES if t[0] == btype_key), "custom")
        self.destroy()
        self.on_confirm({
            "title": title, "description": desc,
            "time_mins": mins, "step_type": step_type,
            "position": position,
        })


# ─── Run Mode ─────────────────────────────────────────────────────────────────
class RunPage(PageBase):
    """Protocol Timers — flexible per-step timer workspace."""

    ST_IDLE      = "idle"
    ST_RUNNING   = "running"
    ST_PAUSED    = "paused"
    ST_COMPLETED = "completed"
    ST_SKIPPED   = "skipped"

    ST_ICON  = {"idle": "○", "running": "▶", "paused": "⏸", "completed": "✓", "skipped": "⤼"}
    ST_COLOR = {
        "idle":      ("#cbd5e1","#475569"),
        "running":   ("#3b82f6","#3b82f6"),
        "paused":    ("#f97316","#f97316"),
        "completed": ("#22c55e","#22c55e"),
        "skipped":   ("#94a3b8","#64748b"),
    }

    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.protocol             = None
        self._step_states         = []
        self._seq_mode            = False
        self._seq_idx             = -1
        self._auto_start          = False
        self._timeline            = []
        self._session_start_ts    = None
        self._temp_block_indices  = []
        self._autosave_job        = None
        self._build()
        # Offer to resume if a checkpoint exists from a previous crash / close
        if RUNTIME_FILE.exists():
            self.after(600, self._offer_resume)

    def _build(self):
        self.grid_columnconfigure(0, minsize=230, weight=0)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Title bar
        tb = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG)
        tb.grid(row=0, column=0, columnspan=2, sticky="ew")
        tb.grid_columnconfigure(1, weight=1)
        label(tb, "Protocol Timers", size=16, weight="bold").grid(
            row=0, column=0, padx=20, pady=12, sticky="w")
        self._lbl_selected = label(tb, "Choose a protocol from the left panel",
                                   size=12, color=T3)
        self._lbl_selected.grid(row=0, column=1, padx=8, sticky="w")
        self._btn_seq = btn(tb, "▶  Start from Step 1", self._start_sequential,
                           color=GREEN, text_color=("#fff","#fff"),
                           width=190, height=32, size=12)
        self._auto_var = ctk.BooleanVar(value=False)
        self._cb_auto  = ctk.CTkCheckBox(
            tb, text="Auto-start next", variable=self._auto_var,
            font=(FONT, 11), checkbox_width=16, checkbox_height=16,
            command=lambda: setattr(self, "_auto_start", self._auto_var.get()))
        self._btn_finish = ctk.CTkButton(
            tb, text="Finish Session", command=self._maybe_finish,
            fg_color="transparent", text_color=T3, font=(FONT, 11),
            width=110, height=28, border_width=1, border_color=CARD_B,
            corner_radius=R_XS)
        self._btn_add_block = ctk.CTkButton(
            tb, text="＋ Block", command=self._add_temp_block,
            fg_color=("#e2e8f0","#334155"), text_color=T1,
            font=(FONT, 11), width=80, height=28, corner_radius=R_XS)

        # Left protocol panel
        self._proto_panel = ctk.CTkScrollableFrame(
            self, fg_color=("#f1f5f9","#1e293b"), corner_radius=R_XL, width=220)
        self._proto_panel.grid(row=1, column=0, sticky="nsew")
        self._proto_panel.grid_columnconfigure(0, weight=1)

        # Right main content
        self._main = ScrollFrame(self)
        self._main.grid(row=1, column=1, sticky="nsew")
        self._main.grid_columnconfigure(0, weight=1)

    # ── PROTOCOL PANEL ────────────────────────────────────────────────────────
    def _refresh_proto_panel(self):
        for w in self._proto_panel.winfo_children(): w.destroy()
        label(self._proto_panel, "Protocols", size=11, color=T3,
              weight="bold").pack(anchor="w", padx=12, pady=(10, 4))
        protocols = sorted(self.app.protocols,
                           key=lambda p: p.get("updatedAt", 0), reverse=True)
        if not protocols:
            label(self._proto_panel, "No protocols yet.", size=11,
                  color=T3).pack(anchor="w", padx=12, pady=4)
            return
        for p in protocols:
            self._proto_row(p)

    def _proto_row(self, proto):
        is_sel = bool(self.protocol and self.protocol["id"] == proto["id"])
        bg    = SB_ACT if is_sel else ("#e8eef4", "#1e293b")
        hover = SB_HOV
        tc    = ("#fff", "#fff") if is_sel else T1
        n     = len(proto.get("steps", []))
        t     = total_mins(proto.get("steps", []))
        name  = proto.get("name", "Untitled")
        info  = f"{n} steps  ·  {fmt_mins(t)}"
        # CTkButton gives reliable click regardless of child widget event absorption
        row_btn = ctk.CTkButton(
            self._proto_panel,
            text=f"{name}\n{info}",
            command=lambda pid=proto["id"]: self._select_proto(pid),
            fg_color=bg,
            hover_color=hover,
            text_color=tc,
            font=(FONT, 12),
            anchor="w",
            corner_radius=R_SM,
            height=54,
        )
        row_btn.pack(fill="x", padx=4, pady=2)
        if is_sel:
            for i, step in enumerate(proto.get("steps", [])):
                self._proto_step_row(step, i)

    def _proto_step_row(self, step, idx):
        st = (self._step_states[idx]["status"]
              if idx < len(self._step_states) else self.ST_IDLE)
        bg  = ("#dbeafe","#1e3a5f") if st == self.ST_RUNNING else "transparent"
        row = ctk.CTkFrame(self._proto_panel, fg_color=bg,
                           corner_radius=R_XS, cursor="hand2")
        row.pack(fill="x", padx=(18, 4), pady=1)
        row.bind("<Button-1>", lambda e, i=idx: self._scroll_to_step(i))
        dot_color = self.ST_COLOR.get(st, self.ST_COLOR[self.ST_IDLE])
        dot = ctk.CTkFrame(row, fg_color=dot_color,
                           width=18, height=18, corner_radius=R_SM)
        dot.pack(side="left", padx=(8, 5), pady=6)
        dot.bind("<Button-1>", lambda e, i=idx: self._scroll_to_step(i))
        label(dot, self.ST_ICON.get(st, "○"), size=8,
              color=("#fff","#fff")).place(relx=.5, rely=.5, anchor="center")
        m = (step.get("handsOnMinutes", 0) + step.get("waitMinutes", 0) +
             step.get("bufferMinutes", 0))
        _st = step.get('title') or 'Untitled'
        txt = f"{idx+1}. {(_st[:22] + '…') if len(_st) > 22 else _st}"
        lbl = label(row, txt, size=10, color=T1)
        lbl.pack(side="left", pady=6)
        lbl.bind("<Button-1>", lambda e, i=idx: self._scroll_to_step(i))
        if m:
            tl = label(row, fmt_mins(m), size=9, color=T3)
            tl.pack(side="right", padx=(0, 8))
            tl.bind("<Button-1>", lambda e, i=idx: self._scroll_to_step(i))

    def _select_proto(self, pid):
        proto = next((p for p in self.app.protocols if p["id"] == pid), None)
        if proto: self._select_proto_obj(proto)

    def _select_proto_obj(self, proto):
        self._cancel_all_timers()
        self.protocol = proto
        steps = proto.get("steps", [])
        self._step_states = []
        for s in steps:
            stype  = s.get("type", "custom")
            ho_m   = s.get("handsOnMinutes", 0)
            wt_m   = s.get("waitMinutes", 0)
            buf_m  = s.get("bufferMinutes", 0)
            # Countdown timer covers wait + buffer only; hands-on is an estimate
            if wt_m > 0 or STEP_TIMER_MODE.get(stype) == "countdown":
                countdown_secs = int((wt_m + buf_m) * 60)
            else:
                countdown_secs = 0
            self._step_states.append({
                "status":              self.ST_IDLE,
                "timer_secs":          float(countdown_secs),
                "original_secs":       countdown_secs,
                "adjusted_total_secs": countdown_secs,
                "elapsed_secs":        0,
                "timer_secs_at_start": countdown_secs,
                "adjusted":            False,
                "notes":               "",
                "timer_job":           None,
                "timer_lbl":           None,
                "status_dot":          None,
                "prog_var":            None,
                "btn_start":           None,
                "card":                None,
                # Wall-clock anchors — set on start/resume, used in tick
                "start_mono":          None,
                "start_remaining":     float(countdown_secs),
                # Optional hands-on stopwatch
                "ho_elapsed_secs":     0,
                "ho_timer_job":        None,
                "ho_timer_lbl":        None,
                "ho_btn":              None,
                "ho_status":           "idle",
                "ho_start_mono":       None,
                "ho_start_elapsed":    0,
                # Undo-complete state
                "_pre_complete_status": self.ST_IDLE,
                "_undo_job":            None,
                "_ctrl_norm":           None,
                "_ctrl_done":           None,
                "_undo_lbl":            None,
                "_ctrl_container":      None,
            })
        self._seq_mode           = False
        self._seq_idx            = -1
        self._temp_block_indices = []
        self._timeline           = []
        self._session_start_ts   = now_ts()
        self._log(f"Opened protocol: {proto.get('name','?')}")
        self._refresh_proto_panel()
        self._render_steps()
        self._show_topbar_controls()
        self._autosave()

    def _used_secs(self, sr):
        if sr["original_secs"] > 0:
            # Wall-clock: remaining is always accurate, derive elapsed from it
            elapsed = max(0, int(sr["original_secs"] - sr.get("timer_secs", sr["original_secs"])))
        else:
            elapsed = 0
        ho_elapsed = max(0, int(sr.get("ho_elapsed_secs", 0)))
        return elapsed + ho_elapsed

    def _apply_adjusted_time_to_step(self, step, adjusted_secs):
        adjusted_mins = max(0, round(adjusted_secs / 60))
        stype = step.get("type", "custom")
        if step.get("waitMinutes", 0) > 0 or STEP_TIMER_MODE.get(stype) == "countdown":
            # Adjust wait time; leave hands-on estimate unchanged
            step["waitMinutes"] = adjusted_mins
        else:
            step["handsOnMinutes"] = adjusted_mins
            step["waitMinutes"] = 0
        step["bufferMinutes"] = 0

    def _log(self, text):
        self._timeline.append({
            "ts":   now_ts(),
            "time": fmt_time(now_ts()),
            "text": text,
        })

    def _cancel_all_timers(self):
        for sr in self._step_states:
            if sr.get("timer_job"):
                self.after_cancel(sr["timer_job"])
                sr["timer_job"] = None
            if sr.get("ho_timer_job"):
                self.after_cancel(sr["ho_timer_job"])
                sr["ho_timer_job"] = None
            if sr.get("_undo_job"):
                self.after_cancel(sr["_undo_job"])
                sr["_undo_job"] = None

    def _show_topbar_controls(self):
        if self.protocol:
            self._lbl_selected.configure(text=self.protocol.get("name",""))
            self._btn_seq.grid(row=0, column=2, padx=(0, 8))
            self._cb_auto.grid(row=0, column=3, padx=(0, 8))
            self._btn_finish.grid(row=0, column=4, padx=(0, 16))
            self._btn_add_block.grid(row=0, column=5, padx=(0, 8))
        else:
            self._lbl_selected.configure(
                text="Choose a protocol from the left panel")
            for w in (self._btn_seq, self._cb_auto, self._btn_finish, self._btn_add_block):
                try: w.grid_remove()
                except Exception: pass

    # ── MAIN CONTENT ──────────────────────────────────────────────────────────
    def _show_empty(self):
        for w in self._main.winfo_children(): w.destroy()
        self._main.grid_columnconfigure(0, weight=1)
        label(self._main, "Choose a protocol to view and run step timers.",
              size=14, color=T3).grid(row=0, column=0, pady=100)
        self.after(80, self._rebind_scroll)

    def _render_steps(self):
        for w in self._main.winfo_children(): w.destroy()
        if not self.protocol: return
        steps = self.protocol.get("steps", [])
        if not steps:
            label(self._main, "This protocol has no steps.", size=13,
                  color=T3).grid(row=0, column=0, pady=60)
            return
        hdr = ctk.CTkFrame(self._main, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 8))
        label(hdr, f"{len(steps)} steps  ·  Total: {fmt_mins(total_mins(steps))}",
              size=13, color=T2).pack(side="left")
        for i, step in enumerate(steps):
            self._step_card(i, step)
        ctk.CTkFrame(self._main, height=40, fg_color="transparent").grid(
            row=len(steps)+1, column=0)
        self.after(80, self._rebind_scroll)

    def _step_card(self, idx, step):
        sr    = self._step_states[idx]
        stype = step.get("type","custom")
        bg_c  = STEP_COLORS.get(stype, STEP_COLORS["custom"])
        bd_c  = STEP_BADGES.get(stype,  STEP_BADGES["custom"])

        outer = ctk.CTkFrame(self._main, fg_color="transparent", corner_radius=R_MD)
        outer.grid(row=idx+1, column=0, sticky="ew", padx=20, pady=(0, 14))
        outer.grid_columnconfigure(0, weight=1)
        sr["card"] = outer

        card = ctk.CTkFrame(outer, fg_color=bg_c, corner_radius=R_LG,
                            border_width=2, border_color=bd_c)
        card.pack(fill="x")

        if step.get("tempBlock"):
            badge_row = ctk.CTkFrame(card, fg_color="transparent")
            badge_row.pack(fill="x", padx=16, pady=(6,0))
            chip = ctk.CTkFrame(badge_row, fg_color=("#fef3c7","#451a03"), corner_radius=R_XS)
            chip.pack(side="left")
            label(chip, "⚡ Added during run", size=10,
                  color=("#b45309","#fde68a")).pack(padx=8, pady=2)

        # Header
        hdr = ctk.CTkFrame(card, fg_color="transparent")
        hdr.pack(fill="x", padx=16, pady=(14, 6))
        num_f = ctk.CTkFrame(hdr, fg_color=bd_c, width=28, height=28, corner_radius=R_SM)
        num_f.pack(side="left", padx=(0, 8))
        num_f.pack_propagate(False)
        label(num_f, str(idx+1), size=12, weight="bold",
              color=("#fff","#fff")).place(relx=.5, rely=.5, anchor="center")
        tbadge = ctk.CTkFrame(hdr, fg_color=bd_c, corner_radius=R_XS)
        tbadge.pack(side="left")
        label(tbadge, STEP_LABELS.get(stype,""), size=10,
              color=("#fff","#fff")).pack(padx=6, pady=3)
        sr["status_dot"] = ctk.CTkFrame(hdr, fg_color=self.ST_COLOR[sr["status"]],
                                         width=10, height=10, corner_radius=R_XS)
        sr["status_dot"].pack(side="left", padx=8)
        planned_m = (step.get("handsOnMinutes",0) + step.get("waitMinutes",0) +
                     step.get("bufferMinutes",0))
        label(hdr, fmt_mins(planned_m), size=12, color=T2).pack(side="right")

        label(card, step.get("title") or "Untitled", size=16,
              weight="bold").pack(anchor="w", padx=16, pady=(0, 4))
        if step.get("description"):
            label(card, step["description"], size=13, color=T2).pack(
                anchor="w", padx=16, pady=(0, 6))

        # Conditions
        cond_items = (
            ([f"🌡 {step['temperature']}"] if step.get("temperature") else []) +
            ([f"🔄 {step['centrifugeCondition']}"] if step.get("centrifugeCondition") else []) +
            ([f"↻ {step['shakingRotation']}"] if step.get("shakingRotation") else []) +
            step.get("equipment", [])
        )
        if cond_items:
            cr = ctk.CTkFrame(card, fg_color="transparent")
            cr.pack(fill="x", padx=16, pady=(0, 8))
            for txt in cond_items:
                chip = ctk.CTkFrame(cr, fg_color=CARD, corner_radius=R_SM,
                                    border_width=1, border_color=CARD_B)
                chip.pack(side="left", padx=(0, 6))
                label(chip, txt, size=12).pack(padx=8, pady=4)

        # Time breakdown
        time_parts = []
        if step.get("handsOnMinutes"): time_parts.append(f"Hands-on: {fmt_mins(step['handsOnMinutes'])}")
        if step.get("waitMinutes"):    time_parts.append(f"Wait: {fmt_mins(step['waitMinutes'])}")
        if step.get("bufferMinutes"):  time_parts.append(f"Buffer: {fmt_mins(step['bufferMinutes'])}")
        if time_parts:
            label(card, "  ·  ".join(time_parts), size=11, color=T3).pack(
                anchor="w", padx=16, pady=(0, 8))

        # Reagents
        if step.get("reagents"):
            rg_f = ctk.CTkFrame(card, fg_color=CARD, corner_radius=R_SM,
                                 border_width=1, border_color=CARD_B)
            rg_f.pack(fill="x", padx=16, pady=(0, 10))
            label(rg_f, "Reagents", size=10, color=T2,
                  weight="bold").pack(anchor="w", padx=12, pady=(6, 2))
            for rg in step["reagents"]:
                t_str = f"• {reagent_display_text(rg)}"
                label(rg_f, t_str.strip(), size=12).pack(anchor="w", padx=16)
            ctk.CTkFrame(rg_f, height=6, fg_color="transparent").pack()

        # Warning
        if step.get("warnings"):
            warn_f = ctk.CTkFrame(card, fg_color=("#fef9c3","#422006"),
                                   corner_radius=R_SM, border_width=1,
                                   border_color=("#fde047","#854d0e"))
            warn_f.pack(fill="x", padx=16, pady=(0, 8))
            label(warn_f, f"⚠  {step['warnings']}", size=12,
                  color=("#92400e","#fde68a")).pack(anchor="w", padx=12, pady=6)

        # Checklist
        if step.get("checklist"):
            for c in step["checklist"]:
                v = ctk.BooleanVar(value=False)
                ctk.CTkCheckBox(card, text=c["text"], variable=v,
                                font=(FONT,12), checkbox_width=18,
                                checkbox_height=18).pack(anchor="w", padx=20, pady=2)
            ctk.CTkFrame(card, height=4, fg_color="transparent").pack()

        # Timer section — layout depends on step timer type
        timer_type = _step_timer_type(step)
        ho_m       = step.get("handsOnMinutes", 0)
        stype_key  = step.get("type", "custom")

        if timer_type == "countdown":
            timer_lbl_txt = TIMER_LABEL.get(stype_key, "Waiting timer")
            timer_f = ctk.CTkFrame(card, fg_color=CARD, corner_radius=R_MD,
                                   border_width=1, border_color=CARD_B)
            timer_f.pack(fill="x", padx=16, pady=(4, 12))
            t_inner = ctk.CTkFrame(timer_f, fg_color="transparent")
            t_inner.pack(fill="x", padx=16, pady=(14, 4))

            label(t_inner, timer_lbl_txt, size=11, color=T2,
                  weight="bold").pack(anchor="center")
            sr["timer_lbl"] = label(t_inner, fmt_secs(sr["timer_secs"]),
                                    size=48, weight="bold")
            sr["timer_lbl"].pack(anchor="center")
            orig_txt = (f"Protocol: {fmt_secs(sr['original_secs'])}"
                        if sr["original_secs"] else "No time set")
            label(t_inner, orig_txt, size=10, color=T3).pack(anchor="center")
            sr["prog_var"] = ctk.DoubleVar(value=1.0)
            if sr["original_secs"] > 0:
                ctk.CTkProgressBar(t_inner, variable=sr["prog_var"],
                                   height=6, corner_radius=R_XS).pack(fill="x", pady=(8, 0))
            if ho_m > 0:
                label(t_inner, f"Hands-on estimate: {fmt_mins(ho_m)}",
                      size=11, color=T2).pack(anchor="center", pady=(6, 0))

            # Adjustment row
            adj = ctk.CTkFrame(timer_f, fg_color="transparent")
            adj.pack(pady=(6, 10))
            label(adj, "Adjust:", size=10, color=T2).pack(side="left", padx=(0, 4))
            for dm, atxt in [(-5,"-5m"),(-1,"-1m"),(1,"+1m"),(5,"+5m")]:
                btn(adj, atxt, lambda i=idx, d=dm: self._adjust(i, d*60),
                    color=("#e2e8f0","#334155"), text_color=T1,
                    width=48, height=26, size=11).pack(side="left", padx=2)
            btn(adj, "Set…", lambda i=idx: self._set_custom(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=44, height=26, size=11).pack(side="left", padx=2)
            btn(adj, "Reset Timer", lambda i=idx: self._reset_timer(i),
                color=("#e2e8f0","#334155"), text_color=T2,
                width=72, height=26, size=11).pack(side="left", padx=2)

            # Controls — container swaps between _ctrl_norm and _ctrl_done
            ctrl_container = ctk.CTkFrame(card, fg_color="transparent")
            ctrl_container.pack(fill="x", padx=16, pady=(0, 12))
            sr["_ctrl_container"] = ctrl_container

            # Normal controls (shown when not completed)
            ctrl_norm = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            ctrl_norm.pack(fill="x")
            sr["_ctrl_norm"] = ctrl_norm
            sr["btn_start"] = btn(ctrl_norm, "▶  Start",
                                  lambda i=idx: self._start_step(i),
                                  color=ACC, text_color=("#fff","#fff"),
                                  width=120, height=42, size=13)
            sr["btn_start"].pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "✓  Complete", lambda i=idx: self._complete_step(i),
                color=GREEN, text_color=("#fff","#fff"),
                width=130, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "⤼  Skip", lambda i=idx: self._skip_step(i),
                color=("#e2e8f0","#334155"), text_color=T2,
                width=70, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=42, size=13).pack(side="left")

            # Done controls (shown after complete/skip — hidden initially)
            ctrl_done = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            sr["_ctrl_done"] = ctrl_done
            sr["_undo_lbl"] = label(ctrl_done, "✓ Completed — Undo (10s)",
                                    size=12, color=GREEN)
            sr["_undo_lbl"].pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↩  Undo", lambda i=idx: self._undo_complete(i),
                color=ORANGE, text_color=("#fff","#fff"),
                width=90, height=36, size=12).pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=36, size=12).pack(side="left")

        elif timer_type == "hands_on_only":
            # No countdown — show estimate and optional stopwatch
            timer_f = ctk.CTkFrame(card, fg_color=CARD, corner_radius=R_MD,
                                   border_width=1, border_color=CARD_B)
            timer_f.pack(fill="x", padx=16, pady=(4, 12))
            t_inner = ctk.CTkFrame(timer_f, fg_color="transparent")
            t_inner.pack(fill="x", padx=16, pady=(14, 10))

            label(t_inner, f"Hands-on estimate: {fmt_mins(ho_m)}",
                  size=22, weight="bold", color=T1).pack(anchor="center")
            label(t_inner, "No countdown — complete when finished",
                  size=11, color=T3).pack(anchor="center", pady=(2, 10))

            ho_row = ctk.CTkFrame(t_inner, fg_color="transparent")
            ho_row.pack(anchor="center")
            sr["ho_timer_lbl"] = label(ho_row, "00:00", size=18, color=T2)
            sr["ho_timer_lbl"].pack(side="left", padx=(0, 8))
            sr["ho_btn"] = btn(ho_row, "⏱  Track hands-on time",
                               lambda i=idx: self._start_ho(i),
                               color=("#e2e8f0","#334155"), text_color=T1,
                               width=170, height=32, size=11)
            sr["ho_btn"].pack(side="left")

            # Controls — container swaps between _ctrl_norm and _ctrl_done
            ctrl_container = ctk.CTkFrame(card, fg_color="transparent")
            ctrl_container.pack(fill="x", padx=16, pady=(0, 12))
            sr["_ctrl_container"] = ctrl_container

            ctrl_norm = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            ctrl_norm.pack(fill="x")
            sr["_ctrl_norm"] = ctrl_norm
            sr["btn_start"] = btn(ctrl_norm, "▶  Mark Active",
                                  lambda i=idx: self._start_step(i),
                                  color=("#e2e8f0","#334155"), text_color=T1,
                                  width=130, height=42, size=13)
            sr["btn_start"].pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "✓  Complete Step", lambda i=idx: self._complete_step(i),
                color=GREEN, text_color=("#fff","#fff"),
                width=150, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "⤼  Skip", lambda i=idx: self._skip_step(i),
                color=("#e2e8f0","#334155"), text_color=T2,
                width=70, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=42, size=13).pack(side="left")

            ctrl_done = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            sr["_ctrl_done"] = ctrl_done
            sr["_undo_lbl"] = label(ctrl_done, "✓ Completed — Undo (10s)",
                                    size=12, color=GREEN)
            sr["_undo_lbl"].pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↩  Undo", lambda i=idx: self._undo_complete(i),
                color=ORANGE, text_color=("#fff","#fff"),
                width=90, height=36, size=12).pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=36, size=12).pack(side="left")

        else:
            # No timer at all (note / checklist / decision)
            sr["btn_start"] = None
            ctrl_container = ctk.CTkFrame(card, fg_color="transparent")
            ctrl_container.pack(fill="x", padx=16, pady=(8, 12))
            sr["_ctrl_container"] = ctrl_container

            ctrl_norm = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            ctrl_norm.pack(fill="x")
            sr["_ctrl_norm"] = ctrl_norm
            btn(ctrl_norm, "✓  Complete", lambda i=idx: self._complete_step(i),
                color=GREEN, text_color=("#fff","#fff"),
                width=130, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "⤼  Skip", lambda i=idx: self._skip_step(i),
                color=("#e2e8f0","#334155"), text_color=T2,
                width=70, height=42, size=13).pack(side="left", padx=(0, 8))
            btn(ctrl_norm, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=42, size=13).pack(side="left")

            ctrl_done = ctk.CTkFrame(ctrl_container, fg_color="transparent")
            sr["_ctrl_done"] = ctrl_done
            sr["_undo_lbl"] = label(ctrl_done, "✓ Completed — Undo (10s)",
                                    size=12, color=GREEN)
            sr["_undo_lbl"].pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↩  Undo", lambda i=idx: self._undo_complete(i),
                color=ORANGE, text_color=("#fff","#fff"),
                width=90, height=36, size=12).pack(side="left", padx=(0, 8))
            btn(ctrl_done, "↺  Reset Block", lambda i=idx: self._reset_block(i),
                color=("#e2e8f0","#334155"), text_color=T1,
                width=110, height=36, size=12).pack(side="left")

        # Notes
        notes_wrap = ctk.CTkFrame(card, fg_color="transparent")
        notes_wrap.pack(fill="x", padx=16, pady=(4, 14))
        label(notes_wrap, "Notes", size=10, color=T3).pack(anchor="w", pady=(0, 3))
        nb = textbox(notes_wrap, width=600, height=44)
        nb.insert("0.0", sr.get("notes",""))
        nb.pack(fill="x")
        def save_note(ev=None, i=idx):
            self._step_states[i]["notes"] = nb.get("0.0","end").strip()
            self._autosave()
        nb.bind("<KeyRelease>", save_note)

        # Restore visual state from sr["status"] so a mid-run rebuild
        # (e.g. adding a temp block) doesn't reset all cards to idle.
        self._restore_step_visual(idx)

    # ── TIMER ENGINE ──────────────────────────────────────────────────────────
    def _start_step(self, idx):
        sr = self._step_states[idx]
        step_title   = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        is_countdown = sr["original_secs"] > 0
        if sr["status"] == self.ST_RUNNING:
            if is_countdown:
                # Snapshot accurate remaining BEFORE cancelling tick
                if sr.get("start_mono") is not None:
                    elapsed = time.monotonic() - sr["start_mono"]
                    sr["timer_secs"] = max(0.0,
                        sr.get("start_remaining", sr["timer_secs"]) - elapsed)
                if sr.get("timer_job"):
                    self.after_cancel(sr["timer_job"])
                    sr["timer_job"] = None
            sr["status"] = self.ST_PAUSED
            if sr.get("btn_start"):
                sr["btn_start"].configure(
                    text="▶  Resume" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)
            if sr.get("status_dot"):
                sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_PAUSED])
            self._log(f"Paused Step {idx+1}: {step_title}")
            self._refresh_proto_panel()
            self._autosave()
        elif sr["status"] in (self.ST_IDLE, self.ST_PAUSED):
            if sr["status"] == self.ST_IDLE:
                if is_countdown:
                    sr["timer_secs_at_start"] = sr["timer_secs"]
                self._log(f"Started Step {idx+1}: {step_title}")
            else:
                self._log(f"Resumed Step {idx+1}: {step_title}")
            sr["status"] = self.ST_RUNNING
            if sr.get("btn_start"):
                sr["btn_start"].configure(
                    text="⏸  Pause" if is_countdown else "⏸  Active",
                    fg_color=("#e2e8f0","#334155"), text_color=T1)
            if sr.get("status_dot"):
                sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_RUNNING])
            self._refresh_proto_panel()
            self._autosave()
            if is_countdown:
                # Anchor wall clock so tick computes remaining from real time
                sr["start_mono"]      = time.monotonic()
                sr["start_remaining"] = sr["timer_secs"]
                self._tick_step(idx)

    def _start_ho(self, idx):
        """Toggle the optional hands-on count-up stopwatch."""
        sr         = self._step_states[idx]
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        if sr.get("ho_status") == "running":
            # Snapshot accurate elapsed BEFORE cancelling tick
            if sr.get("ho_start_mono") is not None:
                elapsed = time.monotonic() - sr["ho_start_mono"]
                sr["ho_elapsed_secs"] = int(sr.get("ho_start_elapsed", 0) + elapsed)
            if sr.get("ho_timer_job"):
                self.after_cancel(sr["ho_timer_job"])
                sr["ho_timer_job"] = None
            sr["ho_status"] = "stopped"
            if sr.get("ho_btn"):
                sr["ho_btn"].configure(text="▶  Resume hands-on",
                                       fg_color=("#e2e8f0","#334155"),
                                       text_color=T1)
            self._log(f"Paused hands-on tracking — Step {idx+1}: {step_title}")
        else:
            if sr.get("ho_status") == "idle":
                self._log(f"Started hands-on tracking — Step {idx+1}: {step_title}")
                if sr["status"] == self.ST_IDLE:
                    sr["status"] = self.ST_RUNNING
                    if sr.get("status_dot"):
                        sr["status_dot"].configure(
                            fg_color=self.ST_COLOR[self.ST_RUNNING])
                    self._refresh_proto_panel()
            else:
                self._log(f"Resumed hands-on tracking — Step {idx+1}: {step_title}")
            sr["ho_status"]       = "running"
            # Anchor wall clock for count-up tick
            sr["ho_start_mono"]    = time.monotonic()
            sr["ho_start_elapsed"] = sr.get("ho_elapsed_secs", 0)
            if sr.get("ho_btn"):
                sr["ho_btn"].configure(text="⏸  Pause tracking",
                                       fg_color=ORANGE,
                                       text_color=("#fff","#fff"))
            self._tick_ho(idx)

    def _tick_ho(self, idx):
        """Count-up tick for the optional hands-on stopwatch (wall-clock based)."""
        sr = self._step_states[idx]
        if sr.get("ho_status") != "running":
            return
        elapsed = time.monotonic() - sr.get("ho_start_mono", time.monotonic())
        sr["ho_elapsed_secs"] = int(sr.get("ho_start_elapsed", 0) + elapsed)
        if sr.get("ho_timer_lbl"):
            sr["ho_timer_lbl"].configure(text=fmt_secs(sr["ho_elapsed_secs"]))
        sr["ho_timer_job"] = self.after(500, lambda i=idx: self._tick_ho(i))

    def _tick_step(self, idx):
        sr = self._step_states[idx]
        if sr["status"] != self.ST_RUNNING:
            return
        # Derive remaining from real wall-clock time so it stays accurate
        # across sleep/wake, minimize, screen lock, or any UI freeze.
        elapsed   = time.monotonic() - sr.get("start_mono", time.monotonic())
        remaining = sr.get("start_remaining", sr["timer_secs"]) - elapsed
        sr["timer_secs"] = max(0.0, remaining)
        if sr["timer_secs"] <= 0:
            sr["timer_secs"] = 0.0
            self._update_timer_display(idx)
            self._on_step_expired(idx)
            return
        self._update_timer_display(idx)
        # Refresh every 500 ms — quick enough to look smooth, cheap enough
        # to not affect app performance.
        sr["timer_job"] = self.after(500, lambda i=idx: self._tick_step(i))

    def _on_step_expired(self, idx):
        sr = self._step_states[idx]
        sr["timer_job"] = None
        sr["status"]    = self.ST_PAUSED
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        self._log(f"Timer expired — Step {idx+1}: {step_title}")
        if sr.get("btn_start"):
            sr["btn_start"].configure(text="▶  Resume", fg_color=ACC,
                                      text_color=("#fff","#fff"))
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_PAUSED])
        if sr.get("timer_lbl"):
            sr["timer_lbl"].configure(text="00:00", text_color=DANGER)
        self._refresh_proto_panel()
        try: self.bell()
        except Exception: pass

    def _update_timer_display(self, idx):
        sr = self._step_states[idx]
        if not sr.get("timer_lbl"): return
        secs = sr["timer_secs"]
        col  = DANGER if secs <= 0 else (ORANGE if secs <= 60 else T1)
        sr["timer_lbl"].configure(text=fmt_secs(secs), text_color=col)
        planned = sr.get("adjusted_total_secs", sr["original_secs"])
        if planned > 0 and sr.get("prog_var"):
            sr["prog_var"].set(min(max(secs / planned, 0.0), 1.0))

    def _adjust(self, idx, delta_secs):
        sr = self._step_states[idx]
        # If running, get the accurate real-time remaining first
        if sr["status"] == self.ST_RUNNING and sr.get("start_mono") is not None:
            elapsed   = time.monotonic() - sr["start_mono"]
            current   = sr.get("start_remaining", sr["timer_secs"]) - elapsed
            sr["timer_secs"] = max(0.0, current + delta_secs)
        else:
            sr["timer_secs"] = max(0.0, sr["timer_secs"] + delta_secs)
        # Re-anchor wall clock so next tick is based on the new value
        sr["start_mono"]      = time.monotonic()
        sr["start_remaining"] = sr["timer_secs"]
        sr["adjusted_total_secs"] = max(0, sr.get("adjusted_total_secs", sr["original_secs"]) + delta_secs)
        sr["adjusted"]   = True
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        sign = "+" if delta_secs >= 0 else ""
        self._log(f"Step {idx+1} timer adjusted {sign}{delta_secs//60}m — {step_title}")
        self._update_timer_display(idx)

    def _set_custom(self, idx):
        dialog = ctk.CTkInputDialog(text="Set timer (MM:SS or minutes):",
                                    title="Set Timer")
        val = dialog.get_input()
        if not val: return
        try:
            val = val.strip()
            if ":" in val:
                p = val.split(":")
                secs = int(p[0]) * 60 + int(p[1])
            else:
                secs = int(float(val) * 60)
        except (ValueError, IndexError):
            return
        sr = self._step_states[idx]
        # Elapsed up to this moment = original - current remaining
        elapsed_so_far = max(0, sr["original_secs"] - sr.get("timer_secs", sr["original_secs"]))
        sr["timer_secs"]          = max(0.0, float(secs))
        sr["start_mono"]          = time.monotonic()
        sr["start_remaining"]     = sr["timer_secs"]
        sr["adjusted_total_secs"] = elapsed_so_far + secs
        sr["adjusted"]            = True
        self._update_timer_display(idx)

    def _reset_timer(self, idx):
        sr = self._step_states[idx]
        if sr.get("timer_job"):
            self.after_cancel(sr["timer_job"])
            sr["timer_job"] = None
        if sr.get("ho_timer_job"):
            self.after_cancel(sr["ho_timer_job"])
            sr["ho_timer_job"] = None
        sr["timer_secs"]         = float(sr["original_secs"])
        sr["adjusted_total_secs"] = sr["original_secs"]
        sr["elapsed_secs"]       = 0
        sr["adjusted"]           = False
        sr["start_mono"]         = None
        sr["start_remaining"]    = float(sr["original_secs"])
        sr["ho_elapsed_secs"]    = 0
        sr["ho_status"]          = "idle"
        sr["ho_start_mono"]      = None
        sr["ho_start_elapsed"]   = 0
        if sr.get("ho_timer_lbl"):
            sr["ho_timer_lbl"].configure(text="00:00")
        if sr.get("ho_btn"):
            sr["ho_btn"].configure(text="⏱  Track hands-on time",
                                   fg_color=("#e2e8f0","#334155"),
                                   text_color=T1)
        # Only update button/dot if the step is not already completed or skipped
        if sr["status"] not in (self.ST_COMPLETED, self.ST_SKIPPED):
            sr["status"] = self.ST_IDLE
            if sr.get("btn_start"):
                is_countdown = sr["original_secs"] > 0
                sr["btn_start"].configure(
                    text="▶  Start" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)
            if sr.get("status_dot"):
                sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_IDLE])
        self._update_timer_display(idx)
        self._refresh_proto_panel()

    def _complete_step(self, idx):
        sr = self._step_states[idx]
        # Cancel running timers
        if sr.get("timer_job"):
            self.after_cancel(sr["timer_job"])
            sr["timer_job"] = None
        if sr.get("ho_timer_job"):
            self.after_cancel(sr["ho_timer_job"])
            sr["ho_timer_job"] = None
        # Save pre-complete status for undo
        sr["_pre_complete_status"] = sr["status"]
        sr["status"] = self.ST_COMPLETED
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        used = self._used_secs(sr)
        self._log(f"Completed Step {idx+1}: {step_title}  ({fmt_secs(used)} used)")
        # Update visual indicators
        if sr.get("btn_start"):
            sr["btn_start"].configure(text="✓ Done", fg_color=GREEN,
                                      text_color=("#fff","#fff"))
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_COMPLETED])
        if sr.get("timer_lbl"):
            sr["timer_lbl"].configure(text_color=GREEN)
        # Swap to done controls with undo countdown
        if sr.get("_ctrl_norm") and sr.get("_ctrl_done"):
            sr["_ctrl_norm"].pack_forget()
            sr["_ctrl_done"].pack(fill="x")
        if sr.get("_undo_lbl"):
            sr["_undo_lbl"].configure(text="✓ Completed — Undo (10s)", text_color=GREEN)
        # Schedule undo window — cancel after 10 seconds
        if sr.get("_undo_job"):
            self.after_cancel(sr["_undo_job"])
        sr["_undo_job"] = self.after(10000, lambda i=idx: self._confirm_complete(i))
        self._undo_countdown_tick(idx, 9)
        self._refresh_proto_panel()
        self._autosave()
        # Only ask to save adjusted time for countdown steps with actual timers
        if sr["adjusted"] and sr["original_secs"] > 0:
            self._ask_save_adjusted(idx)
        if self._seq_mode and idx == self._seq_idx:
            self._advance_sequential()

    def _skip_step(self, idx):
        """Mark step as skipped (instant, with undo window)."""
        sr = self._step_states[idx]
        if sr.get("timer_job"):
            self.after_cancel(sr["timer_job"])
            sr["timer_job"] = None
        if sr.get("ho_timer_job"):
            self.after_cancel(sr["ho_timer_job"])
            sr["ho_timer_job"] = None
        sr["_pre_complete_status"] = sr["status"]
        sr["status"] = self.ST_SKIPPED
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        self._log(f"Skipped Step {idx+1}: {step_title}")
        if sr.get("btn_start"):
            sr["btn_start"].configure(text="⤼ Skipped",
                                      fg_color=self.ST_COLOR[self.ST_SKIPPED],
                                      text_color=("#fff","#fff"))
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_SKIPPED])
        if sr.get("timer_lbl"):
            sr["timer_lbl"].configure(text_color=self.ST_COLOR[self.ST_SKIPPED][0])
        # Swap to done controls with undo countdown
        if sr.get("_ctrl_norm") and sr.get("_ctrl_done"):
            sr["_ctrl_norm"].pack_forget()
            sr["_ctrl_done"].pack(fill="x")
        if sr.get("_undo_lbl"):
            sr["_undo_lbl"].configure(text="⤼ Skipped — Undo (10s)",
                                      text_color=self.ST_COLOR[self.ST_SKIPPED][0])
        if sr.get("_undo_job"):
            self.after_cancel(sr["_undo_job"])
        sr["_undo_job"] = self.after(10000, lambda i=idx: self._confirm_complete(i))
        self._undo_countdown_tick(idx, 9)
        self._refresh_proto_panel()
        self._autosave()
        if self._seq_mode and idx == self._seq_idx:
            self._advance_sequential()

    def _undo_complete(self, idx):
        """Undo a completion or skip within the 10-second window."""
        sr = self._step_states[idx]
        if sr["status"] not in (self.ST_COMPLETED, self.ST_SKIPPED):
            return
        was_status = sr["status"]
        # Cancel undo expiry timer
        if sr.get("_undo_job"):
            self.after_cancel(sr["_undo_job"])
            sr["_undo_job"] = None
        # Restore previous status
        prev = sr.get("_pre_complete_status", self.ST_IDLE)
        sr["status"] = prev
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        action = "Completion" if was_status == self.ST_COMPLETED else "Skip"
        self._log(f"{action} undone — Step {idx+1}: {step_title}")
        # Swap back to normal controls
        if sr.get("_ctrl_done") and sr.get("_ctrl_norm"):
            sr["_ctrl_done"].pack_forget()
            sr["_ctrl_norm"].pack(fill="x")
        # Restore visual indicators
        is_countdown = sr["original_secs"] > 0
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[prev])
        if sr.get("timer_lbl"):
            sr["timer_lbl"].configure(text_color=T1)
        if sr.get("btn_start"):
            if prev == self.ST_RUNNING:
                # Step was running — restore running appearance but keep paused (safe)
                sr["status"] = self.ST_PAUSED
                sr["btn_start"].configure(
                    text="▶  Resume" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)
                if sr.get("status_dot"):
                    sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_PAUSED])
            elif prev == self.ST_PAUSED:
                sr["btn_start"].configure(
                    text="▶  Resume" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)
            else:  # IDLE
                sr["btn_start"].configure(
                    text="▶  Start" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)
        self._refresh_proto_panel()

    def _confirm_complete(self, idx):
        """Called after 10-second undo window expires — lock in completion."""
        sr = self._step_states[idx]
        sr["_undo_job"] = None
        if sr.get("_undo_lbl"):
            if sr["status"] == self.ST_COMPLETED:
                sr["_undo_lbl"].configure(text="✓ Done", text_color=GREEN)
            elif sr["status"] == self.ST_SKIPPED:
                sr["_undo_lbl"].configure(text="⤼ Skipped",
                                          text_color=self.ST_COLOR[self.ST_SKIPPED][0])

    def _undo_countdown_tick(self, idx, remaining):
        """Update the undo countdown label each second during the 10s window."""
        sr = self._step_states[idx]
        if sr["status"] not in (self.ST_COMPLETED, self.ST_SKIPPED):
            return
        if sr.get("_undo_lbl"):
            if sr["status"] == self.ST_COMPLETED:
                sr["_undo_lbl"].configure(
                    text=f"✓ Completed — Undo ({remaining}s)", text_color=GREEN)
            else:
                sr["_undo_lbl"].configure(
                    text=f"⤼ Skipped — Undo ({remaining}s)",
                    text_color=self.ST_COLOR[self.ST_SKIPPED][0])
        if remaining > 1:
            self.after(1000, lambda i=idx, r=remaining-1: self._undo_countdown_tick(i, r))

    def _reset_block(self, idx):
        """Full block reset: clears all state, cancels timers, keeps notes."""
        sr = self._step_states[idx]
        # Cancel all pending jobs
        for key in ("timer_job", "ho_timer_job", "_undo_job"):
            if sr.get(key):
                self.after_cancel(sr[key])
                sr[key] = None
        # Reset status and all timer state
        sr["status"]               = self.ST_IDLE
        sr["timer_secs"]           = float(sr["original_secs"])
        sr["adjusted_total_secs"]  = sr["original_secs"]
        sr["elapsed_secs"]         = 0
        sr["adjusted"]             = False
        sr["start_mono"]           = None
        sr["start_remaining"]      = float(sr["original_secs"])
        sr["ho_elapsed_secs"]      = 0
        sr["ho_status"]            = "idle"
        sr["ho_start_mono"]        = None
        sr["ho_start_elapsed"]     = 0
        sr["_pre_complete_status"] = self.ST_IDLE
        # Reset optional HO stopwatch widgets
        if sr.get("ho_timer_lbl"):
            sr["ho_timer_lbl"].configure(text="00:00")
        if sr.get("ho_btn"):
            sr["ho_btn"].configure(text="⏱  Track hands-on time",
                                   fg_color=("#e2e8f0","#334155"), text_color=T1)
        # Swap back to normal controls
        if sr.get("_ctrl_done") and sr.get("_ctrl_norm"):
            sr["_ctrl_done"].pack_forget()
            sr["_ctrl_norm"].pack(fill="x")
        # Restore button/dot to idle
        is_countdown = sr["original_secs"] > 0
        if sr.get("btn_start"):
            sr["btn_start"].configure(
                text="▶  Start" if is_countdown else "▶  Mark Active",
                fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                text_color=("#fff","#fff") if is_countdown else T1)
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[self.ST_IDLE])
        if sr.get("timer_lbl"):
            sr["timer_lbl"].configure(text_color=T1)
        self._update_timer_display(idx)
        step_title = self.protocol["steps"][idx].get("title","?") if self.protocol else "?"
        self._log(f"Reset Block — Step {idx+1}: {step_title}")
        self._refresh_proto_panel()
        self._autosave()

    def _restore_step_visual(self, idx):
        """Re-apply the correct visual state after step cards are rebuilt.

        Called at the end of _step_card() so that steps which were running /
        paused / completed retain their appearance when the full card list is
        recreated (e.g. after adding a temp block mid-run).
        """
        sr = self._step_states[idx]
        st = sr["status"]
        is_countdown = sr["original_secs"] > 0

        # Status dot
        if sr.get("status_dot"):
            sr["status_dot"].configure(fg_color=self.ST_COLOR[st])

        if st in (self.ST_COMPLETED, self.ST_SKIPPED):
            # Swap to done-controls panel
            if sr.get("_ctrl_norm") and sr.get("_ctrl_done"):
                sr["_ctrl_norm"].pack_forget()
                sr["_ctrl_done"].pack(fill="x")
            if sr.get("_undo_lbl"):
                if sr.get("_undo_job"):
                    # Undo window still open — show generic "Undo" prompt
                    action = "✓ Completed" if st == self.ST_COMPLETED else "⤼ Skipped"
                    col    = GREEN if st == self.ST_COMPLETED else self.ST_COLOR[self.ST_SKIPPED][0]
                    sr["_undo_lbl"].configure(text=f"{action} — Undo", text_color=col)
                else:
                    if st == self.ST_COMPLETED:
                        sr["_undo_lbl"].configure(text="✓ Done", text_color=GREEN)
                    else:
                        sr["_undo_lbl"].configure(
                            text="⤼ Skipped",
                            text_color=self.ST_COLOR[self.ST_SKIPPED][0])
        elif st == self.ST_RUNNING:
            if sr.get("btn_start"):
                sr["btn_start"].configure(
                    text="⏸  Pause" if is_countdown else "⏸  Active",
                    fg_color=("#e2e8f0","#334155"), text_color=T1)
        elif st == self.ST_PAUSED:
            if sr.get("btn_start"):
                sr["btn_start"].configure(
                    text="▶  Resume" if is_countdown else "▶  Mark Active",
                    fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                    text_color=("#fff","#fff") if is_countdown else T1)

        # Restore HO stopwatch button label
        if sr.get("ho_status") == "running" and sr.get("ho_btn"):
            sr["ho_btn"].configure(text="⏸  Pause tracking",
                                   fg_color=ORANGE, text_color=("#fff","#fff"))
        elif sr.get("ho_status") == "stopped" and sr.get("ho_btn"):
            sr["ho_btn"].configure(text="▶  Resume hands-on",
                                   fg_color=("#e2e8f0","#334155"), text_color=T1)

        # Refresh countdown display (label text + progress bar)
        self._update_timer_display(idx)
        if sr.get("ho_timer_lbl"):
            sr["ho_timer_lbl"].configure(
                text=fmt_secs(sr.get("ho_elapsed_secs", 0)))

    def _ask_save_adjusted(self, idx):
        sr   = self._step_states[idx]
        step = self.protocol["steps"][idx]
        adj_total = sr.get("adjusted_total_secs", sr["original_secs"])
        orig_m = round(sr["original_secs"] / 60, 1)
        adj_m  = round(adj_total / 60, 1)

        d = ctk.CTkToplevel(self)
        d.withdraw()
        d.title("Save Adjusted Time?")
        d.resizable(False, False)

        f = ctk.CTkFrame(d, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=20)
        label(f, f"Timer was adjusted for:", size=13, weight="bold").pack(anchor="w")
        label(f, step.get("title","?"), size=13, color=T2).pack(anchor="w", pady=(0, 4))
        label(f, f"Original: {fmt_mins(orig_m)}  →  Adjusted to: {fmt_mins(adj_m)}",
              size=12, color=T2).pack(anchor="w", pady=(0, 14))

        def choose(val):
            if val == "update":
                self._apply_adjusted_time_to_step(step, adj_total)
                self.protocol["updatedAt"] = now_ts()
                save_protocols(self.app.protocols)
            elif val == "copy":
                import copy as _copy
                np2 = _copy.deepcopy(self.protocol)
                np2["id"]   = new_id()
                np2["name"] = np2.get("name","") + " (v2)"
                self._apply_adjusted_time_to_step(np2["steps"][idx], adj_total)
                np2["createdAt"] = now_ts()
                np2["updatedAt"] = now_ts()
                self.app.protocols.insert(0, np2)
                save_protocols(self.app.protocols)
                self._refresh_proto_panel()
            elif val == "default":
                self._apply_adjusted_time_to_step(step, adj_total)
                self.protocol["updatedAt"] = now_ts()
                save_protocols(self.app.protocols)
            d.destroy()

        for txt, val in [
            ("Keep only for this session",         "keep"),
            ("Update this step in the protocol",   "update"),
            ("Save as a new protocol version",     "copy"),
            ("Save as my personal default speed",  "default"),
        ]:
            btn(f, txt, lambda v=val: choose(v),
                color=ACC if val=="keep" else ("#e2e8f0","#334155"),
                text_color=("#fff","#fff") if val=="keep" else T1,
                width=370, height=36, size=12).pack(anchor="w", pady=3)
        _show_dialog(d, self, 430, 230)

    # ── SEQUENTIAL MODE ────────────────────────────────────────────────────────
    def _start_sequential(self):
        if not self._step_states: return
        self._seq_mode = True
        self._seq_idx  = 0
        self._cancel_all_timers()
        for i, sr in enumerate(self._step_states):
            if sr["status"] not in (self.ST_COMPLETED, self.ST_SKIPPED):
                if sr.get("ho_timer_job"):
                    self.after_cancel(sr["ho_timer_job"])
                    sr["ho_timer_job"] = None
                if sr.get("_undo_job"):
                    self.after_cancel(sr["_undo_job"])
                    sr["_undo_job"] = None
                sr["status"]               = self.ST_IDLE
                sr["_pre_complete_status"] = self.ST_IDLE
                sr["timer_secs"]           = float(sr["original_secs"])
                sr["adjusted_total_secs"]  = sr["original_secs"]
                sr["elapsed_secs"]         = 0
                sr["adjusted"]             = False
                sr["start_mono"]           = None
                sr["start_remaining"]      = float(sr["original_secs"])
                sr["ho_elapsed_secs"]      = 0
                sr["ho_status"]            = "idle"
                sr["ho_start_mono"]        = None
                sr["ho_start_elapsed"]     = 0
                if sr.get("ho_timer_lbl"):
                    sr["ho_timer_lbl"].configure(text="00:00")
                if sr.get("ho_btn"):
                    sr["ho_btn"].configure(text="⏱  Track hands-on time",
                                           fg_color=("#e2e8f0","#334155"),
                                           text_color=T1)
                # Restore normal ctrl frame if it was swapped to done
                if sr.get("_ctrl_done") and sr.get("_ctrl_norm"):
                    sr["_ctrl_done"].pack_forget()
                    sr["_ctrl_norm"].pack(fill="x")
                is_countdown = sr["original_secs"] > 0
                if sr.get("btn_start"):
                    sr["btn_start"].configure(
                        text="▶  Start" if is_countdown else "▶  Mark Active",
                        fg_color=ACC if is_countdown else ("#e2e8f0","#334155"),
                        text_color=("#fff","#fff") if is_countdown else T1)
                if sr.get("status_dot"):
                    sr["status_dot"].configure(
                        fg_color=self.ST_COLOR[self.ST_IDLE])
                self._update_timer_display(i)
        self._start_step(0)
        self._scroll_to_step(0)

    def _advance_sequential(self):
        nxt = self._seq_idx + 1
        if nxt >= len(self._step_states):
            self._seq_mode = False
            self._seq_idx  = -1
            return
        self._seq_idx = nxt
        self._scroll_to_step(nxt)
        if self._auto_start:
            self._start_step(nxt)

    def _scroll_to_step(self, idx):
        sr = (self._step_states[idx]
              if idx < len(self._step_states) else None)
        if not sr or not sr.get("card"): return
        try:
            self.update_idletasks()
            canvas = self._main._parent_canvas
            fh = self._main.winfo_height()
            if fh > 0:
                frac = sr["card"].winfo_y() / fh
                canvas.yview_moveto(max(0.0, frac - 0.05))
        except Exception:
            pass

    # ── SESSION PERSISTENCE (crash recovery) ─────────────────────────────────
    _AUTOSAVE_DELAY_MS = 1500   # debounce: at most one write per 1.5 s
    _autosave_job: "int | None" = None

    def _autosave(self):
        """Debounce-save: schedules a checkpoint write in 1.5 s.
        If called again before the timer fires, the previous one is cancelled
        so we never spam the disk during rapid state changes."""
        if getattr(self, "_autosave_job", None):
            self.after_cancel(self._autosave_job)
        self._autosave_job = self.after(
            self._AUTOSAVE_DELAY_MS, self._write_runtime_session)

    def _write_runtime_session(self):
        """Serialize current run state to RUNTIME_FILE for crash recovery."""
        self._autosave_job = None
        if not self.protocol:
            _discard_runtime_session()
            return
        step_states_data = []
        for sr in self._step_states:
            # Snapshot accurate remaining for running timers
            if sr["status"] == self.ST_RUNNING and sr.get("start_mono") is not None:
                elapsed = time.monotonic() - sr["start_mono"]
                live_remaining = max(0.0,
                    sr.get("start_remaining", sr["timer_secs"]) - elapsed)
            else:
                live_remaining = sr["timer_secs"]
            step_states_data.append({
                "status":              sr["status"],
                "timer_secs":          live_remaining,
                "original_secs":       sr["original_secs"],
                "adjusted_total_secs": sr.get("adjusted_total_secs", sr["original_secs"]),
                "elapsed_secs":        sr.get("elapsed_secs", 0),
                "timer_secs_at_start": sr.get("timer_secs_at_start", sr["original_secs"]),
                "adjusted":            sr.get("adjusted", False),
                "notes":               sr.get("notes", ""),
                "ho_elapsed_secs":     sr.get("ho_elapsed_secs", 0),
                "ho_status":           sr.get("ho_status", "idle"),
                "is_temp":             sr.get("is_temp", False),
            })
        payload = {
            "version":           1,
            "saved_at_ts":       now_ts(),
            "protocol_id":       self.protocol["id"],
            "protocol_snapshot": self.protocol,   # includes temp blocks added mid-run
            "session_start_ts":  self._session_start_ts,
            "timeline":          self._timeline,
            "step_states":       step_states_data,
            "seq_mode":          self._seq_mode,
            "seq_idx":           self._seq_idx,
        }
        try:
            RUNTIME_FILE.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass   # never crash the app due to a save failure

    def _offer_resume(self):
        """Shown on startup when a crash-recovery file exists."""
        try:
            data = json.loads(RUNTIME_FILE.read_text(encoding="utf-8"))
        except Exception:
            _discard_runtime_session()
            return
        if data.get("version") != 1:
            _discard_runtime_session()
            return
        proto_name = (data.get("protocol_snapshot") or {}).get("name", "?")
        saved_ts   = data.get("saved_at_ts", 0)
        try:
            saved_str = datetime.fromtimestamp(
                saved_ts / 1000).strftime("%Y-%m-%d %H:%M")
        except Exception:
            saved_str = "?"

        d = ctk.CTkToplevel(self)
        d.withdraw()
        d.title("Resume Previous Session?")
        d.resizable(False, False)
        f = ctk.CTkFrame(d, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=24, pady=24)
        label(f, "Resume previous session?", size=15,
              weight="bold").pack(anchor="w")
        label(f, f"Protocol:   {proto_name}",
              size=13, color=T2).pack(anchor="w", pady=(8, 2))
        label(f, f"Last saved:  {saved_str}",
              size=12, color=T3).pack(anchor="w", pady=(0, 16))

        def _resume():
            d.destroy()
            self._restore_from_runtime(data)

        def _discard():
            d.destroy()
            _discard_runtime_session()

        btn(f, "▶  Resume Session", _resume,
            color=GREEN, text_color=("#fff","#fff"),
            width=200, height=42).pack(anchor="w", pady=(0, 8))
        btn(f, "Discard", _discard,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=110, height=38).pack(anchor="w")
        _show_dialog(d, self, 420, 260)

    def _restore_from_runtime(self, data):
        """Rebuild RunPage state from a crash-recovery snapshot."""
        proto = data.get("protocol_snapshot")
        if not proto:
            _discard_runtime_session()
            return

        self._cancel_all_timers()
        self.protocol           = proto
        self._session_start_ts  = data.get("session_start_ts") or now_ts()
        self._timeline          = list(data.get("timeline", []))
        self._seq_mode          = data.get("seq_mode", False)
        self._seq_idx           = int(data.get("seq_idx", -1))
        self._temp_block_indices = []

        now_ts_val           = now_ts()
        saved_ts             = data.get("saved_at_ts", now_ts_val)
        elapsed_since_save   = max(0.0, (now_ts_val - saved_ts) / 1000.0)  # seconds

        saved_states = data.get("step_states", [])
        steps        = proto.get("steps", [])
        self._step_states = []

        for i, step in enumerate(steps):
            saved = saved_states[i] if i < len(saved_states) else {}
            stype  = step.get("type", "custom")
            wt_m   = step.get("waitMinutes", 0)
            buf_m  = step.get("bufferMinutes", 0)
            if wt_m > 0 or STEP_TIMER_MODE.get(stype) == "countdown":
                default_secs = int((wt_m + buf_m) * 60)
            else:
                default_secs = 0

            saved_status = saved.get("status", self.ST_IDLE)
            saved_secs   = float(saved.get("timer_secs", default_secs))

            # Adjust for time elapsed while the app was closed
            if saved_status == self.ST_RUNNING:
                restored_secs   = max(0.0, saved_secs - elapsed_since_save)
                # Restore as PAUSED — user resumes manually; avoids silent overrun
                restored_status = self.ST_PAUSED
            else:
                restored_secs   = saved_secs
                restored_status = saved_status

            if saved.get("is_temp"):
                self._temp_block_indices.append(i)

            self._step_states.append({
                "status":               restored_status,
                "timer_secs":           float(restored_secs),
                "original_secs":        int(saved.get("original_secs", default_secs)),
                "adjusted_total_secs":  int(saved.get("adjusted_total_secs", default_secs)),
                "elapsed_secs":         int(saved.get("elapsed_secs", 0)),
                "timer_secs_at_start":  float(saved.get("timer_secs_at_start", default_secs)),
                "adjusted":             bool(saved.get("adjusted", False)),
                "notes":                saved.get("notes", ""),
                "timer_job":            None,
                "timer_lbl":            None,
                "status_dot":           None,
                "prog_var":             None,
                "btn_start":            None,
                "card":                 None,
                "is_temp":              bool(saved.get("is_temp", False)),
                "start_mono":           None,
                "start_remaining":      float(restored_secs),
                "ho_elapsed_secs":      int(saved.get("ho_elapsed_secs", 0)),
                "ho_timer_job":         None,
                "ho_timer_lbl":         None,
                "ho_btn":               None,
                "ho_status":            "idle",   # hands-on always resets on restore
                "ho_start_mono":        None,
                "ho_start_elapsed":     0,
                "_pre_complete_status": self.ST_IDLE,
                "_undo_job":            None,
                "_ctrl_norm":           None,
                "_ctrl_done":           None,
                "_undo_lbl":            None,
                "_ctrl_container":      None,
            })

        self._log(
            f"Session restored (was saved {datetime.fromtimestamp(saved_ts/1000).strftime('%H:%M')})")
        _discard_runtime_session()   # clear the file — we've consumed it
        self._refresh_proto_panel()
        self._render_steps()
        self._show_topbar_controls()
        self._autosave()   # write fresh checkpoint immediately

    # ── FINISH SESSION ─────────────────────────────────────────────────────────
    def _maybe_finish(self):
        if not self.protocol: return
        incomplete = [i for i, sr in enumerate(self._step_states)
                      if sr["status"] not in (self.ST_COMPLETED, self.ST_SKIPPED)]
        if incomplete:
            self._show_finish_dialog(incomplete)
        else:
            self._cancel_all_timers()
            self._log("Session finished")
            self._show_save_session_dialog()

    def _show_finish_dialog(self, incomplete_indices):
        """Show 3-option finish dialog when not all steps are complete/skipped."""
        d = ctk.CTkToplevel(self)
        d.withdraw()
        d.title("Finish Session")
        d.resizable(False, False)

        f = ctk.CTkFrame(d, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=20)

        n_incomplete = len(incomplete_indices)
        n_total      = len(self._step_states)
        label(f, "Finish Session", size=14, weight="bold").pack(anchor="w")
        label(f, f"{n_incomplete} of {n_total} steps are incomplete.",
              size=12, color=T2).pack(anchor="w", pady=(4, 14))

        def _cancel():
            d.destroy()

        def _save_incomplete():
            d.destroy()
            self._cancel_all_timers()
            self._log("Session finished (incomplete)")
            self._show_save_session_dialog()

        def _skip_and_save():
            for i in incomplete_indices:
                sr = self._step_states[i]
                if sr.get("_undo_job"):
                    self.after_cancel(sr["_undo_job"])
                    sr["_undo_job"] = None
                sr["status"] = self.ST_SKIPPED
            d.destroy()
            self._cancel_all_timers()
            self._log("Session finished (remaining marked as skipped)")
            self._show_save_session_dialog()

        btn(f, "Cancel — continue session", _cancel,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=340, height=38, size=12).pack(anchor="w", pady=(0, 6))
        btn(f, "Save as incomplete session", _save_incomplete,
            color=ACC, text_color=("#fff","#fff"),
            width=340, height=38, size=12).pack(anchor="w", pady=(0, 6))
        btn(f, f"Mark {n_incomplete} remaining as skipped and save", _skip_and_save,
            color=("#64748b","#475569"), text_color=("#fff","#fff"),
            width=340, height=38, size=12).pack(anchor="w")
        _show_dialog(d, self, 400, 230)

    def _show_save_session_dialog(self):
        d = ctk.CTkToplevel(self)
        d.withdraw()
        d.title("Save Session to Lab Notebook")
        d.resizable(False, False)

        f = ctk.CTkFrame(d, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=20)

        label(f, "Save to Lab Notebook", size=14, weight="bold").pack(
            anchor="w", pady=(0, 12))

        label(f, "Session title:", size=12, color=T2).pack(anchor="w")
        default_title = f"{self.protocol.get('name','?')} — {fmt_date(now_ts())}"
        title_var = ctk.StringVar(value=default_title)
        entry(f, textvariable=title_var, width=400, height=36).pack(
            anchor="w", pady=(3, 10))

        label(f, "Observations / deviations:", size=12, color=T2).pack(anchor="w")
        obs_box = textbox(f, width=400, height=80)
        obs_box.pack(anchor="w", pady=(3, 0))

        brow = ctk.CTkFrame(f, fg_color="transparent")
        brow.pack(anchor="w", pady=(14, 0))

        def _save():
            t = title_var.get().strip() or default_title
            obs = obs_box.get("0.0", "end").strip()
            d.destroy()
            self._finish_session(t, obs)

        def _skip():
            d.destroy()
            self._finish_session(default_title, "")

        btn(brow, "Save to Lab Notebook", _save,
            color=GREEN, text_color=("#fff","#fff"),
            width=190, height=38).pack(side="left", padx=(0, 8))
        btn(brow, "Skip", _skip,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=70, height=38).pack(side="left")
        _show_dialog(d, self, 460, 300)

    def _completion_status(self, sr):
        """Map internal step status to a lab-notebook-friendly string."""
        s = sr["status"]
        if s == self.ST_COMPLETED: return "completed"
        if s == self.ST_SKIPPED:   return "skipped"
        return "incomplete"

    def _finish_session(self, title, observations):
        now = now_ts()
        steps = self.protocol.get("steps", [])
        step_records = []
        for i, (step, sr) in enumerate(zip(steps, self._step_states)):
            used = self._used_secs(sr)
            step_records.append({
                "stepId":      step.get("id",""),
                "stepTitle":   step.get("title","?"),
                "stepType":    step.get("type","custom"),
                "status":      self._completion_status(sr),
                "plannedSecs": sr["original_secs"],
                "usedSecs":    used,
                "notes":       sr.get("notes",""),
            })
        session = {
            "id":               new_id(),
            "title":            title,
            "protocolId":       self.protocol["id"],
            "protocolName":     self.protocol.get("name","?"),
            "protocolSnapshot": self.protocol,
            "startedAt":        self._session_start_ts,
            "endedAt":          now,
            "actualDuration":   round((now - self._session_start_ts) / 60000) if self._session_start_ts else None,
            "timeline":         self._timeline,
            "stepRecords":      step_records,
            "observations":     observations,
            "tags":             [],
            "notes":            "",
        }
        runs = self.app.runs
        runs.insert(0, session)
        self.app.runs = runs
        save_runs(runs)
        _discard_runtime_session()   # session complete — clear crash-recovery file
        has_temp = any(s.get("tempBlock") for s in self.protocol.get("steps",[]))
        if has_temp:
            self._ask_save_temp_blocks()
        else:
            self._render_summary()

    def _render_summary(self):
        for w in self._main.winfo_children(): w.destroy()
        f = self._main
        f.grid_columnconfigure(0, weight=1)
        r = 0
        label(f, "Session Summary", size=22, weight="bold").grid(
            row=r, column=0, pady=(24, 4), sticky="w", padx=20); r += 1
        label(f, self.protocol.get("name",""), size=14, color=T2).grid(
            row=r, column=0, pady=(0, 16), sticky="w", padx=20); r += 1
        steps     = self.protocol.get("steps", [])
        done_n    = sum(1 for sr in self._step_states if sr["status"] == self.ST_COMPLETED)
        skipped_n = sum(1 for sr in self._step_states if sr["status"] == self.ST_SKIPPED)
        t_orig    = sum(sr["original_secs"] for sr in self._step_states)
        t_used    = sum(self._used_secs(sr) for sr in self._step_states)
        summ = card_frame(f)
        summ.grid(row=r, column=0, sticky="ew", padx=20, pady=(0, 16)); r += 1
        summ.grid_columnconfigure((0,1,2,3), weight=1)
        done_str = f"{done_n}/{len(steps)}"
        if skipped_n:
            done_str += f"  ({skipped_n} skipped)"
        for col_i, (title, val) in enumerate([
            ("Steps Completed", done_str),
            ("Planned Total",   fmt_mins(round(t_orig/60, 1))),
            ("Time Used",       fmt_mins(round(t_used/60, 1))),
        ]):
            cf = ctk.CTkFrame(summ, fg_color="transparent")
            cf.grid(row=0, column=col_i, padx=16, pady=16, sticky="ew")
            label(cf, title, size=11, color=T2).pack(anchor="center")
            label(cf, val, size=20, weight="bold").pack(anchor="center")
        label(f, "Step Breakdown", size=14, weight="bold", color=T2).grid(
            row=r, column=0, sticky="w", padx=20, pady=(4, 4)); r += 1
        for i, (step, sr) in enumerate(zip(steps, self._step_states)):
            sc = card_frame(f)
            sc.grid(row=r, column=0, sticky="ew", padx=20, pady=(0, 6)); r += 1
            srow = ctk.CTkFrame(sc, fg_color="transparent")
            srow.pack(fill="x", padx=14, pady=10)
            if sr["status"] == self.ST_COMPLETED:
                ic = GREEN
            elif sr["status"] == self.ST_SKIPPED:
                ic = self.ST_COLOR[self.ST_SKIPPED][0]
            else:
                ic = T3
            label(srow, self.ST_ICON.get(sr["status"], "○"), size=14,
                  color=ic).pack(side="left", padx=(0, 8))
            label(srow, f"{i+1}. {step.get('title','?')}", size=13,
                  weight="bold").pack(side="left")
            used = self._used_secs(sr)
            label(srow, f"{fmt_secs(used)} / {fmt_secs(sr['original_secs'])}",
                  size=12, color=T2).pack(side="right")
        brow = ctk.CTkFrame(f, fg_color="transparent")
        brow.grid(row=r, column=0, pady=24); r += 1
        btn(brow, "Back to Steps", self._render_steps,
            color=ACC, text_color=("#fff","#fff"),
            width=140, height=40).pack(side="left", padx=8)
        btn(brow, "Choose Another Protocol", self._reset_selection,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=200, height=40).pack(side="left", padx=8)

    def _reset_selection(self):
        self._cancel_all_timers()
        self.protocol     = None
        self._step_states = []
        self._seq_mode    = False
        self._temp_block_indices = []
        self._show_topbar_controls()
        self._refresh_proto_panel()
        self._show_empty()

    def _rebind_scroll(self, widget=None):
        """Trigger immediate scroll re-binding on the main content ScrollFrame."""
        try:
            self._main._refresh_scroll_bindings()
        except Exception:
            pass

    def _find_current_step_idx(self):
        """Find first running step, or last completed, or 0."""
        for i, sr in enumerate(self._step_states):
            if sr["status"] == self.ST_RUNNING: return i
        for i in range(len(self._step_states)-1, -1, -1):
            if self._step_states[i]["status"] == self.ST_COMPLETED: return i+1
        return len(self._step_states)

    def _add_temp_block(self):
        if not self.protocol:
            messagebox.showwarning("No Protocol", "Select a protocol first."); return
        def _on_confirmed(data):
            block = {
                "id": new_id(),
                "order": 0,
                "title": data["title"],
                "type": data["step_type"],
                "description": data["description"],
                "reagents": [], "equipment": [],
                "handsOnMinutes": data["time_mins"],
                "waitMinutes": 0, "bufferMinutes": 0,
                "temperature": "", "centrifugeCondition": "", "shakingRotation": "",
                "checklist": [], "notes": "", "warnings": "", "substeps": [],
                "tempBlock": True, "addedAt": now_ts(),
            }
            _temp_secs = int(data["time_mins"] * 60)
            new_state = {
                "status":               self.ST_IDLE,
                "timer_secs":           float(_temp_secs),
                "original_secs":        _temp_secs,
                "adjusted_total_secs":  _temp_secs,
                "elapsed_secs":         0,
                "timer_secs_at_start":  _temp_secs,
                "adjusted":             False,
                "notes":                "",
                "timer_job":            None,
                "timer_lbl":            None,
                "status_dot":           None,
                "prog_var":             None,
                "btn_start":            None,
                "card":                 None,
                "is_temp":              True,
                # Wall-clock anchors
                "start_mono":           None,
                "start_remaining":      float(_temp_secs),
                # Hands-on stopwatch
                "ho_elapsed_secs":      0,
                "ho_timer_job":         None,
                "ho_timer_lbl":         None,
                "ho_btn":               None,
                "ho_status":            "idle",
                "ho_start_mono":        None,
                "ho_start_elapsed":     0,
                # Undo-complete state
                "_pre_complete_status": self.ST_IDLE,
                "_undo_job":            None,
                "_ctrl_norm":           None,
                "_ctrl_done":           None,
                "_undo_lbl":            None,
                "_ctrl_container":      None,
            }
            steps = list(self.protocol.get("steps", []))
            pos   = data.get("position", "end")
            if pos == "after_current":
                ins = self._find_current_step_idx()
            elif pos == "before_current":
                ins = max(0, self._find_current_step_idx() - 1)
            else:
                ins = len(steps)
            steps.insert(ins, block)
            self._step_states.insert(ins, new_state)
            self.protocol["steps"] = steps
            if ins not in self._temp_block_indices:
                self._temp_block_indices.append(ins)
            self._log(f"Added temp block: {block['title']}")
            # Defer so dialog window visually closes before the heavy rebuild
            self.after(0, self._render_steps)
            self._autosave()
        AddBlockDialog(self, _on_confirmed)

    def _ask_save_temp_blocks(self):
        d = ctk.CTkToplevel(self)
        d.withdraw()
        d.title("New blocks added during run")
        d.resizable(False, False)
        f = ctk.CTkFrame(d, fg_color=BG, corner_radius=R_XL)
        f.pack(fill="both", expand=True, padx=20, pady=20)
        label(f, "You added blocks during this run.", size=14, weight="bold").pack(anchor="w")
        label(f, "What would you like to do with them?", size=13, color=T2).pack(anchor="w", pady=(4,12))
        temp_steps = [s for s in self.protocol.get("steps",[]) if s.get("tempBlock")]
        label(f, f"{len(temp_steps)} temporary block(s)", size=12, color=T3).pack(anchor="w", pady=(0,10))
        def _keep():
            d.destroy(); self._render_summary()
        def _add_to_proto():
            d.destroy()
            protocols = self.app.protocols
            for i, p in enumerate(protocols):
                if p["id"] == self.protocol["id"]:
                    protocols[i]["steps"] = [
                        {k:v for k,v in s.items() if k not in ("tempBlock","addedAt")}
                        for s in self.protocol["steps"]
                    ]
                    protocols[i]["updatedAt"] = now_ts()
                    break
            save_protocols(protocols)
            self.app.protocols = protocols
            self._render_summary()
        for txt, cmd in [
            ("Keep only in notebook record (don't update protocol)", _keep),
            ("Add temporary blocks to current protocol permanently",  _add_to_proto),
        ]:
            btn(f, txt, cmd, color=ACC, text_color=("#fff","#fff"),
                width=400, height=36).pack(anchor="w", pady=3)
        _show_dialog(d, self, 460, 240)

    def start(self, protocol):
        """External entry point: select a protocol and show its step timers."""
        self._select_proto_obj(protocol)

    def refresh(self):
        self._refresh_proto_panel()
        if not self.protocol:
            self._show_topbar_controls()
            self._show_empty()

# ─── Export helpers ───────────────────────────────────────────────────────────
def _export_session_pdf(session, filepath):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib import colors as rl_colors
    except ImportError:
        return False

    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    body  = styles["BodyText"]
    head1 = styles["Heading1"]
    story = []

    title = session.get("title") or session.get("protocolName","Untitled Session")
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=rl_colors.HexColor("#e2e8f0")))
    story.append(Spacer(1, 10))

    date_str = fmt_date(session.get("startedAt", 0))
    story.append(Paragraph(f"<b>Protocol:</b> {session.get('protocolName','?')}", body))
    story.append(Paragraph(f"<b>Date:</b> {date_str}", body))
    story.append(Spacer(1, 14))

    # Step records (new format) or old stepRuns format
    step_records = session.get("stepRecords", [])
    if step_records:
        story.append(Paragraph("Step Summary", head1))
        for sr in step_records:
            used    = sr.get("usedSecs", 0)
            planned = sr.get("plannedSecs", 0)
            icon    = "✓" if sr.get("status") == "completed" else "○"
            line    = f"{icon}  <b>{sr.get('stepTitle','?')}</b>"
            if planned: line += f"  — planned {round(planned/60,1)} min"
            if used:    line += f", used {round(used/60,1)} min"
            story.append(Paragraph(line, body))
            if sr.get("notes"):
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;Notes: {sr['notes']}", body))
        story.append(Spacer(1, 14))

    # Timeline
    timeline = session.get("timeline", [])
    if timeline:
        story.append(Paragraph("Timeline", head1))
        for ev in timeline:
            story.append(Paragraph(f"{ev.get('time','')}  —  {ev.get('text','')}", body))
        story.append(Spacer(1, 14))

    # Observations
    obs = session.get("observations","")
    if obs:
        story.append(Paragraph("Observations", head1))
        story.append(Paragraph(obs.replace("\n","<br/>"), body))
        story.append(Spacer(1, 14))

    # Notes field (blank for writing)
    story.append(Paragraph("Additional Notes", head1))
    story.append(Paragraph(session.get("notes",""), body))

    try:
        doc.build(story)
        return True
    except Exception:
        return False


def _export_session_docx(session, filepath):
    try:
        from docx import Document
    except ImportError:
        return False

    doc = Document()

    title = session.get("title") or session.get("protocolName","Untitled Session")
    doc.add_heading(title, level=0)

    date_str = fmt_date(session.get("startedAt", 0))
    p = doc.add_paragraph()
    p.add_run("Protocol: ").bold = True
    p.add_run(session.get("protocolName","?"))
    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run(date_str)
    doc.add_paragraph()

    step_records = session.get("stepRecords", [])
    if step_records:
        doc.add_heading("Step Summary", level=1)
        for sr in step_records:
            used    = sr.get("usedSecs", 0)
            planned = sr.get("plannedSecs", 0)
            icon    = "✓" if sr.get("status") == "completed" else "○"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{icon} {sr.get('stepTitle','?')}").bold = True
            if planned: p.add_run(f"  — {round(planned/60,1)} min planned")
            if used:    p.add_run(f", {round(used/60,1)} min used")
            if sr.get("notes"):
                doc.add_paragraph(f"    Notes: {sr['notes']}")

    timeline = session.get("timeline", [])
    if timeline:
        doc.add_heading("Timeline", level=1)
        for ev in timeline:
            doc.add_paragraph(f"{ev.get('time','')}  —  {ev.get('text','')}")

    obs = session.get("observations","")
    if obs:
        doc.add_heading("Observations", level=1)
        doc.add_paragraph(obs)

    doc.add_heading("Additional Notes", level=1)
    notes = session.get("notes","")
    doc.add_paragraph(notes if notes else " ")

    try:
        doc.save(filepath)
        return True
    except Exception:
        return False


# ─── Lab Notebook ──────────────────────────────────────────────────────────────
class HistoryPage(PageBase):
    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        self._filter_text = ""
        self._build()

    def _build(self):
        # Header + filter bar
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(28, 8))
        hdr.grid_columnconfigure(1, weight=1)

        label(hdr, "Lab Notebook", size=24, weight="bold").grid(
            row=0, column=0, sticky="w")

        self._filter_var = ctk.StringVar()
        filter_box = ctk.CTkEntry(hdr, textvariable=self._filter_var,
                                  placeholder_text="Filter by protocol name…",
                                  font=(FONT, 13), height=34, corner_radius=R_SM)
        filter_box.grid(row=0, column=1, sticky="ew", padx=(20, 0))
        self._filter_var.trace_add("write", lambda *_: self.refresh())

        self.scroll = ScrollFrame(self)
        self.scroll.grid(row=1, column=0, sticky="nsew", padx=28, pady=(0, 28))
        self.scroll.grid_columnconfigure(0, weight=1)

    def refresh(self):
        for w in self.scroll.winfo_children(): w.destroy()
        query = self._filter_var.get().strip().lower()
        runs  = [r for r in self.app.runs
                 if query in r.get("protocolName","").lower()
                 or query in r.get("title","").lower()] if query else self.app.runs
        if not runs:
            msg = "No sessions match the filter." if query else "No experiment sessions yet.\nFinish a Run Mode session to save one."
            label(self.scroll, msg, size=13, color=T3).grid(
                row=0, column=0, pady=60)
            return
        for i, r in enumerate(runs):
            self._session_card(r, i)

    def _session_card(self, session, row_idx):
        card = card_frame(self.scroll)
        card.grid(row=row_idx, column=0, sticky="ew", pady=(0, 10))
        card.grid_columnconfigure(0, weight=1)

        # ── Header row ────────────────────────────────────────────────────────
        top = ctk.CTkFrame(card, fg_color="transparent")
        top.pack(fill="x", padx=16, pady=(14, 6))
        top.grid_columnconfigure(0, weight=1)

        title = session.get("title") or session.get("protocolName","?")
        label(top, title, size=14, weight="bold").grid(row=0, column=0, sticky="w")

        date_str = fmt_date(session.get("startedAt", 0))
        label(top, date_str, size=11, color=T3).grid(row=1, column=0, sticky="w")

        # Export + Delete buttons
        btn_row = ctk.CTkFrame(top, fg_color="transparent")
        btn_row.grid(row=0, column=1, rowspan=2, sticky="e", padx=(8,0))

        def do_export_pdf(s=session):
            from tkinter import filedialog
            default = (s.get("title") or s.get("protocolName","session")).replace(" ","_")
            path = filedialog.asksaveasfilename(
                defaultextension=".pdf", filetypes=[("PDF","*.pdf")],
                initialfile=f"{default}.pdf")
            if path:
                ok = _export_session_pdf(s, path)
                messagebox.showinfo("Export", "PDF saved." if ok else "Export failed.")

        def do_export_docx(s=session):
            from tkinter import filedialog
            default = (s.get("title") or s.get("protocolName","session")).replace(" ","_")
            path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word","*.docx")],
                initialfile=f"{default}.docx")
            if path:
                ok = _export_session_docx(s, path)
                messagebox.showinfo("Export", "Word file saved." if ok else "Export failed.")

        btn(btn_row, "PDF", do_export_pdf,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=52, height=28, size=11).pack(side="left", padx=2)
        btn(btn_row, "Word", do_export_docx,
            color=("#e2e8f0","#334155"), text_color=T1,
            width=54, height=28, size=11).pack(side="left", padx=2)

        def del_session(sid=session["id"]): self._delete(sid)
        ctk.CTkButton(btn_row, text="✕", width=28, height=28,
                      fg_color=DANGER, font=(FONT, 12),
                      command=del_session, corner_radius=R_XS).pack(side="left", padx=(4,0))

        separator(card).pack(fill="x", padx=12)

        # ── Step records ──────────────────────────────────────────────────────
        step_records = session.get("stepRecords", [])
        old_step_runs = session.get("stepRuns", [])
        steps_src = session.get("protocolSnapshot",{}).get("steps",[])

        if step_records:
            sr_f = ctk.CTkFrame(card, fg_color="transparent")
            sr_f.pack(fill="x", padx=16, pady=(8, 4))
            for sr in step_records:
                row_f = ctk.CTkFrame(sr_f, fg_color="transparent")
                row_f.pack(fill="x", pady=1)
                dot_col = GREEN if sr.get("status") == "completed" else ("#94a3b8","#475569")
                dot = ctk.CTkFrame(row_f, fg_color=dot_col,
                                   width=14, height=14, corner_radius=R_XS)
                dot.pack(side="left", padx=(0, 8))
                label(row_f, sr.get("stepTitle","?"), size=12).pack(side="left")
                planned_m = round(sr.get("plannedSecs",0)/60, 1)
                used_m    = round(sr.get("usedSecs",0)/60, 1)
                if planned_m:
                    info = f"{used_m}m used / {planned_m}m planned"
                    label(row_f, info, size=11, color=T3).pack(side="right")
                if sr.get("notes"):
                    label(row_f, f"  ↳ {sr['notes']}", size=11,
                          color=T2).pack(side="left", padx=8)
        elif old_step_runs and steps_src:
            # backward-compat with old format
            sr_f = ctk.CTkFrame(card, fg_color="transparent")
            sr_f.pack(fill="x", padx=16, pady=(8, 4))
            for step, sr in zip(steps_src, old_step_runs):
                row_f = ctk.CTkFrame(sr_f, fg_color="transparent")
                row_f.pack(fill="x", pady=1)
                dot_col = GREEN if sr.get("completed") else ("#94a3b8","#475569")
                dot = ctk.CTkFrame(row_f, fg_color=dot_col,
                                   width=14, height=14, corner_radius=R_XS)
                dot.pack(side="left", padx=(0, 8))
                label(row_f, step.get("title","?"), size=12).pack(side="left")
                if sr.get("plannedMinutes"):
                    label(row_f, f"{sr.get('plannedMinutes',0)} min planned",
                          size=11, color=T3).pack(side="right")

        # ── Timeline (collapsible) ─────────────────────────────────────────────
        timeline = session.get("timeline", [])
        if timeline:
            tl_visible = ctk.BooleanVar(value=False)
            tl_toggle  = ctk.CTkButton(
                card, text="▸  Timeline", font=(FONT, 11),
                fg_color="transparent", text_color=T3,
                hover_color=("#f1f5f9","#1e293b"), anchor="w",
                height=28, corner_radius=R_XS)
            tl_toggle.pack(anchor="w", padx=16, pady=(4, 0))

            tl_frame = ctk.CTkFrame(card, fg_color=("#f8fafc","#1e293b"),
                                     corner_radius=R_SM)

            def toggle_tl(tl_frame=tl_frame, btn=tl_toggle, v=tl_visible):
                v.set(not v.get())
                if v.get():
                    tl_frame.pack(fill="x", padx=16, pady=(4, 8))
                    btn.configure(text="▾  Timeline")
                else:
                    tl_frame.pack_forget()
                    btn.configure(text="▸  Timeline")

            tl_toggle.configure(command=toggle_tl)

            for ev in timeline:
                tl_row = ctk.CTkFrame(tl_frame, fg_color="transparent")
                tl_row.pack(fill="x", padx=12, pady=1)
                label(tl_row, ev.get("time",""), size=10, color=T3,
                      weight="bold").pack(side="left", padx=(0, 10))
                label(tl_row, ev.get("text",""), size=11, color=T1).pack(side="left")

        # ── Observations ──────────────────────────────────────────────────────
        obs = session.get("observations","")
        if obs:
            obs_f = ctk.CTkFrame(card, fg_color=("#f0fdf4","#052e16"),
                                  corner_radius=R_SM, border_width=1,
                                  border_color=("#bbf7d0","#166534"))
            obs_f.pack(fill="x", padx=16, pady=(4, 10))
            label(obs_f, "Observations", size=10, color=("#15803d","#4ade80"),
                  weight="bold").pack(anchor="w", padx=12, pady=(6, 2))
            label(obs_f, obs, size=12, color=T1, wraplength=640).pack(
                anchor="w", padx=12, pady=(0, 8))
        else:
            ctk.CTkFrame(card, height=10, fg_color="transparent").pack()

    def _delete(self, sid):
        if not messagebox.askyesno("Delete", "Delete this session record?"): return
        self.app.runs = [r for r in self.app.runs if r["id"] != sid]
        save_runs(self.app.runs)
        self.refresh()

# ─── Settings ─────────────────────────────────────────────────────────────────
class SettingsPage(PageBase):
    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self._build()

    def _build(self):
        scroll = ScrollFrame(self)
        scroll.grid(row=0, column=0, sticky="nsew", padx=28, pady=28)
        scroll.grid_columnconfigure(0, weight=1)

        r = 0
        label(scroll, "Settings", size=24, weight="bold").grid(row=r, column=0, sticky="w", pady=(0,16)); r+=1

        # Appearance
        ap = card_frame(scroll)
        ap.grid(row=r, column=0, sticky="ew", pady=(0,12)); r+=1
        label(ap, "Appearance", size=14, weight="bold").pack(anchor="w", padx=16, pady=(12,8))
        row_ap = ctk.CTkFrame(ap, fg_color="transparent")
        row_ap.pack(fill="x", padx=16, pady=(0,12))
        label(row_ap, "Theme", size=13).pack(side="left")
        seg = ctk.CTkSegmentedButton(row_ap, values=["Light", "Dark", "System"],
                                      font=(FONT,13), command=self._set_theme)
        seg.pack(side="right")
        current = ctk.get_appearance_mode()
        seg.set(current.capitalize() if current.capitalize() in ["Light","Dark"] else "System")
        self._seg = seg

        # Data
        data_f = card_frame(scroll)
        data_f.grid(row=r, column=0, sticky="ew", pady=(0,12)); r+=1
        label(data_f, "Data Management", size=14, weight="bold").pack(anchor="w", padx=16, pady=(12,8))
        for lbl_txt, cmd, col in [
            ("Export All Data (JSON)", self._export, ACC),
            ("Import Data (JSON)", self._import, ("#64748b","#475569")),
        ]:
            row_d = ctk.CTkFrame(data_f, fg_color="transparent")
            row_d.pack(fill="x", padx=16, pady=(0,8))
            label(row_d, lbl_txt, size=13).pack(side="left")
            btn(row_d, "→", cmd, color=col, width=60, height=30).pack(side="right")

        # Danger
        danger_f = ctk.CTkFrame(scroll, fg_color=("#fff5f5","#2d0a0a"), corner_radius=R_MD,
                                  border_width=1, border_color=("#fca5a5","#7f1d1d"))
        danger_f.grid(row=r, column=0, sticky="ew", pady=(0,12)); r+=1
        label(danger_f, "Danger Zone", size=14, weight="bold", color=DANGER).pack(anchor="w", padx=16, pady=(12,8))
        row_dz = ctk.CTkFrame(danger_f, fg_color="transparent")
        row_dz.pack(fill="x", padx=16, pady=(0,12))
        label(row_dz, "Delete all protocols and run history", size=13).pack(side="left")
        btn(row_dz, "Clear All", self._clear, color=DANGER, width=100, height=30).pack(side="right")

        # About
        about = card_frame(scroll)
        about.grid(row=r, column=0, sticky="ew", pady=(0,12)); r+=1
        label(about, "About BenchFlow", size=14, weight="bold").pack(anchor="w", padx=16, pady=(12,4))
        label(about, "A local-first wet lab protocol manager.\nAll data stays on your device — no cloud, no accounts.", size=13, color=T2).pack(anchor="w", padx=16)
        label(about, f"Data stored at: {APP_DIR}", size=11, color=T3).pack(anchor="w", padx=16, pady=(4,12))

    def _set_theme(self, val):
        ctk.set_appearance_mode(val.lower())

    def _export(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON", "*.json")],
            initialfile=f"benchflow-export-{datetime.now().strftime('%Y-%m-%d')}.json"
        )
        if not path: return
        data = {
            "protocols": self.app.protocols,
            "runs": self.app.runs,
            "categories": self.app.categories,
            "tags": self.app.tags,
            "exportedAt": now_ts()
        }
        Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2))
        messagebox.showinfo("Export", f"Exported to:\n{path}")

    def _import(self):
        path = filedialog.askopenfilename(filetypes=[("JSON", "*.json")])
        if not path: return
        try:
            data = json.loads(Path(path).read_text())
            if "protocols" in data:
                existing_ids = {p["id"] for p in self.app.protocols}
                for p in data["protocols"]:
                    if p["id"] not in existing_ids:
                        self.app.protocols.append(p)
                save_protocols(self.app.protocols)
            if "runs" in data:
                existing_ids = {r["id"] for r in self.app.runs}
                for r in data["runs"]:
                    if r["id"] not in existing_ids:
                        self.app.runs.append(r)
                save_runs(self.app.runs)
            if "categories" in data:
                self.app.categories = sorted({*self.app.categories, *data["categories"]})
                save_categories(self.app.categories)
            if "tags" in data:
                self.app.tags = sorted({*self.app.tags, *data["tags"]})
                save_tags(self.app.tags)
            messagebox.showinfo("Import", "Import successful!")
        except Exception as e:
            messagebox.showerror("Import Error", str(e))

    def _clear(self):
        if not messagebox.askyesno("Clear All", "Delete ALL protocols and history? This cannot be undone."): return
        self.app.protocols = []
        self.app.runs = []
        self.app.categories = []
        self.app.tags = []
        save_protocols([])
        save_runs([])
        save_categories([])
        save_tags([])
        messagebox.showinfo("Cleared", "All data has been deleted.")

# ─── Import ───────────────────────────────────────────────────────────────────
def _extract_pdf_text(path: str) -> str:
    if not _HAS_PDF:
        return "[PyMuPDF not installed]"
    doc = pymupdf.open(path)
    pages = []
    for i, page in enumerate(doc):
        txt = page.get_text()
        if txt.strip():
            pages.append(f"── Page {i+1} ──\n{txt}")
    return "\n\n".join(pages)

def _extract_docx_text(path: str) -> str:
    if not _HAS_DOCX:
        return "[python-docx not installed]"
    doc = _docx_lib.Document(path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    return "\n".join(lines)


class QuickStepDialog(ctk.CTkToplevel):
    """Mini step-creator shown when user clicks 'Create Step from Selection'."""
    def __init__(self, parent, prefill_text: str, on_add):
        super().__init__(parent)
        self.on_add = on_add
        self.withdraw()
        self.title("Create Step")
        self.resizable(False, False)

        f = ctk.CTkFrame(self, fg_color="transparent")
        f.pack(fill="both", expand=True, padx=20, pady=16)
        f.grid_columnconfigure(0, weight=1)

        label(f, "Step Title", size=12, color=T2).grid(row=0, column=0, sticky="w")
        self.e_title = entry(f, "e.g. Wash cells with PBS", width=480)
        self.e_title.grid(row=1, column=0, sticky="ew", pady=(2,10))

        label(f, "Step Type", size=12, color=T2).grid(row=2, column=0, sticky="w")
        self.type_var = ctk.StringVar(value="preparation")
        ctk.CTkOptionMenu(f, variable=self.type_var, values=STEP_TYPES,
                          font=(FONT,13), width=200, corner_radius=R_SM).grid(row=3, column=0, sticky="w", pady=(2,10))

        label(f, "Description (from selection)", size=12, color=T2).grid(row=4, column=0, sticky="w")
        self.tb_desc = textbox(f, width=480, height=100)
        self.tb_desc.insert("0.0", prefill_text)
        self.tb_desc.grid(row=5, column=0, sticky="ew", pady=(2,10))

        time_row = ctk.CTkFrame(f, fg_color="transparent")
        time_row.grid(row=6, column=0, sticky="ew", pady=(0,10))
        for lbl_txt, attr, default in [("Hands-on (min)", "e_ho", "5"),
                                        ("Wait (min)",     "e_wt", "0"),
                                        ("Buffer (min)",   "e_bf", "0")]:
            tf = ctk.CTkFrame(time_row, fg_color="transparent")
            tf.pack(side="left", expand=True, fill="x", padx=(0,8))
            label(tf, lbl_txt, size=11, color=T2).pack(anchor="w")
            e = entry(tf, width=120)
            e.insert(0, default)
            e.pack(anchor="w")
            setattr(self, attr, e)

        cond_row = ctk.CTkFrame(f, fg_color="transparent")
        cond_row.grid(row=7, column=0, sticky="ew", pady=(0,14))
        for lbl_txt, attr, ph in [("Temperature", "e_temp", "e.g. 37°C"),
                                    ("Centrifuge",  "e_cent", "e.g. 300g×5min")]:
            cf = ctk.CTkFrame(cond_row, fg_color="transparent")
            cf.pack(side="left", expand=True, fill="x", padx=(0,8))
            label(cf, lbl_txt, size=11, color=T2).pack(anchor="w")
            e = entry(cf, ph, width=200)
            e.pack(anchor="w")
            setattr(self, attr, e)

        btn_row = ctk.CTkFrame(f, fg_color="transparent")
        btn_row.grid(row=8, column=0, sticky="ew")
        btn(btn_row, "Cancel", self.destroy, color=("#e2e8f0","#334155"), text_color=T1, width=90).pack(side="left")
        btn(btn_row, "Add Step →", self._add, width=120).pack(side="right")
        _show_dialog(self, parent, 540, 440)

    def _add(self):
        s = new_step(0)
        s["title"] = self.e_title.get().strip()
        s["type"] = self.type_var.get()
        s["description"] = self.tb_desc.get("0.0", "end").strip()
        try: s["handsOnMinutes"] = int(self.e_ho.get())
        except: s["handsOnMinutes"] = 0
        try: s["waitMinutes"] = int(self.e_wt.get())
        except: s["waitMinutes"] = 0
        try: s["bufferMinutes"] = int(self.e_bf.get())
        except: s["bufferMinutes"] = 0
        s["temperature"] = self.e_temp.get().strip()
        s["centrifugeCondition"] = self.e_cent.get().strip()
        self.on_add(s)
        self.destroy()


import re as _re

def _parse_steps_from_text(text):
    """Local rule-based step detection — no external API."""
    ACTION_VERBS = {
        "add","mix","incubate","centrifuge","spin","wash","aspirate","discard",
        "resuspend","transfer","vortex","pipette","thaw","chill","place","shake",
        "rotate","warm","spread","load","run","block","stain","image","lyse",
        "harvest","collect","heat","cool","store","prepare","dilute","aliquot",
        "sonicate","homogenize","filter","elute","equilibrate","quench",
    }
    TYPE_KEYWORDS = {
        "incubate":   "incubation", "incubation": "incubation", "overnight":"incubation",
        "centrifuge": "centrifuge", "spin":        "centrifuge",
        "wash":       "wash",
        "mix":        "mixing",     "vortex":      "mixing",     "shake":"mixing",
        "transfer":   "transfer",
        "pipette":    "pipetting",  "pipet":       "pipetting",  "aspirate":"pipetting",
        "resuspend":  "resuspension","resuspension":"resuspension",
        "stain":      "staining",   "staining":    "staining",
        "block":      "blocking",   "blocking":    "blocking",
        "image":      "imaging",    "imaging":     "imaging",
        "lyse":       "lysis",      "lysis":       "lysis",
        "harvest":    "harvest",
        "collect":    "sample_collection","collection":"sample_collection",
        "heat":       "heating",    "warm":        "heating",    "thaw":"heating",
        "cool":       "cooling",    "chill":       "cooling",    "ice":"cooling",
        "store":      "storage",    "storage":     "storage",
        "add":        "reagent_addition",
        "discard":    "transfer",
    }
    TIME_RE  = _re.compile(
        r'(\d+(?:\.\d+)?)\s*'
        r'(seconds?|secs?|minutes?|mins?|hours?|hrs?|h\b)',
        _re.IGNORECASE
    )
    OVERNIGHT_RE = _re.compile(r'overnight', _re.IGNORECASE)
    TEMP_RE  = _re.compile(
        r'(\d+)\s*°?\s*C\b|room\s*temp(?:erature)?|RT\b|on\s+ice\b|4\s*°?\s*C|37\s*°?\s*C',
        _re.IGNORECASE
    )
    CENT_RE  = _re.compile(
        r'(\d[\d,]*)\s*(?:rpm|rcf|×g|xg|x\s*g)',
        _re.IGNORECASE
    )
    STEP_NUM_RE = _re.compile(r'^(?:step\s+)?(\d+)[\.\)]\s+', _re.IGNORECASE)
    BULLET_RE   = _re.compile(r'^[\-\*•]\s+')

    def _infer_type(line_low):
        for kw, stype in TYPE_KEYWORDS.items():
            if kw in line_low:
                return stype
        return "preparation"

    def _parse_time(line):
        m = OVERNIGHT_RE.search(line)
        if m: return 720, "waitMinutes"
        matches = TIME_RE.findall(line)
        if not matches: return 0, "handsOnMinutes"
        val_str, unit = matches[0]
        val = float(val_str)
        u = unit.lower()
        if u.startswith("s"): mins = val / 60
        elif u.startswith("h"): mins = val * 60
        else: mins = val
        return round(mins, 1), "waitMinutes"

    def _parse_temp(line):
        m = TEMP_RE.search(line)
        if not m: return ""
        t = m.group(0).strip()
        for norm, canon in [("room temperature","RT"),("room temp","RT"),("on ice","ice")]:
            if norm.lower() in t.lower(): return canon
        return t

    def _parse_cent(line):
        m = CENT_RE.search(line)
        return m.group(0).strip() if m else ""

    lines = text.splitlines()
    candidates = []
    for raw in lines:
        stripped = raw.strip()
        if not stripped: continue
        if len(stripped) < 8: continue

        is_numbered = bool(STEP_NUM_RE.match(stripped))
        is_bullet   = bool(BULLET_RE.match(stripped))
        low = stripped.lower()
        has_verb = any(
            _re.search(r'\b' + v + r'\b', low)
            for v in ACTION_VERBS
        )

        if is_numbered or is_bullet or has_verb:
            clean = STEP_NUM_RE.sub("", stripped)
            clean = BULLET_RE.sub("", clean).strip()
            candidates.append(clean)

    steps = []
    for i, text_line in enumerate(candidates):
        low = text_line.lower()
        stype = _infer_type(low)
        mins, time_key = _parse_time(text_line)
        temp  = _parse_temp(text_line)
        cent  = _parse_cent(text_line)

        words = text_line.split()
        title_words = words[:min(6, len(words))]
        title = " ".join(title_words).rstrip(".,;:")
        if len(text_line) > len(title) + 2:
            title = title + "…"

        step = new_step(i)
        step["title"] = title
        step["type"]  = stype
        step["description"] = text_line
        step["temperature"]  = temp
        step["centrifugeCondition"] = cent
        if time_key == "waitMinutes":
            step["waitMinutes"]    = mins
            step["handsOnMinutes"] = 0
        else:
            step["handsOnMinutes"] = mins
            step["waitMinutes"]    = 0
        steps.append(step)

    return steps


class ImportPage(PageBase):
    """Two-phase import: (1) source input  →  (2) review & step builder."""

    def __init__(self, parent, app):
        super().__init__(parent, app)
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self._source_text = ""
        self._source_meta = {}   # sourceType, sourceName, sourceUrl, originalFileName
        self._pending_steps: list[dict] = []

        # Two frames stacked — we lift one at a time
        self._phase1 = ctk.CTkFrame(self, fg_color="transparent")
        self._phase1.grid(row=0, column=0, sticky="nsew")
        self._phase1.grid_columnconfigure(0, weight=1)
        self._phase1.grid_rowconfigure(1, weight=1)

        self._phase2 = ctk.CTkFrame(self, fg_color="transparent")
        self._phase2.grid(row=0, column=0, sticky="nsew")
        self._phase2.grid_columnconfigure(0, weight=1)
        self._phase2.grid_rowconfigure(1, weight=1)

        self._build_phase1()
        self._build_phase2()
        self._phase1.tkraise()

    def start(self, tab="text"):
        tab = tab if tab in ("text", "pdf", "docx") else "text"
        self._tab_var.set(tab)
        self._refresh_tab()
        try:
            self.e_proto_name.delete(0, "end")
            self.e_proto_cat.delete(0, "end")
        except Exception:
            pass
        self._source_text = ""
        self._source_meta = {}
        self._pending_steps = []
        self._phase1.tkraise()

    # ── Phase 1: source selection ──────────────────────────────────────────────
    def _build_phase1(self):
        f = self._phase1

        hdr = ctk.CTkFrame(f, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(24,12))
        btn(hdr, "← Library", lambda: self.app.navigate("library"),
            color=("#e2e8f0","#334155"), text_color=T1, width=95, height=32).pack(side="left", padx=(0, 14))
        label(hdr, "Import Protocol", size=24, weight="bold").pack(side="left")

        body = ScrollFrame(f)
        body.grid(row=1, column=0, sticky="nsew", padx=28, pady=(0,24))
        body.grid_columnconfigure(0, weight=1)

        # Tab bar
        tab_row = ctk.CTkFrame(body, fg_color="transparent")
        tab_row.grid(row=0, column=0, sticky="ew", pady=(0,16))
        self._tab_var = ctk.StringVar(value="text")
        for val, lbl_txt in [("text","📝  Paste Text / Web"), ("pdf","📄  PDF File"), ("docx","📝  Word (.docx)")]:
            rb = ctk.CTkRadioButton(tab_row, text=lbl_txt, variable=self._tab_var, value=val,
                                     font=(FONT,13), command=self._refresh_tab)
            rb.pack(side="left", padx=(0,20))

        # Tab content container
        self._tab_card = card_frame(body)
        self._tab_card.grid(row=1, column=0, sticky="ew", pady=(0,12))
        self._tab_card.grid_columnconfigure(0, weight=1)

        # Protocol metadata
        meta_card = card_frame(body)
        meta_card.grid(row=2, column=0, sticky="ew", pady=(0,12))
        meta_card.grid_columnconfigure(1, weight=1)
        label(meta_card, "Protocol Name", size=12, color=T2).grid(row=0, column=0, padx=(14,8), pady=(12,4), sticky="w")
        self.e_proto_name = entry(meta_card, "Protocol name...", width=320)
        self.e_proto_name.grid(row=0, column=1, sticky="w", pady=(12,4))
        label(meta_card, "Category", size=12, color=T2).grid(row=1, column=0, padx=(14,8), pady=(0,12), sticky="w")
        self.e_proto_cat = entry(meta_card, "e.g. Cell Biology", width=200)
        self.e_proto_cat.grid(row=1, column=1, sticky="w", pady=(0,12))

        btn_row = ctk.CTkFrame(body, fg_color="transparent")
        btn_row.grid(row=3, column=0, sticky="ew")
        btn(btn_row, "Cancel", lambda: self.app.navigate("library"),
            color=("#e2e8f0","#334155"), text_color=T1, width=100).pack(side="left")
        btn(btn_row, "Review & Build Steps →", self._go_review, width=200).pack(side="right")

        self._refresh_tab()

    def _refresh_tab(self):
        for w in self._tab_card.winfo_children(): w.destroy()
        tab = self._tab_var.get()
        if tab == "text":   self._build_text_tab(self._tab_card)
        elif tab == "pdf":  self._build_pdf_tab(self._tab_card)
        elif tab == "docx": self._build_docx_tab(self._tab_card)

    def _build_text_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        label(parent, "Source URL (optional)", size=12, color=T2).grid(row=0, column=0, sticky="w", padx=14, pady=(12,2))
        self.e_url = entry(parent, "https://...", width=500)
        self.e_url.grid(row=1, column=0, sticky="ew", padx=14, pady=(0,8))
        label(parent, "Paste protocol text here", size=12, color=T2).grid(row=2, column=0, sticky="w", padx=14)
        self.tb_paste = textbox(parent, width=700, height=260)
        self.tb_paste.grid(row=3, column=0, sticky="ew", padx=14, pady=(4,14))

    def _build_pdf_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        self._pdf_path = ""
        top = ctk.CTkFrame(parent, fg_color="transparent")
        top.grid(row=0, column=0, sticky="ew", padx=14, pady=14)
        self.lbl_pdf = label(top, "No file selected", size=13, color=T3)
        self.lbl_pdf.pack(side="left", expand=True, fill="x")
        btn(top, "Choose PDF…", self._pick_pdf, width=130, height=34).pack(side="right")
        if not _HAS_PDF:
            label(parent, "⚠ PyMuPDF not installed. Run: pip install pymupdf", size=12, color=DANGER).grid(row=1, column=0, padx=14, sticky="w")

    def _pick_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if not path: return
        self._pdf_path = path
        self.lbl_pdf.configure(text=Path(path).name, text_color=T1)

    def _build_docx_tab(self, parent):
        parent.grid_columnconfigure(0, weight=1)
        self._docx_path = ""
        top = ctk.CTkFrame(parent, fg_color="transparent")
        top.grid(row=0, column=0, sticky="ew", padx=14, pady=14)
        self.lbl_docx = label(top, "No file selected", size=13, color=T3)
        self.lbl_docx.pack(side="left", expand=True, fill="x")
        btn(top, "Choose .docx…", self._pick_docx, width=140, height=34).pack(side="right")
        if not _HAS_DOCX:
            label(parent, "⚠ python-docx not installed. Run: pip install python-docx", size=12, color=DANGER).grid(row=1, column=0, padx=14, sticky="w")

    def _pick_docx(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if not path: return
        self._docx_path = path
        self.lbl_docx.configure(text=Path(path).name, text_color=T1)

    def _go_review(self):
        tab = self._tab_var.get()
        raw = ""
        meta = {"sourceType": tab, "sourceName": "", "sourceUrl": "",
                 "importedAt": now_ts(), "rawText": "", "originalFileName": "", "notes": ""}

        if tab == "text":
            raw = self.tb_paste.get("0.0", "end").strip()
            meta["sourceUrl"] = self.e_url.get().strip()
            meta["sourceType"] = "web" if meta["sourceUrl"] else "text"
            meta["sourceName"] = meta["sourceUrl"] or "Pasted text"
        elif tab == "pdf":
            if not self._pdf_path:
                messagebox.showwarning("No file", "Please choose a PDF file first."); return
            raw = _extract_pdf_text(self._pdf_path)
            meta["originalFileName"] = Path(self._pdf_path).name
            meta["sourceName"] = Path(self._pdf_path).name
        elif tab == "docx":
            if not self._docx_path:
                messagebox.showwarning("No file", "Please choose a .docx file first."); return
            raw = _extract_docx_text(self._docx_path)
            meta["originalFileName"] = Path(self._docx_path).name
            meta["sourceName"] = Path(self._docx_path).name

        if not raw.strip():
            messagebox.showwarning("Empty", "No text found. Please check the source."); return

        meta["rawText"] = raw
        self._source_text = raw
        self._source_meta = meta
        # Auto-parse into draft steps
        self._pending_steps = _parse_steps_from_text(raw)

        proto_name = self.e_proto_name.get().strip() or meta["sourceName"] or "Imported Protocol"
        self._review_proto_name = proto_name
        self._review_proto_cat = self.e_proto_cat.get().strip()

        self._load_phase2(raw, proto_name)
        self._phase2.tkraise()

    # ── Phase 2: review & step builder ───────────────────────────────────────
    def _build_phase2(self):
        f = self._phase2

        hdr = ctk.CTkFrame(f, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=28, pady=(20,8))
        hdr.grid_columnconfigure(1, weight=1)

        btn(hdr, "← Back", self._back_to_phase1,
            color=("#e2e8f0","#334155"), text_color=T1, width=90, height=32).grid(row=0, column=0)

        self.lbl_review_title = label(hdr, "Review Protocol", size=18, weight="bold")
        self.lbl_review_title.grid(row=0, column=1, padx=16, sticky="w")

        right_btns = ctk.CTkFrame(hdr, fg_color="transparent")
        right_btns.grid(row=0, column=2)
        btn(right_btns, "Create Step from Selection", self._create_step_from_selection,
            width=220, height=32).pack(side="left", padx=(0,8))
        btn(right_btns, "💾  Save Protocol", self._save_import,
            color=GREEN, width=150, height=32).pack(side="left")

        # Split pane
        pane = ctk.CTkFrame(f, fg_color="transparent")
        pane.grid(row=1, column=0, sticky="nsew", padx=28, pady=(0,20))
        pane.grid_columnconfigure(0, weight=3)
        pane.grid_columnconfigure(1, weight=2)
        pane.grid_rowconfigure(0, weight=1)

        # Left: source text
        left = card_frame(pane)
        left.grid(row=0, column=0, sticky="nsew", padx=(0,8))
        left.grid_columnconfigure(0, weight=1)
        left.grid_rowconfigure(1, weight=1)
        label(left, "Source Text", size=12, color=T2, weight="bold").grid(row=0, column=0, sticky="w", padx=12, pady=(10,4))
        self.source_text_widget = tk.Text(
            left, font=(FONT, 12), wrap="word",
            bg="#f8fafc", fg="#0f172a",
            selectbackground="#bfdbfe", selectforeground="#1e40af",
            relief="flat", padx=8, pady=8,
            cursor="xterm"
        )
        self.source_text_widget.grid(row=1, column=0, sticky="nsew", padx=4, pady=(0,4))
        vsb = ctk.CTkScrollbar(left, command=self.source_text_widget.yview,
                               corner_radius=R_SM,
                               width=12, fg_color="transparent",
                               button_color=("#cbd5e1","#475569"),
                               button_hover_color=("#94a3b8","#64748b"))
        vsb.grid(row=1, column=1, sticky="ns")
        self.source_text_widget.configure(yscrollcommand=vsb.set)

        # Right: steps being built
        right = card_frame(pane)
        right.grid(row=0, column=1, sticky="nsew")
        right.grid_columnconfigure(0, weight=1)
        right.grid_rowconfigure(1, weight=1)
        right_hdr = ctk.CTkFrame(right, fg_color="transparent")
        right_hdr.grid(row=0, column=0, sticky="ew", padx=12, pady=(10,4))
        label(right_hdr, "Steps", size=12, color=T2, weight="bold").pack(side="left")
        self.lbl_step_count = label(right_hdr, "0 steps", size=11, color=T3)
        self.lbl_step_count.pack(side="right")

        self.steps_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        self.steps_scroll.grid(row=1, column=0, sticky="nsew", padx=4, pady=(0,4))
        self.steps_scroll.grid_columnconfigure(0, weight=1)

    def _load_phase2(self, text, proto_name):
        self.lbl_review_title.configure(text=f"Review: {proto_name}")
        self.source_text_widget.configure(state="normal")
        self.source_text_widget.delete("1.0", "end")
        self.source_text_widget.insert("1.0", text)
        # keep editable so user can select text
        self._render_pending_steps()
        if self._pending_steps:
            self.lbl_step_count.configure(
                text=f"{len(self._pending_steps)} draft steps (auto-detected)")

    def _back_to_phase1(self):
        self._phase1.tkraise()

    def _create_step_from_selection(self):
        try:
            sel = self.source_text_widget.get("sel.first", "sel.last").strip()
        except tk.TclError:
            sel = ""
        if not sel:
            messagebox.showinfo("Select Text", "Please select some text in the left panel first.")
            return
        def on_add(step):
            step["order"] = len(self._pending_steps)
            self._pending_steps.append(step)
            self._render_pending_steps()
        QuickStepDialog(self.app, sel, on_add)

    def _render_pending_steps(self):
        for w in self.steps_scroll.winfo_children(): w.destroy()
        self.lbl_step_count.configure(text=f"{len(self._pending_steps)} steps")
        for i, step in enumerate(self._pending_steps):
            self._pending_step_row(step, i)
        if not self._pending_steps:
            label(self.steps_scroll, "Select text on the left,\nthen click\n'Create Step from Selection'",
                  size=12, color=T3).grid(row=0, column=0, pady=40)

    def _pending_step_row(self, step, idx):
        stype = step.get("type","custom")
        bg = STEP_COLORS.get(stype, STEP_COLORS["custom"])
        bd = STEP_BADGES.get(stype, STEP_BADGES["custom"])
        card = ctk.CTkFrame(self.steps_scroll, fg_color=bg, corner_radius=R_SM,
                            border_width=1, border_color=bd)
        card.grid(row=idx, column=0, sticky="ew", pady=(0,5))
        card.grid_columnconfigure(1, weight=1)

        label(card, f"{idx+1}", size=11, color=T2).grid(row=0, column=0, padx=(8,4), pady=8)

        info = ctk.CTkFrame(card, fg_color="transparent")
        info.grid(row=0, column=1, sticky="ew")
        label(info, step.get("title") or "Untitled", size=12, weight="bold").pack(anchor="w")
        total_m = step.get("handsOnMinutes",0)+step.get("waitMinutes",0)+step.get("bufferMinutes",0)
        badge = STEP_BADGES.get(stype, STEP_BADGES["custom"])
        meta_txt = STEP_LABELS.get(stype,"")
        if total_m: meta_txt += f"  ·  {fmt_mins(total_m)}"
        label(info, meta_txt, size=10, color=T2).pack(anchor="w")

        acts = ctk.CTkFrame(card, fg_color="transparent")
        acts.grid(row=0, column=2, padx=(4,8))
        def move_up(i=idx): self._move_ps(i,-1)
        def move_dn(i=idx): self._move_ps(i,1)
        def edit_ps(s=step): self._edit_ps(s)
        def del_ps(i=idx): self._del_ps(i)
        for txt, cmd, col in [("↑",move_up,("#94a3b8","#475569")),("↓",move_dn,("#94a3b8","#475569")),
                               ("Edit",edit_ps,ACC),("✕",del_ps,DANGER)]:
            ctk.CTkButton(acts, text=txt, width=36, height=26, fg_color=col, font=(FONT,11),
                          command=cmd, corner_radius=R_XS).pack(side="left", padx=1)

    def _move_ps(self, idx, delta):
        steps = self._pending_steps
        ni = idx + delta
        if 0 <= ni < len(steps):
            steps[idx], steps[ni] = steps[ni], steps[idx]
            for i, s in enumerate(steps): s["order"] = i
            self._render_pending_steps()

    def _edit_ps(self, step):
        def on_save(updated):
            for i, s in enumerate(self._pending_steps):
                if s["id"] == updated["id"]:
                    self._pending_steps[i] = updated; break
            self._render_pending_steps()
        StepEditorDialog(self.app, step, on_save)

    def _del_ps(self, idx):
        self._pending_steps.pop(idx)
        for i, s in enumerate(self._pending_steps): s["order"] = i
        self._render_pending_steps()

    def _save_import(self):
        if not self._pending_steps:
            messagebox.showwarning("No Steps", "Add at least one step before saving."); return

        source = dict(self._source_meta)
        p = {
            "id": new_id(),
            "name": getattr(self, "_review_proto_name", "Imported Protocol"),
            "category": getattr(self, "_review_proto_cat", ""),
            "description": "",
            "createdAt": now_ts(), "updatedAt": now_ts(),
            "tags": [],
            "steps": [dict(s) for s in self._pending_steps],
            "source": source,
        }
        self.app.protocols.append(p)
        if p["category"] and p["category"] not in self.app.categories:
            self.app.categories.append(p["category"])
            save_categories(self.app.categories)
        save_protocols(self.app.protocols)
        messagebox.showinfo("Saved", f"Protocol '{p['name']}' saved with {len(p['steps'])} steps.")
        self._pending_steps = []
        self._render_pending_steps()
        self._back_to_phase1()
        self.app.navigate("library")

    def refresh(self):
        pass


# ══════════════════════════════════════════════════════════════════════════════
# Schedule — Experiment Calendar
# ══════════════════════════════════════════════════════════════════════════════
class MiniCalendarPicker(ctk.CTkToplevel):
    """Compact month-grid date picker popup.

    Opens centered over *parent*, blocks interaction until the user picks a
    date or presses Escape, then calls ``on_pick(date)`` and destroys itself.
    """

    def __init__(self, parent, initial_date, on_pick):
        super().__init__(parent)
        self._on_pick  = on_pick
        self._year     = initial_date.year
        self._month    = initial_date.month
        self._selected = initial_date
        self._today    = datetime.now().date()
        self._matrix   = []          # filled by _refresh_grid
        self._btns     = []          # flat list, len = 42  (6 rows × 7 cols)

        self.withdraw()
        self.title("Pick a Date")
        self.resizable(False, False)
        self._build()

        # Centre over parent window
        self.update_idletasks()
        px = parent.winfo_rootx()
        py = parent.winfo_rooty()
        pW = parent.winfo_width()
        pH = parent.winfo_height()
        w  = self.winfo_reqwidth()
        h  = self.winfo_reqheight()
        x  = px + (pW - w) // 2
        y  = py + (pH - h) // 2
        self.geometry(f"+{x}+{y}")

        self.deiconify()
        self.lift()
        self.focus_force()
        self.grab_set()
        self.bind("<Escape>", lambda e: self._close())

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build(self):
        outer = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG,
                              border_width=1, border_color=CARD_B)
        outer.pack(fill="both", expand=True, padx=2, pady=2)

        # ── Month navigation ─────────────────────────────────────────────────
        nav = ctk.CTkFrame(outer, fg_color="transparent")
        nav.pack(fill="x", padx=10, pady=(12, 4))
        ctk.CTkButton(
            nav, text="◀", width=30, height=28, font=(FONT, 11),
            fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
            text_color=T1, command=self._prev_month,
        ).pack(side="left")
        self._lbl_month = ctk.CTkLabel(
            nav, text="", font=(FONT, 13, "bold"), text_color=T1)
        self._lbl_month.pack(side="left", expand=True)
        ctk.CTkButton(
            nav, text="▶", width=30, height=28, font=(FONT, 11),
            fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
            text_color=T1, command=self._next_month,
        ).pack(side="right")

        # ── Weekday header row ────────────────────────────────────────────────
        dh = ctk.CTkFrame(outer, fg_color="transparent")
        dh.pack(fill="x", padx=10)
        for abbr in ("Mo", "Tu", "We", "Th", "Fr", "Sa", "Su"):
            ctk.CTkLabel(dh, text=abbr, width=36, height=22,
                          font=(FONT, 10), text_color=T3).pack(side="left")

        # ── 6 × 7 day buttons ─────────────────────────────────────────────────
        gf = ctk.CTkFrame(outer, fg_color="transparent")
        gf.pack(fill="x", padx=10, pady=(2, 12))
        for row in range(6):
            for col in range(7):
                b = ctk.CTkButton(
                    gf, text="", width=36, height=30,
                    font=(FONT, 12), corner_radius=6,
                    fg_color="transparent",
                    hover_color=("#e2e8f0", "#334155"),
                    text_color=T1,
                    command=lambda r=row, c=col: self._pick(r, c),
                )
                b.grid(row=row, column=col, padx=1, pady=1)
                self._btns.append(b)

        self._refresh_grid()

    # ── Grid refresh ──────────────────────────────────────────────────────────
    def _refresh_grid(self):
        self._lbl_month.configure(
            text=datetime(self._year, self._month, 1).strftime("%B %Y"))
        matrix = _calendar.monthcalendar(self._year, self._month)
        while len(matrix) < 6:
            matrix.append([0] * 7)
        self._matrix = matrix

        for row in range(6):
            for col in range(7):
                b = self._btns[row * 7 + col]
                day_num = matrix[row][col]
                if day_num == 0:
                    b.configure(text="", fg_color="transparent",
                                state="disabled", text_color=T3)
                else:
                    d = datetime(self._year, self._month, day_num).date()
                    is_sel   = (d == self._selected)
                    is_today = (d == self._today)
                    if is_sel:
                        fg, tc = ("#3b82f6", "#3b82f6"), ("#ffffff", "#ffffff")
                    elif is_today:
                        fg, tc = ("#dbeafe", "#1e3a5f"), ("#1d4ed8", "#93c5fd")
                    else:
                        fg, tc = "transparent", T1
                    b.configure(text=str(day_num), state="normal",
                                fg_color=fg, text_color=tc)

    # ── Navigation ────────────────────────────────────────────────────────────
    def _prev_month(self):
        if self._month == 1:
            self._month, self._year = 12, self._year - 1
        else:
            self._month -= 1
        self._refresh_grid()

    def _next_month(self):
        if self._month == 12:
            self._month, self._year = 1, self._year + 1
        else:
            self._month += 1
        self._refresh_grid()

    # ── Pick / close ──────────────────────────────────────────────────────────
    def _pick(self, row, col):
        if row >= len(self._matrix):
            return
        day_num = self._matrix[row][col]
        if day_num == 0:
            return
        picked = datetime(self._year, self._month, day_num).date()
        cb = self._on_pick
        self._close()
        cb(picked)

    def _close(self):
        try:
            self.grab_release()
            self.destroy()
        except Exception:
            pass


class ScheduleExperimentDialog(ctk.CTkToplevel):
    """Create or edit a scheduled experiment."""

    def __init__(self, parent, app, existing_exp, on_save):
        super().__init__(parent)
        self.withdraw()
        self.app     = app
        self.exp     = existing_exp
        self.on_save = on_save
        self.is_edit = existing_exp is not None
        self._scheduled_steps = []
        self._exp_end_dt      = None
        self.title("Edit Schedule" if self.is_edit else "Schedule Experiment")
        self.resizable(True, True)
        self._build()
        _show_dialog(self, parent, 530, 600)

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # Header
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 4))
        label(hdr, "Edit Schedule" if self.is_edit else "Schedule Experiment",
              size=14, weight="bold").pack(anchor="w")
        label(hdr, "Automatically generates a full experiment timeline from the protocol.",
              size=11, color=T2).pack(anchor="w")

        # Scrollable body
        body = ScrollFrame(self)
        body.grid(row=1, column=0, sticky="nsew", padx=8, pady=0)

        f = ctk.CTkFrame(body, fg_color="transparent")
        f.pack(fill="x", padx=12, pady=(8, 16))

        # Title
        label(f, "Experiment title:", size=12, color=T2).pack(anchor="w")
        self.e_title = entry(f, "e.g. Western Blot – Sample A", width=460)
        if self.is_edit:
            self.e_title.insert(0, self.exp.get("title", ""))
        self.e_title.pack(anchor="w", pady=(2, 10))

        # Protocol selector
        label(f, "Protocol:", size=12, color=T2).pack(anchor="w")
        proto_names = [p.get("name", "?") for p in self.app.protocols]
        if not proto_names:
            proto_names = ["— No protocols yet —"]
        default_proto = ""
        if self.is_edit:
            default_proto = self.exp.get("protocolName", "")
        self._proto_var = ctk.StringVar(
            value=default_proto or (proto_names[0] if proto_names else ""))
        ctk.CTkOptionMenu(
            f, variable=self._proto_var, values=proto_names,
            font=(FONT, 12), width=460, corner_radius=R_SM,
            command=lambda _: self._update_preview(),
        ).pack(anchor="w", pady=(2, 10))

        # Date
        label(f, "Start date (YYYY-MM-DD):", size=12, color=T2).pack(anchor="w")
        default_date = datetime.now().strftime("%Y-%m-%d")
        if self.is_edit:
            try:
                default_date = datetime.fromtimestamp(
                    self.exp["plannedStart"] / 1000).strftime("%Y-%m-%d")
            except Exception:
                pass
        self.e_date = entry(f, "YYYY-MM-DD", width=220)
        self.e_date.insert(0, default_date)
        self.e_date.pack(anchor="w", pady=(2, 10))
        self.e_date.bind("<FocusOut>", lambda _: self._update_preview())

        # Start time
        label(f, "Start time (HH:MM, 24-hour):", size=12, color=T2).pack(anchor="w")
        default_time = "09:00"
        if self.is_edit:
            try:
                default_time = datetime.fromtimestamp(
                    self.exp["plannedStart"] / 1000).strftime("%H:%M")
            except Exception:
                pass
        self.e_time = entry(f, "09:00", width=120)
        self.e_time.insert(0, default_time)
        self.e_time.pack(anchor="w", pady=(2, 10))
        self.e_time.bind("<FocusOut>", lambda _: self._update_preview())

        # Notes
        label(f, "Notes (optional):", size=12, color=T2).pack(anchor="w")
        self.tb_notes = textbox(f, width=460, height=56)
        if self.is_edit:
            self.tb_notes.insert("0.0", self.exp.get("notes", ""))
        self.tb_notes.pack(anchor="w", pady=(2, 10))

        # Timeline preview panel
        self._preview_wrap = ctk.CTkFrame(
            f, fg_color=("#eef3fc", "#1e3358"), corner_radius=R_MD,
            border_width=1, border_color=("#c7d7f5", "#2d4a7a"))
        self._preview_wrap.pack(fill="x", pady=(0, 4))
        self._preview_inner = ctk.CTkFrame(self._preview_wrap, fg_color="transparent")
        self._preview_inner.pack(fill="x", padx=12, pady=10)
        label(self._preview_inner, "Select a protocol and fill dates to preview.",
              size=11, color=T3).pack(anchor="w")

        # Sticky footer
        footer = ctk.CTkFrame(self, fg_color=CARD, border_width=1,
                               border_color=CARD_B, corner_radius=0)
        footer.grid(row=2, column=0, sticky="ew")
        brow = ctk.CTkFrame(footer, fg_color="transparent")
        brow.pack(fill="x", padx=20, pady=12)
        btn(brow, "Cancel", self.destroy,
            color=("#e2e8f0", "#334155"), text_color=T1, width=90).pack(side="left")
        btn(brow, "Save" if self.is_edit else "Schedule", self._confirm,
            color=GREEN, text_color=("#fff", "#fff"), width=130).pack(side="right")

        self._update_preview()

    def _get_selected_proto(self):
        name = self._proto_var.get()
        return next((p for p in self.app.protocols if p.get("name") == name), None)

    def _update_preview(self):
        proto = self._get_selected_proto()
        if not proto:
            return
        try:
            start_dt = datetime.strptime(
                f"{self.e_date.get().strip()} {self.e_time.get().strip()}",
                "%Y-%m-%d %H:%M")
        except ValueError:
            return

        # Build scheduled steps
        self._scheduled_steps = []
        cur = start_dt
        for step in proto.get("steps", []):
            dur = (step.get("handsOnMinutes", 0) + step.get("waitMinutes", 0) +
                   step.get("bufferMinutes", 0)) or 5
            end = cur + timedelta(minutes=dur)
            self._scheduled_steps.append({
                "id":             new_id(),
                "protocolStepId": step.get("id", ""),
                "title":          step.get("title", "?") or "Untitled",
                "type":           step.get("type", "custom"),
                "plannedStart":   int(cur.timestamp() * 1000),
                "plannedEnd":     int(end.timestamp() * 1000),
                "durationMinutes": dur,
                "handsOnMinutes": step.get("handsOnMinutes", 0),
                "waitMinutes":    step.get("waitMinutes", 0),
                "bufferMinutes":  step.get("bufferMinutes", 0),
                "notes":          "",
            })
            cur = end
        self._exp_end_dt = cur
        self._render_preview(start_dt, cur)

    def _render_preview(self, start_dt, end_dt):
        for w in self._preview_inner.winfo_children():
            w.destroy()
        dur_min = int((end_dt - start_dt).total_seconds() / 60)

        hrow = ctk.CTkFrame(self._preview_inner, fg_color="transparent")
        hrow.pack(fill="x", pady=(0, 6))
        label(hrow, "Timeline Preview", size=11, weight="bold").pack(side="left")
        label(hrow,
              f"{start_dt.strftime('%I:%M %p')} → {end_dt.strftime('%I:%M %p')}  ·  {fmt_mins(dur_min)}",
              size=11, color=T2).pack(side="right")

        cat_colors = {"hands_on": "#3b82f6", "waiting": "#f97316",
                      "machine": "#8b5cf6", "note": "#eab308"}
        MAX_SHOW = 6
        for ss in self._scheduled_steps[:MAX_SHOW]:
            try:
                s = datetime.fromtimestamp(ss["plannedStart"] / 1000)
                e = datetime.fromtimestamp(ss["plannedEnd"] / 1000)
                t_str = f"{s.strftime('%I:%M')}–{e.strftime('%I:%M %p')}"
            except Exception:
                t_str = "?"
            row = ctk.CTkFrame(self._preview_inner, fg_color="transparent")
            row.pack(fill="x", pady=1)
            dot_c = cat_colors.get(_step_sched_cat(ss.get("type", "custom")), "#94a3b8")
            dot = ctk.CTkFrame(row, fg_color=dot_c, width=7, height=7, corner_radius=4)
            dot.pack(side="left", padx=(0, 7), pady=7)
            dot.pack_propagate(False)
            label(row, ss.get("title", "?"), size=11).pack(side="left")
            label(row, f"{t_str}  ·  {fmt_mins(ss['durationMinutes'])}",
                  size=10, color=T3).pack(side="right")
        if len(self._scheduled_steps) > MAX_SHOW:
            n = len(self._scheduled_steps) - MAX_SHOW
            label(self._preview_inner, f"  … and {n} more steps",
                  size=10, color=T3).pack(anchor="w", pady=(2, 0))

    def _confirm(self):
        title = self.e_title.get().strip()
        if not title:
            messagebox.showwarning("Missing title", "Please enter an experiment title.")
            return
        proto = self._get_selected_proto()
        if not proto:
            messagebox.showwarning("No protocol", "Please select a protocol.")
            return
        try:
            start_dt = datetime.strptime(
                f"{self.e_date.get().strip()} {self.e_time.get().strip()}",
                "%Y-%m-%d %H:%M")
        except ValueError:
            messagebox.showwarning("Invalid date/time", "Use YYYY-MM-DD and HH:MM format.")
            return
        if not self._scheduled_steps:
            self._update_preview()
        start_ts = int(start_dt.timestamp() * 1000)
        end_ts   = (int(self._exp_end_dt.timestamp() * 1000)
                    if self._exp_end_dt else start_ts + 3600000)
        # Build timelineBlocks from scheduledSteps
        tl_blocks = []
        for ss in self._scheduled_steps:
            tl_blocks.append({
                "id":                   ss.get("id", new_id()),
                "blockType":            "protocol_step",
                "title":                ss.get("title", "Step"),
                "type":                 ss.get("type", "custom"),
                "startTime":            ss.get("plannedStart", start_ts),
                "endTime":              ss.get("plannedEnd",   start_ts),
                "durationMinutes":      ss.get("durationMinutes", 5),
                "handsOnMinutes":       ss.get("handsOnMinutes", 0),
                "waitMinutes":          ss.get("waitMinutes",    0),
                "notes":                ss.get("notes", ""),
                "status":               "planned",
                "sourceProtocolStepId": ss.get("protocolStepId", ""),
                "isParallelTask":       False,
                "parallelWithBlockId":  None,
                "keepTime":             True,
            })
        # If editing and the experiment already has timelineBlocks with user edits,
        # preserve them rather than regenerating from scratch.
        if self.is_edit and self.exp.get("timelineBlocks"):
            existing_ids = {b["sourceProtocolStepId"] for b in self.exp["timelineBlocks"]
                            if b.get("blockType") == "protocol_step"}
            gen_ids      = {b["sourceProtocolStepId"] for b in tl_blocks}
            # Only replace if protocol changed (different step IDs)
            if existing_ids == gen_ids:
                tl_blocks = self.exp["timelineBlocks"]

        exp = {
            "id":               self.exp["id"] if self.is_edit else new_id(),
            "title":            title,
            "protocolId":       proto["id"],
            "protocolName":     proto.get("name", "?"),
            "protocolSnapshot": proto,
            "date":             start_dt.strftime("%Y-%m-%d"),
            "plannedStart":     start_ts,
            "plannedEnd":       end_ts,
            "scheduledSteps":   self._scheduled_steps,
            "timelineBlocks":   tl_blocks,
            "notes":            self.tb_notes.get("0.0", "end").strip(),
            "status":           (self.exp.get("status", "planned")
                                 if self.is_edit else "planned"),
        }
        self.destroy()
        self.on_save(exp)


# ══════════════════════════════════════════════════════════════════════════════
# BlockEditDialog — edit any timeline block (type-aware fields)
# ══════════════════════════════════════════════════════════════════════════════
class BlockEditDialog(ctk.CTkToplevel):
    """Edit a single timeline block with type-aware fields."""

    def __init__(self, parent, block, on_save):
        super().__init__(parent)
        self.withdraw()
        self._block   = block
        self._on_save = on_save
        self.title("Edit Block")
        self.resizable(True, False)
        self._build()
        _show_dialog(self, parent, 460, 520)
        self.bind("<Escape>", lambda e: self._close())

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # ── Header ────────────────────────────────────────────────────────────
        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 4))
        btype = self._block.get("blockType", "protocol_step")
        _, accent, _ = BLOCK_TYPE_COLORS.get(btype, ("#f8fafc", "#64748b", "#d1d5db"))
        label(hdr, "Edit Block", size=14, weight="bold").pack(side="left")
        badge_f = ctk.CTkFrame(hdr, fg_color=accent, corner_radius=R_XS)
        badge_f.pack(side="left", padx=(10, 0))
        ctk.CTkLabel(badge_f, text=BLOCK_TYPE_LABELS.get(btype, btype),
                     font=(FONT, 10), text_color=("#fff", "#fff")).pack(padx=7, pady=2)

        # ── Scrollable body ───────────────────────────────────────────────────
        body = ScrollFrame(self)
        body.grid(row=1, column=0, sticky="nsew", padx=8, pady=0)
        f = ctk.CTkFrame(body, fg_color="transparent")
        f.pack(fill="x", padx=12, pady=(8, 16))

        def fld(lbl_txt, var, w=400):
            r = ctk.CTkFrame(f, fg_color="transparent")
            r.pack(fill="x", pady=(0, 8))
            label(r, lbl_txt, size=11, color=T3).pack(anchor="w")
            e = ctk.CTkEntry(r, textvariable=var, font=(FONT, 12), height=32, width=w,
                             corner_radius=R_SM, border_color=CARD_B)
            e.pack(anchor="w", pady=(2, 0))
            return e

        # Title
        self._title_var = ctk.StringVar(value=self._block.get("title", ""))
        title_e = fld("Title", self._title_var)
        title_e.focus()

        # Block type
        label(f, "Block type:", size=11, color=T3).pack(anchor="w")
        type_vals = list(BLOCK_TYPE_LABELS.values())
        cur_type  = BLOCK_TYPE_LABELS.get(self._block.get("blockType", "protocol_step"),
                                          "Protocol Step")
        self._type_var = ctk.StringVar(value=cur_type)
        ctk.CTkOptionMenu(
            f, variable=self._type_var, values=type_vals,
            font=(FONT, 12), width=200, corner_radius=R_SM,
            command=self._on_type_change,
        ).pack(anchor="w", pady=(2, 8))

        # Duration
        self._dur_var = ctk.StringVar(value=str(self._block.get("durationMinutes", 5)))
        fld("Duration (minutes)", self._dur_var, w=120)

        # Protocol-step-only fields (hands-on / wait)
        self._ps_frame = ctk.CTkFrame(f, fg_color="transparent")
        ho = self._block.get("handsOnMinutes", 0)
        wt = self._block.get("waitMinutes",    0)
        self._ho_var = ctk.StringVar(value=str(ho))
        self._wt_var = ctk.StringVar(value=str(wt))
        lf = self._ps_frame
        label(lf, "Hands-on (minutes):", size=11, color=T3).pack(anchor="w")
        ctk.CTkEntry(lf, textvariable=self._ho_var, font=(FONT, 12), height=32,
                     width=120, corner_radius=R_SM, border_color=CARD_B
                     ).pack(anchor="w", pady=(2, 8))
        label(lf, "Wait / incubation (minutes):", size=11, color=T3).pack(anchor="w")
        ctk.CTkEntry(lf, textvariable=self._wt_var, font=(FONT, 12), height=32,
                     width=120, corner_radius=R_SM, border_color=CARD_B
                     ).pack(anchor="w", pady=(2, 8))
        if self._block.get("blockType", "protocol_step") == "protocol_step":
            self._ps_frame.pack(fill="x")

        # Status
        label(f, "Status:", size=11, color=T3).pack(anchor="w")
        self._status_var = ctk.StringVar(value=self._block.get("status", "planned"))
        ctk.CTkOptionMenu(
            f, variable=self._status_var,
            values=["planned", "done", "skipped", "canceled", "modified"],
            font=(FONT, 12), width=200, corner_radius=R_SM,
        ).pack(anchor="w", pady=(2, 8))

        # Notes
        self._notes_var = ctk.StringVar(value=self._block.get("notes", ""))
        fld("Notes (optional)", self._notes_var)

        # ── Footer ────────────────────────────────────────────────────────────
        footer = ctk.CTkFrame(self, fg_color=CARD, border_width=1,
                               border_color=CARD_B, corner_radius=0)
        footer.grid(row=2, column=0, sticky="ew")
        brow = ctk.CTkFrame(footer, fg_color="transparent")
        brow.pack(fill="x", padx=20, pady=12)
        btn(brow, "Cancel", self._close,
            color=("#e2e8f0", "#334155"), text_color=T1,
            height=34, width=90).pack(side="left")
        btn(brow, "Save", self._save,
            color=ACC, text_color=("#fff", "#fff"),
            height=34, width=90).pack(side="right")

    def _on_type_change(self, val):
        key = next((k for k, v in BLOCK_TYPE_LABELS.items() if v == val), "custom")
        if key == "protocol_step":
            self._ps_frame.pack(fill="x")
        else:
            self._ps_frame.pack_forget()

    def _save(self):
        try:
            dur = max(1, int(self._dur_var.get()))
        except ValueError:
            dur = self._block.get("durationMinutes", 5)
        type_label = self._type_var.get()
        block_type = next((k for k, v in BLOCK_TYPE_LABELS.items() if v == type_label),
                          self._block.get("blockType", "custom"))
        updated = dict(
            self._block,
            title           = self._title_var.get().strip() or "Block",
            blockType       = block_type,
            durationMinutes = dur,
            status          = self._status_var.get(),
            notes           = self._notes_var.get().strip(),
        )
        if block_type == "protocol_step":
            try: updated["handsOnMinutes"] = max(0, int(self._ho_var.get()))
            except ValueError: pass
            try: updated["waitMinutes"]    = max(0, int(self._wt_var.get()))
            except ValueError: pass
        cb = self._on_save
        self._close()
        cb(updated)

    def _close(self):
        try:
            self.grab_release()
            self.destroy()
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# ScheduleBlockDialog — choose type + fill basic fields for a new timeline block
# ══════════════════════════════════════════════════════════════════════════════
class ScheduleBlockDialog(ctk.CTkToplevel):
    """Create a new schedule timeline block (type + title + duration)."""

    def __init__(self, parent, initial_type, on_add):
        super().__init__(parent)
        self.withdraw()
        self._on_add      = on_add
        self._init_type   = initial_type
        self.title("Add Block")
        self.resizable(False, False)
        self._build()
        _show_dialog(self, parent, 420, 420)
        self.bind("<Escape>", lambda e: self._close())

    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        hdr = ctk.CTkFrame(self, fg_color="transparent")
        hdr.grid(row=0, column=0, sticky="ew", padx=20, pady=(16, 4))
        label(hdr, "Add Block to Timeline", size=14, weight="bold").pack(anchor="w")

        body = ScrollFrame(self)
        body.grid(row=1, column=0, sticky="nsew", padx=8, pady=0)
        f = ctk.CTkFrame(body, fg_color="transparent")
        f.pack(fill="x", padx=12, pady=(8, 16))

        # Block type
        label(f, "Block type:", size=11, color=T3).pack(anchor="w")
        type_vals = list(BLOCK_TYPE_LABELS.values())
        init_lbl  = BLOCK_TYPE_LABELS.get(self._init_type, "Custom")
        self._type_var = ctk.StringVar(value=init_lbl)
        ctk.CTkOptionMenu(
            f, variable=self._type_var, values=type_vals,
            font=(FONT, 12), width=220, corner_radius=R_SM,
            command=self._on_type_change,
        ).pack(anchor="w", pady=(2, 10))

        # Quick-pick suggestions for break/task
        self._sugg_frame = ctk.CTkFrame(f, fg_color="transparent")
        self._sugg_frame.pack(fill="x", pady=(0, 6))
        self._refresh_suggestions(self._init_type)

        # Title
        label(f, "Title:", size=11, color=T3).pack(anchor="w")
        self._title_var = ctk.StringVar()
        self._title_entry = ctk.CTkEntry(
            f, textvariable=self._title_var, font=(FONT, 12),
            height=32, width=380, corner_radius=R_SM, border_color=CARD_B,
            placeholder_text="e.g. Lunch break")
        self._title_entry.pack(anchor="w", pady=(2, 10))
        self._title_entry.focus()

        # Duration
        label(f, "Duration (minutes):", size=11, color=T3).pack(anchor="w")
        self._dur_var = ctk.StringVar(value="30")
        ctk.CTkEntry(f, textvariable=self._dur_var, font=(FONT, 12),
                     height=32, width=120, corner_radius=R_SM, border_color=CARD_B
                     ).pack(anchor="w", pady=(2, 10))

        # Notes
        label(f, "Notes (optional):", size=11, color=T3).pack(anchor="w")
        self._notes_var = ctk.StringVar()
        ctk.CTkEntry(f, textvariable=self._notes_var, font=(FONT, 12),
                     height=32, width=380, corner_radius=R_SM, border_color=CARD_B,
                     placeholder_text="Optional description"
                     ).pack(anchor="w", pady=(2, 10))

        footer = ctk.CTkFrame(self, fg_color=CARD, border_width=1,
                               border_color=CARD_B, corner_radius=0)
        footer.grid(row=2, column=0, sticky="ew")
        brow = ctk.CTkFrame(footer, fg_color="transparent")
        brow.pack(fill="x", padx=20, pady=12)
        btn(brow, "Cancel", self._close,
            color=("#e2e8f0", "#334155"), text_color=T1,
            height=34, width=90).pack(side="left")
        btn(brow, "Add", self._add,
            color=ACC, text_color=("#fff", "#fff"),
            height=34, width=90).pack(side="right")

    _SUGGESTIONS = {
        "break":    ["Lunch break", "Dinner break", "Rest", "Meeting",
                     "Waiting for reagent", "Waiting for equipment"],
        "task":     ["Prepare buffer", "Check cells", "Thaw reagent",
                     "Clean bench", "Set up imaging", "Make solution"],
        "note":     ["General note", "Observation", "Reminder"],
        "decision": ["Skip step", "Repeat wash", "Discard sample",
                     "Continue tomorrow"],
        "custom":   ["Custom block"],
    }

    def _on_type_change(self, val):
        key = next((k for k, v in BLOCK_TYPE_LABELS.items() if v == val), "custom")
        self._refresh_suggestions(key)

    def _refresh_suggestions(self, block_type):
        for w in self._sugg_frame.winfo_children():
            w.destroy()
        suggestions = self._SUGGESTIONS.get(block_type, [])
        if not suggestions:
            return
        row = ctk.CTkFrame(self._sugg_frame, fg_color="transparent")
        row.pack(anchor="w")
        label(row, "Quick pick:", size=10, color=T3).pack(side="left", padx=(0, 6))
        for s in suggestions[:4]:
            ctk.CTkButton(
                row, text=s, width=0, height=22,
                fg_color=("#e2e8f0", "#334155"), hover_color=("#dbeafe", "#1e3a5f"),
                text_color=T2, font=(FONT, 10), corner_radius=R_XS,
                command=lambda v=s: self._title_var.set(v),
            ).pack(side="left", padx=(0, 4))

    def _add(self):
        title = self._title_var.get().strip()
        if not title:
            return
        try:
            dur = max(1, int(self._dur_var.get()))
        except ValueError:
            dur = 30
        type_label = self._type_var.get()
        block_type = next((k for k, v in BLOCK_TYPE_LABELS.items() if v == type_label),
                          "custom")
        new_blk = {
            "id":              new_id(),
            "blockType":       block_type,
            "title":           title,
            "type":            "custom",
            "startTime":       0,   # will be recalculated
            "endTime":         0,
            "durationMinutes": dur,
            "handsOnMinutes":  0,
            "waitMinutes":     0,
            "notes":           self._notes_var.get().strip(),
            "status":          "planned",
            "sourceProtocolStepId": "",
            "isParallelTask":  False,
            "parallelWithBlockId": None,
            "keepTime":        True,
        }
        cb = self._on_add
        self._close()
        cb(new_blk)

    def _close(self):
        try:
            self.grab_release()
            self.destroy()
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# SkipBlockDialog — ask whether to remove or keep time when skipping
# ══════════════════════════════════════════════════════════════════════════════
class SkipBlockDialog(ctk.CTkToplevel):
    """Ask user how to handle time for a skipped block."""

    def __init__(self, parent, block_title, on_choose):
        super().__init__(parent)
        self.withdraw()
        self._on_choose = on_choose
        self.title("Skip Block")
        self.resizable(False, False)
        self._build(block_title)
        _show_dialog(self, parent, 400, 240)
        self.bind("<Escape>", lambda e: self._close())

    def _build(self, block_title):
        f = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG,
                          border_width=1, border_color=CARD_B)
        f.pack(fill="both", expand=True, padx=4, pady=4)
        f.grid_columnconfigure(0, weight=1)

        label(f, f"Skip: {block_title[:40]}", size=13, weight="bold").pack(
            padx=16, pady=(14, 4), anchor="w")
        label(f, "Should this skipped block still occupy time?",
              size=11, color=T2).pack(padx=16, pady=(0, 12), anchor="w")

        separator(f).pack(fill="x", padx=12, pady=(0, 12))

        btn(f, "Remove time  (shift later blocks earlier)",
            lambda: self._choose(False),
            color=("#e2e8f0", "#334155"), text_color=T1,
            height=34, width=340).pack(padx=16, pady=(0, 6))
        btn(f, "Keep time  (show as grayed-out placeholder)",
            lambda: self._choose(True),
            color=("#fef3c7", "#451a03"), text_color=("#78350f", "#fef08a"),
            height=34, width=340).pack(padx=16, pady=(0, 6))
        btn(f, "Cancel", self._close,
            color="transparent", text_color=T3,
            height=28, width=80).pack(pady=(4, 10))

    def _choose(self, keep_time):
        cb = self._on_choose
        self._close()
        cb(keep_time)

    def _close(self):
        try:
            self.grab_release()
            self.destroy()
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# BlockContextMenu — right-click popup menu for a block card
# ══════════════════════════════════════════════════════════════════════════════
class BlockContextMenu(ctk.CTkToplevel):
    """Floating right-click context menu for a timeline block."""

    def __init__(self, parent, x, y, actions):
        """
        actions: list of (label_str, command) or None for separator.
        """
        super().__init__(parent)
        self.withdraw()
        self.overrideredirect(True)
        self._build(actions)
        self.update_idletasks()
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        w  = self.winfo_reqwidth()
        h  = self.winfo_reqheight()
        x  = min(x, sw - w - 4)
        y  = min(y, sh - h - 4)
        self.geometry(f"+{x}+{y}")
        self.deiconify()
        self.lift()
        self.focus_set()
        # Dismiss on click outside
        self.bind("<FocusOut>", lambda e: self._close())
        self.bind("<Escape>",   lambda e: self._close())

    def _build(self, actions):
        f = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_SM,
                          border_width=1, border_color=CARD_B)
        f.pack(fill="both", expand=True, padx=1, pady=1)
        for item in actions:
            if item is None:
                ctk.CTkFrame(f, height=1, fg_color=CARD_B, corner_radius=0
                             ).pack(fill="x", padx=4, pady=2)
            else:
                lbl, cmd = item
                danger = lbl.startswith("Delete") or lbl.startswith("Cancel")
                ctk.CTkButton(
                    f, text=lbl, anchor="w",
                    fg_color="transparent",
                    hover_color=("#eff6ff", "#1e3a5f") if not danger else ("#fee2e2", "#450a0a"),
                    text_color=DANGER if danger else T1,
                    font=(FONT, 12), height=30, corner_radius=R_XS,
                    command=lambda c=cmd: (self._close(), c()),
                ).pack(fill="x", padx=4, pady=1)

    def _close(self):
        try:
            self.destroy()
        except Exception:
            pass


class SchedulePage(PageBase):
    """Experiment Calendar — schedule and view protocols on a time grid."""

    # ── Calendar geometry constants ───────────────────────────────────────────
    HOUR_START   = 7    # first hour shown (7 AM)
    HOUR_END     = 22   # last hour shown  (10 PM)
    PX_PER_HR    = 64   # canvas pixels per hour; 1 px ≈ 56.25 s
    TIME_W       = 68   # width of time-label column (px)
    MIN_SNAP     = 15   # drag snap interval (minutes)

    # Color palettes — (light fill, medium fill, accent/border)
    PALETTES = [
        ("#dbeafe", "#bfdbfe", "#3b82f6"),
        ("#dcfce7", "#bbf7d0", "#22c55e"),
        ("#ede9fe", "#ddd6fe", "#8b5cf6"),
        ("#fce7f3", "#fbcfe8", "#ec4899"),
        ("#cffafe", "#a5f3fc", "#06b6d4"),
        ("#fef3c7", "#fde68a", "#f59e0b"),
        ("#fee2e2", "#fecaca", "#ef4444"),
    ]

    # Step category → (card bg tuple, border hex)
    _CAT_STYLE = {
        "hands_on": (("#dbeafe", "#1e3a5f"), "#93c5fd"),
        "waiting":  (("#fff7ed", "#431407"), "#fed7aa"),
        "machine":  (("#f5f3ff", "#1e1b4b"), "#c4b5fd"),
        "note":     (("#fefce8", "#422006"), "#fef08a"),
    }
    _CAT_DOT = {
        "hands_on": "#3b82f6",
        "waiting":  "#f97316",
        "machine":  "#8b5cf6",
        "note":     "#eab308",
    }

    # Step block colours: (fill, border, accent) per category
    _STEP_COLORS = {
        "hands_on": ("#dbeafe", "#93c5fd", "#3b82f6"),
        "waiting":  ("#fff7ed", "#fed7aa", "#f97316"),
        "machine":  ("#f5f3ff", "#c4b5fd", "#8b5cf6"),
        "note":     ("#fefce8", "#fde68a", "#eab308"),
    }

    def __init__(self, parent, app):
        super().__init__(parent, app)
        self._anchor         = datetime.now().date()  # the calendar's "focus" date
        self._view           = "workweek"  # "day" | "workweek" | "week"
        self._selected_id    = None
        self._drag           = None
        self._layout         = {}
        self._detail_visible = False
        self._tooltip_job    = None         # pending after() ID for hover tooltip
        self._build()

    # ── refresh ───────────────────────────────────────────────────────────────
    def refresh(self):
        self.app.schedule = load_schedule()
        existing_ids = {e["id"] for e in self.app.schedule}
        self._render_calendar()
        # Restore detail panel if the previously selected experiment still exists
        if self._selected_id and self._detail_visible:
            if self._selected_id in existing_ids:
                self._show_detail(self._selected_id)
            else:
                self._selected_id    = None
                self._detail_visible = False

    # ── build UI ──────────────────────────────────────────────────────────────
    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=0)
        self.grid_rowconfigure(1, weight=1)

        # ── Toolbar ───────────────────────────────────────────────────────────
        tb = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG)
        tb.grid(row=0, column=0, columnspan=2, sticky="ew", padx=12, pady=(12, 6))
        tb.grid_columnconfigure(1, weight=1)

        # Left: nav buttons
        nav_f = ctk.CTkFrame(tb, fg_color="transparent")
        nav_f.grid(row=0, column=0, padx=12, pady=10, sticky="w")
        btn(nav_f, "◀", lambda: self._nav_period(-1),
            color=("#e2e8f0", "#334155"), text_color=T1, width=34, height=32, size=12
            ).pack(side="left", padx=(0, 4))
        btn(nav_f, "Today", self._goto_today,
            color=("#e2e8f0", "#334155"), text_color=T1, width=60, height=32, size=12
            ).pack(side="left", padx=(0, 4))
        btn(nav_f, "▶", lambda: self._nav_period(1),
            color=("#e2e8f0", "#334155"), text_color=T1, width=34, height=32, size=12
            ).pack(side="left", padx=(0, 8))
        btn(nav_f, "📅", self._go_to_date,
            color=("#e2e8f0", "#334155"), text_color=T1, width=34, height=32, size=12
            ).pack(side="left")

        # Center: date range label (read-only)
        self._lbl_range = label(tb, "", size=14, weight="bold")
        self._lbl_range.grid(row=0, column=1, padx=8)

        # Right: view selector + add button
        right_f = ctk.CTkFrame(tb, fg_color="transparent")
        right_f.grid(row=0, column=2, padx=12, pady=10, sticky="e")
        self._view_var = ctk.StringVar(value="workweek")
        for v, lbl_txt in [("day", "Day"), ("workweek", "Work Week"), ("week", "Week")]:
            ctk.CTkRadioButton(
                right_f, text=lbl_txt, variable=self._view_var, value=v,
                font=(FONT, 12), radiobutton_width=16, radiobutton_height=16,
                command=self._on_view_change,
            ).pack(side="left", padx=8)
        btn(right_f, "＋  Schedule Experiment", self._new_experiment,
            color=ACC, text_color=("#fff", "#fff"), width=180, height=32, size=12
            ).pack(side="left", padx=(18, 0))

        # ── Calendar frame ─────────────────────────────────────────────────────
        # Layout: col 0 = header canvas + body canvas (same width)
        #         col 1 = scrollbar (fixed width = 14 px)
        # Having both canvases in col 0 guarantees they are always the same
        # pixel width → day column lines in header align perfectly with the grid.
        cal_outer = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG,
                                  border_width=1, border_color=CARD_B)
        cal_outer.grid(row=1, column=0, sticky="nsew", padx=12, pady=(0, 12))
        cal_outer.grid_columnconfigure(0, weight=1)
        cal_outer.grid_columnconfigure(1, minsize=14, weight=0)
        cal_outer.grid_rowconfigure(1, weight=1)

        # ── Non-scrolling day-header canvas (row 0, col 0) ────────────────────
        # Height fixed at HDR_H; drawn with exact same col_w math as body canvas.
        self._hdr_canvas = tk.Canvas(cal_outer, height=56, bg="#ffffff",
                                      highlightthickness=0)
        self._hdr_canvas.grid(row=0, column=0, sticky="ew")
        self._hdr_canvas.bind("<Configure>", self._on_hdr_resize)
        self._hdr_canvas.bind("<Button-1>",  self._on_hdr_click)
        self._hdr_canvas.bind("<Motion>",    self._on_hdr_motion)
        self._hdr_canvas.bind("<Leave>",
                              lambda e: self._hdr_canvas.config(cursor="arrow"))

        # ── Scrollable body canvas (row 1, col 0) ─────────────────────────────
        self._canvas = tk.Canvas(cal_outer, bg="#f8fafc",
                                  highlightthickness=0, cursor="arrow")
        self._canvas.grid(row=1, column=0, sticky="nsew")

        # ── Scrollbar (row 1, col 1) ───────────────────────────────────────────
        vsb = ctk.CTkScrollbar(cal_outer, orientation="vertical",
                                command=self._canvas.yview,
                                button_color=("#cbd5e1", "#475569"),
                                button_hover_color=("#94a3b8", "#64748b"))
        vsb.grid(row=1, column=1, sticky="ns")
        self._canvas.configure(yscrollcommand=vsb.set)

        self._canvas.bind("<Configure>",        self._on_canvas_resize)
        self._canvas.bind("<MouseWheel>",       self._canvas_wheel)
        self._canvas.bind("<Button-4>",         lambda e: self._canvas.yview_scroll(-3, "units"))
        self._canvas.bind("<Button-5>",         lambda e: self._canvas.yview_scroll(3, "units"))
        self._canvas.bind("<Button-1>",         self._on_click)
        self._canvas.bind("<Double-Button-1>",  self._on_double_click)
        self._canvas.bind("<B1-Motion>",        self._on_drag)
        self._canvas.bind("<ButtonRelease-1>",  self._on_release)
        self._canvas.bind("<Motion>",           self._on_canvas_motion)
        self._canvas.bind("<Leave>",            self._on_canvas_leave)

        # ── Detail panel (shown in col 1 of SchedulePage when an exp is clicked)
        self._detail_frame = ctk.CTkFrame(self, fg_color=CARD, corner_radius=R_LG,
                                           border_width=1, border_color=CARD_B, width=300)
        self._detail_frame.grid_propagate(False)

        # Scroll to ~8 AM once the canvas has been laid out
        self.after(400, lambda: self._canvas.yview_moveto(
            max(0.0, (8 - self.HOUR_START) / (self.HOUR_END - self.HOUR_START) - 0.05)))

    # ── Wheel ─────────────────────────────────────────────────────────────────
    def _canvas_wheel(self, event):
        delta = getattr(event, "delta", 0)
        if abs(delta) >= 120:
            units = int(-delta / 120) * 3
        else:
            units = int(-delta / 3) or (-1 if delta > 0 else 1)
        self._canvas.yview_scroll(units, "units")

    # ── Navigation ────────────────────────────────────────────────────────────
    def _nav_period(self, direction):
        """Move forward/backward one period: 1 day in day view, 1 week otherwise."""
        if self._view == "day":
            self._anchor += timedelta(days=direction)
        else:
            self._anchor += timedelta(weeks=direction)
        self._render_calendar()

    # Keep old name as alias so any existing references survive
    def _nav_week(self, direction):
        self._nav_period(direction)

    def _goto_today(self):
        self._anchor = datetime.now().date()
        self._render_calendar()

    def _on_view_change(self):
        self._view = self._view_var.get()
        self._render_calendar()

    def _go_to_date(self):
        """Open a mini calendar picker so the user can jump to any date."""
        def on_pick(d):
            self._anchor = d
            self._render_calendar()
        MiniCalendarPicker(self, self._anchor, on_pick)

    # ── Date helpers ──────────────────────────────────────────────────────────
    def _get_days(self):
        anchor = self._anchor
        if self._view == "day":
            return [anchor]
        # For week views, show the Monday-anchored week
        monday = anchor - timedelta(days=anchor.weekday())
        if self._view == "workweek":
            return [monday + timedelta(days=i) for i in range(5)]
        else:  # "week" — Monday through Sunday
            return [monday + timedelta(days=i) for i in range(7)]

    def _ts_to_y(self, ts_ms):
        dt = datetime.fromtimestamp(ts_ms / 1000)
        mins = (dt.hour - self.HOUR_START) * 60 + dt.minute
        return mins * self.PX_PER_HR / 60

    def _y_to_ts(self, y, date):
        total_min = int(y * 60 / self.PX_PER_HR) + self.HOUR_START * 60
        total_min = round(total_min / self.MIN_SNAP) * self.MIN_SNAP
        total_min = max(self.HOUR_START * 60, min((self.HOUR_END - 1) * 60, total_min))
        dt = datetime(date.year, date.month, date.day, total_min // 60, total_min % 60)
        return int(dt.timestamp() * 1000)

    def _exp_palette(self, exp_id):
        return self.PALETTES[abs(hash(exp_id)) % len(self.PALETTES)]

    # ── Rendering ─────────────────────────────────────────────────────────────
    @staticmethod
    def _day_str(d):
        """Return day number without zero-padding (cross-platform)."""
        return str(d.day)

    def _render_calendar(self):
        days = self._get_days()
        # Build range label without platform-specific strftime codes (%-d)
        if len(days) == 1:
            d = days[0]
            self._lbl_range.configure(
                text=f"{d.strftime('%A, %B')} {self._day_str(d)}, {d.year}")
        else:
            d0, d1 = days[0], days[-1]
            start_s = f"{d0.strftime('%b')} {self._day_str(d0)}"
            if d0.month == d1.month:
                end_s = f"{self._day_str(d1)}, {d1.year}"
            else:
                end_s = f"{d1.strftime('%b')} {self._day_str(d1)}, {d1.year}"
            self._lbl_range.configure(text=f"{start_s} – {end_s}")
        self._draw_headers(days)
        self._draw_grid(days)

    def _on_hdr_resize(self, event=None):
        """Header canvas was resized — redraw headers so they stay aligned."""
        days = self._get_days()
        self._draw_headers(days)

    def _on_hdr_click(self, event):
        """Click on a day column header → set anchor to that date."""
        days  = self._get_days()
        cw    = max(self._hdr_canvas.winfo_width(), 300)
        col_w = max(80, (cw - self.TIME_W) // len(days))
        x     = event.x
        if x < self.TIME_W:
            return  # click landed in the time-label gutter
        col_i = (x - self.TIME_W) // col_w
        if 0 <= col_i < len(days):
            self._anchor = days[col_i]
            self._render_calendar()

    def _on_hdr_motion(self, event):
        """Change cursor to hand2 over clickable day columns."""
        days  = self._get_days()
        cw    = max(self._hdr_canvas.winfo_width(), 300)
        col_w = max(80, (cw - self.TIME_W) // len(days))
        x     = event.x
        if x >= self.TIME_W:
            col_i = (x - self.TIME_W) // col_w
            if 0 <= col_i < len(days):
                self._hdr_canvas.config(cursor="hand2")
                return
        self._hdr_canvas.config(cursor="arrow")

    def _draw_headers(self, days):
        """Draw day-column headers on _hdr_canvas.

        Visual states (in priority order):
          • anchor + today  → filled blue circle, white bold text
          • anchor only     → filled blue circle, white bold text
          • today only      → light-blue filled circle + blue ring, blue bold text
          • neither         → no circle, dark text
        Circle geometry: center=(cx, y_num), radius=r — oval and text share
        the same centre point so they are always perfectly aligned.
        """
        c      = self._hdr_canvas
        cw     = max(c.winfo_width(), 300)
        h      = 56                       # fixed header height
        n      = len(days)
        col_w  = max(80, (cw - self.TIME_W) // n)
        today  = datetime.now().date()
        anchor = self._anchor             # currently selected / focused date

        # Geometry constants — both oval and text use (cx, y_num) as centre
        y_abbr = 15   # y for weekday abbreviation
        y_num  = 39   # y for date number (and circle centre)
        r      = 13   # circle radius

        c.delete("all")

        # ── Backgrounds ──────────────────────────────────────────────────────
        c.create_rectangle(0, 0, cw, h, fill="#ffffff", outline="")
        c.create_rectangle(0, 0, self.TIME_W, h, fill="#f8fafc", outline="")

        for i, d in enumerate(days):
            is_today  = (d == today)
            is_anchor = (d == anchor)
            cx = self.TIME_W + i * col_w + col_w // 2   # column centre x

            # ── Weekday abbreviation (e.g. "Mon") ────────────────────────────
            abbr_color = "#3b82f6" if is_today else "#94a3b8"
            c.create_text(cx, y_abbr, anchor="center",
                           text=d.strftime("%a"),
                           font=(FONT, 10), fill=abbr_color)

            # ── Date number + optional circle ────────────────────────────────
            if is_anchor:
                # Selected date → solid blue circle, white bold text
                c.create_oval(cx - r, y_num - r, cx + r, y_num + r,
                               fill="#3b82f6", outline="",
                               tags="circle")
                c.create_text(cx, y_num, anchor="center",
                               text=self._day_str(d),
                               font=(FONT, 13, "bold"), fill="#ffffff",
                               tags="top_text")
            elif is_today:
                # Today (not selected) → light-blue fill with blue ring
                c.create_oval(cx - r, y_num - r, cx + r, y_num + r,
                               fill="#dbeafe", outline="#3b82f6", width=2,
                               tags="circle")
                c.create_text(cx, y_num, anchor="center",
                               text=self._day_str(d),
                               font=(FONT, 13, "bold"), fill="#1d4ed8",
                               tags="top_text")
            else:
                # Plain date number — no circle
                c.create_text(cx, y_num, anchor="center",
                               text=self._day_str(d),
                               font=(FONT, 13), fill="#0f172a")

        # Raise all text on top of circles so they render above the ovals
        c.tag_raise("top_text")

        # ── Bottom border + time-axis divider ─────────────────────────────────
        c.create_line(0, h - 1, cw, h - 1, fill="#e2e8f0", width=1)
        c.create_line(self.TIME_W, 0, self.TIME_W, h, fill="#cbd5e1", width=2)

    def _on_canvas_resize(self, event):
        days = self._get_days()
        self._draw_headers(days)  # keep header in sync with body width
        self._draw_grid(days)

    def _draw_grid(self, days):
        c = self._canvas
        c.delete("all")

        n_days  = len(days)
        cw      = max(c.winfo_width(), 300)
        ch      = (self.HOUR_END - self.HOUR_START) * self.PX_PER_HR + 20
        col_w   = max(80, (cw - self.TIME_W) // n_days)
        today   = datetime.now().date()

        self._layout = {"days": days, "col_w": col_w, "ch": ch}
        c.configure(scrollregion=(0, 0, cw, ch))

        # ── Day column backgrounds ─────────────────────────────────────────────
        anchor = self._anchor
        for i, d in enumerate(days):
            x1 = self.TIME_W + i * col_w
            x2 = self.TIME_W + (i + 1) * col_w
            if d == today:
                fill = "#eff6ff"           # today → blue tint
            elif d == anchor:
                fill = "#f0f9ff"           # selected (not today) → lighter blue tint
            else:
                fill = "#f8fafc" if i % 2 == 0 else "#f1f5f9"
            c.create_rectangle(x1, 0, x2, ch, fill=fill, outline="")

        # ── Hour grid lines + time labels ──────────────────────────────────────
        for h in range(self.HOUR_START, self.HOUR_END + 1):
            y = (h - self.HOUR_START) * self.PX_PER_HR
            # Label
            if h < self.HOUR_END:
                hr12 = h % 12 or 12
                ampm = "AM" if h < 12 else "PM"
                c.create_text(self.TIME_W - 6, y + 2, anchor="ne",
                               text=f"{hr12}:00 {ampm}",
                               font=(FONT, 9), fill="#94a3b8")
            # Full-hour line
            c.create_line(self.TIME_W, y, cw, y, fill="#e2e8f0", width=1)
            # Half-hour dashed
            if h < self.HOUR_END:
                yh = y + self.PX_PER_HR // 2
                c.create_line(self.TIME_W, yh, cw, yh,
                               fill="#f1f5f9", width=1, dash=(3, 5))

        # ── Column separators ──────────────────────────────────────────────────
        for i in range(n_days + 1):
            x = self.TIME_W + i * col_w
            c.create_line(x, 0, x, ch, fill="#e2e8f0", width=1)
        c.create_line(self.TIME_W, 0, self.TIME_W, ch, fill="#cbd5e1", width=2)

        # ── Current time indicator ─────────────────────────────────────────────
        now = datetime.now()
        if (now.date() in days and self.HOUR_START <= now.hour < self.HOUR_END):
            now_y = ((now.hour - self.HOUR_START) * 60 + now.minute) * self.PX_PER_HR / 60
            col_i = days.index(now.date())
            x1 = self.TIME_W + col_i * col_w
            x2 = self.TIME_W + (col_i + 1) * col_w
            c.create_line(x1, now_y, x2, now_y, fill="#ef4444", width=2)
            c.create_oval(x1 - 5, now_y - 5, x1 + 5, now_y + 5,
                           fill="#ef4444", outline="white", width=1)

        # ── Experiment blocks ──────────────────────────────────────────────────
        self._draw_all_experiments(days, col_w, ch)

    def _draw_all_experiments(self, days, col_w, ch):
        day_idx = {d: i for i, d in enumerate(days)}
        for exp in self.app.schedule:
            if exp.get("status") == "canceled":
                continue
            self._draw_exp_block(exp, day_idx, col_w, ch)

    def _draw_exp_block(self, exp, day_idx, col_w, ch, y_offset=0):
        c = self._canvas
        try:
            start_dt = datetime.fromtimestamp(exp["plannedStart"] / 1000)
            end_dt   = datetime.fromtimestamp(exp["plannedEnd"] / 1000)
        except (KeyError, TypeError, ValueError):
            return
        d = start_dt.date()
        if d not in day_idx:
            return
        i     = day_idx[d]
        x1    = self.TIME_W + i * col_w + 3
        x2    = self.TIME_W + (i + 1) * col_w - 3
        y1    = max(0, self._ts_to_y(exp["plannedStart"])) + y_offset
        y2    = min(ch, max(y1 + 28, self._ts_to_y(exp["plannedEnd"]) + y_offset))

        fill, mid, accent = self._exp_palette(exp["id"])
        is_sel = (exp["id"] == self._selected_id)
        outline = "#1d4ed8" if is_sel else accent
        bw      = 2        if is_sel else 1
        tag     = f"exp_{exp['id']}"

        c.create_rectangle(x1, y1, x2, y2,
                            fill=fill, outline=outline, width=bw,
                            tags=("exp_block", tag))
        # Accent bar on left edge
        c.create_rectangle(x1, y1, x1 + 5, y2,
                            fill=accent, outline="",
                            tags=("exp_block", tag))
        height = y2 - y1
        if height >= 18:
            c.create_text(x1 + 9, y1 + 4, anchor="nw",
                           text=exp.get("title", "Experiment"),
                           font=(FONT, 10, "bold"), fill="#1e293b",
                           width=max(10, x2 - x1 - 14),
                           tags=("exp_block", tag))
        if height >= 34:
            dur_min = int((exp["plannedEnd"] - exp["plannedStart"]) / 60000)
            time_str = (f"{start_dt.strftime('%I:%M')}–"
                        f"{end_dt.strftime('%I:%M %p')}  ·  {fmt_mins(dur_min)}")
            c.create_text(x1 + 9, y1 + 19, anchor="nw",
                           text=time_str, font=(FONT, 9),
                           fill="#475569", width=max(10, x2 - x1 - 14),
                           tags=("exp_block", tag))

    # ── Mouse interaction ──────────────────────────────────────────────────────
    def _find_exp_id(self, cx, cy):
        """Return the experiment ID of the topmost calendar block at (cx, cy)."""
        _SKIP = {"block"}          # "exp_block" → candidate "block" — skip it
        for item in reversed(self._canvas.find_overlapping(cx - 1, cy - 1, cx + 1, cy + 1)):
            for t in self._canvas.gettags(item):
                if t.startswith("exp_"):
                    candidate = t[4:]
                    if candidate not in _SKIP:
                        return candidate
        return None

    def _on_click(self, event):
        cx = self._canvas.canvasx(event.x)
        cy = self._canvas.canvasy(event.y)
        exp_id = self._find_exp_id(cx, cy)
        if exp_id:
            exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
            self._selected_id = exp_id
            self._drag = {
                "exp_id":     exp_id,
                "start_cx":   cx, "start_cy": cy,
                "dragging":   False,
                "orig_start": exp["plannedStart"] if exp else None,
                "orig_end":   exp["plannedEnd"]   if exp else None,
            }
            self._draw_grid(self._get_days())
            self._show_detail(exp_id)
        else:
            self._drag = None
            if self._selected_id:
                self._selected_id = None
                self._hide_detail()
                self._draw_grid(self._get_days())

    def _on_double_click(self, event):
        """Double-click a calendar block → open the experiment editor."""
        self._drag = None          # cancel drag state from the first click
        cx = self._canvas.canvasx(event.x)
        cy = self._canvas.canvasy(event.y)
        exp_id = self._find_exp_id(cx, cy)
        if exp_id:
            self._edit_exp(exp_id)

    def _on_drag(self, event):
        if not self._drag:
            return
        cy = self._canvas.canvasy(event.y)
        dy = cy - self._drag["start_cy"]
        if abs(dy) <= 4 and not self._drag["dragging"]:
            return
        self._drag["dragging"] = True
        self._canvas.config(cursor="fleur")
        layout = self._layout
        if not layout:
            return
        delta_ms  = round(dy * 60 / self.PX_PER_HR / self.MIN_SNAP) * self.MIN_SNAP * 60000
        new_start = self._drag["orig_start"] + delta_ms
        new_end   = self._drag["orig_end"]   + delta_ms
        exp = next((e for e in self.app.schedule
                    if e["id"] == self._drag["exp_id"]), None)
        if exp:
            tmp     = dict(exp, plannedStart=new_start, plannedEnd=new_end)
            day_idx = {d: i for i, d in enumerate(layout["days"])}
            self._canvas.delete(f"exp_{exp['id']}")
            self._draw_exp_block(tmp, day_idx, layout["col_w"], layout["ch"])

    def _on_release(self, event):
        self._canvas.config(cursor="arrow")
        if not self._drag:
            return
        drag = self._drag
        self._drag = None
        if not drag.get("dragging"):
            return
        cy       = self._canvas.canvasy(event.y)
        dy       = cy - drag["start_cy"]
        delta_ms = round(dy * 60 / self.PX_PER_HR / self.MIN_SNAP) * self.MIN_SNAP * 60000
        exp_id   = drag["exp_id"]
        for exp in self.app.schedule:
            if exp["id"] == exp_id:
                dur_ms = exp["plannedEnd"] - exp["plannedStart"]
                exp["plannedStart"] += delta_ms
                exp["plannedEnd"]    = exp["plannedStart"] + dur_ms
                # Shift all timeline blocks
                for blk in exp.get("timelineBlocks", []):
                    blk["startTime"] = blk.get("startTime", 0) + delta_ms
                    blk["endTime"]   = blk.get("endTime",   0) + delta_ms
                # Keep legacy scheduledSteps in sync too
                for ss in exp.get("scheduledSteps", []):
                    ss["plannedStart"] = ss.get("plannedStart", 0) + delta_ms
                    ss["plannedEnd"]   = ss.get("plannedEnd",   0) + delta_ms
                break
        save_schedule(self.app.schedule)
        self._draw_grid(self._get_days())
        if self._detail_visible and self._selected_id == exp_id:
            self._show_detail(exp_id)

    # ── Hover / tooltip ───────────────────────────────────────────────────────
    def _on_canvas_motion(self, event):
        """Update cursor and schedule a hover tooltip over experiment blocks."""
        if self._drag and self._drag.get("dragging"):
            return                  # don't interfere with active drag
        cx = self._canvas.canvasx(event.x)
        cy = self._canvas.canvasy(event.y)
        if self._tooltip_job:
            self.after_cancel(self._tooltip_job)
            self._tooltip_job = None
        self._canvas.delete("tooltip")
        exp_id = self._find_exp_id(cx, cy)
        if exp_id:
            self._canvas.config(cursor="hand2")
            self._tooltip_job = self.after(
                650, lambda: self._show_exp_tip(cx, cy, exp_id))
        else:
            self._canvas.config(cursor="arrow")

    def _on_canvas_leave(self, event):
        if self._tooltip_job:
            self.after_cancel(self._tooltip_job)
            self._tooltip_job = None
        self._canvas.delete("tooltip")
        self._canvas.config(cursor="arrow")

    def _show_exp_tip(self, cx, cy, exp_id):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        try:
            s   = datetime.fromtimestamp(exp["plannedStart"] / 1000)
            e   = datetime.fromtimestamp(exp["plannedEnd"]   / 1000)
            dur = int((exp["plannedEnd"] - exp["plannedStart"]) / 60000)
            blocks = exp.get("timelineBlocks", exp.get("scheduledSteps", []))
            n      = len(blocks)
            lines = [
                exp.get("title", "Experiment"),
                f"{s.strftime('%I:%M %p')} – {e.strftime('%I:%M %p')}",
                f"Duration: {fmt_mins(dur)}  ·  {n} block{'s' if n != 1 else ''}",
                "Click to view details  ·  Double-click to edit  ·  Drag to move",
            ]
        except Exception:
            lines = [exp.get("title", "?")]
        self._draw_tooltip(cx, cy, lines)

    def _draw_tooltip(self, cx, cy, lines):
        """Render a dark tooltip card directly on the body canvas."""
        c = self._canvas
        c.delete("tooltip")
        if not lines:
            return
        px, py, lh = 10, 6, 15
        tw = max(140, max(len(l) for l in lines) * 7 + px * 2)
        th = len(lines) * lh + py * 2
        # Keep inside visible canvas area
        vis_x2 = c.canvasx(c.winfo_width())
        x = cx + 14
        if x + tw > vis_x2:
            x = cx - tw - 6
        y = cy - th // 2
        # Background card
        c.create_rectangle(x, y, x + tw, y + th,
                            fill="#1e293b", outline="#334155", width=1,
                            tags="tooltip")
        # Blue left bar
        c.create_rectangle(x + 1, y + 1, x + 4, y + th - 1,
                            fill="#3b82f6", outline="", tags="tooltip")
        for i, line in enumerate(lines):
            c.create_text(x + px, y + py + i * lh, anchor="nw", text=line,
                           font=(FONT, 9, "bold") if i == 0 else (FONT, 9),
                           fill="#f8fafc" if i == 0 else "#94a3b8",
                           tags="tooltip")
        c.tag_raise("tooltip")

    # ── Detail panel ──────────────────────────────────────────────────────────
    def _show_detail(self, exp_id):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        self._detail_visible = True
        self._tl_exp_id = exp_id     # store for rebuild

        p = self._detail_frame
        p.grid(row=1, column=1, sticky="nsew", padx=(0, 12), pady=12)
        for w in p.winfo_children():
            w.destroy()

        # ── Fixed header (title + close button) ──────────────────────────────
        hdr = ctk.CTkFrame(p, fg_color="transparent")
        hdr.pack(side="top", fill="x", padx=12, pady=(12, 4))
        ctk.CTkLabel(hdr, text=exp.get("title", "?"),
                      font=(FONT, 14, "bold"), text_color=T1,
                      anchor="w", wraplength=210).pack(side="left", fill="x", expand=True)
        ctk.CTkButton(hdr, text="✕", command=self._hide_detail,
                       fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
                       text_color=T3, width=26, height=26,
                       font=(FONT, 13), corner_radius=R_SM).pack(side="right")

        # ── Fixed status + time metadata ──────────────────────────────────────
        meta_f = ctk.CTkFrame(p, fg_color="transparent")
        meta_f.pack(side="top", fill="x", padx=12, pady=(0, 4))
        st = exp.get("status", "planned")
        st_bg = {"planned": ACC, "in_progress": ORANGE,
                 "completed": GREEN, "canceled": ("#94a3b8", "#64748b")}.get(st, T3)
        sf = ctk.CTkFrame(meta_f, fg_color=st_bg, corner_radius=R_XS)
        sf.pack(side="left")
        ctk.CTkLabel(sf, text=st.replace("_", " ").title(),
                      font=(FONT, 10), text_color=("#fff", "#fff")).pack(padx=6, pady=2)
        try:
            s_dt = datetime.fromtimestamp(exp["plannedStart"] / 1000)
            e_dt = datetime.fromtimestamp(exp["plannedEnd"]   / 1000)
            dur  = int((exp["plannedEnd"] - exp["plannedStart"]) / 60000)
            t_str = (f"  {s_dt.strftime('%b')} {s_dt.day}  "
                     f"{s_dt.strftime('%I:%M').lstrip('0')}–"
                     f"{e_dt.strftime('%I:%M %p').lstrip('0')}  {fmt_mins(dur)}")
            ctk.CTkLabel(meta_f, text=t_str, font=(FONT, 10), text_color=T3,
                          anchor="w").pack(side="left")
        except Exception:
            pass

        ctk.CTkFrame(p, height=1, fg_color=CARD_B, corner_radius=0).pack(
            side="top", fill="x", padx=12, pady=(4, 0))

        # ── Fixed action buttons (bottom) ─────────────────────────────────────
        acts = ctk.CTkFrame(p, fg_color="transparent")
        acts.pack(side="bottom", fill="x", padx=12, pady=(4, 10))
        ctk.CTkFrame(p, height=1, fg_color=CARD_B, corner_radius=0).pack(
            side="bottom", fill="x", padx=12, pady=(0, 4))

        proto = next((pr for pr in self.app.protocols
                       if pr["id"] == exp.get("protocolId")), None)
        if proto:
            ctk.CTkButton(
                acts, text="▶  Run Mode",
                command=lambda pr=proto: (
                    self.app.pages["run"].start(pr), self.app.navigate("run")),
                fg_color=GREEN, hover_color=("#16a34a", "#22c55e"),
                text_color=("#fff", "#fff"), font=(FONT, 11),
                height=30, corner_radius=R_SM,
            ).pack(fill="x", pady=(0, 4))

        brow = ctk.CTkFrame(acts, fg_color="transparent")
        brow.pack(fill="x")
        brow.grid_columnconfigure((0, 1), weight=1)
        ctk.CTkButton(
            brow, text="✎ Edit",
            command=lambda eid=exp_id: self._edit_exp(eid),
            fg_color=("#e2e8f0", "#334155"), hover_color=("#cbd5e1", "#475569"),
            text_color=T1, font=(FONT, 11), height=28, corner_radius=R_SM,
        ).grid(row=0, column=0, sticky="ew", padx=(0, 3))
        ctk.CTkButton(
            brow, text="✕ Cancel",
            command=lambda eid=exp_id: self._cancel_exp(eid),
            fg_color=("#e2e8f0", "#334155"), hover_color=("#fecaca", "#7f1d1d"),
            text_color=DANGER, font=(FONT, 11), height=28, corner_radius=R_SM,
        ).grid(row=0, column=1, sticky="ew", padx=(3, 0))

        # ── Scrollable timeline (fills remaining space) ───────────────────────
        self._tl_frame = ctk.CTkScrollableFrame(
            p, fg_color="transparent",
            scrollbar_button_color=("#cbd5e1", "#475569"),
            scrollbar_button_hover_color=("#94a3b8", "#64748b"),
        )
        self._tl_frame.pack(fill="both", expand=True)
        self._rebuild_detail_timeline(exp_id)

    def _rebuild_detail_timeline(self, exp_id):
        """Clear and redraw the block list inside self._tl_frame."""
        if not hasattr(self, "_tl_frame"):
            return
        try:
            if not self._tl_frame.winfo_exists():
                return
        except Exception:
            return
        tl = self._tl_frame
        for w in tl.winfo_children():
            w.destroy()

        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        blocks = exp.get("timelineBlocks", [])

        # Notes (if any)
        if exp.get("notes"):
            nf = ctk.CTkFrame(tl, fg_color=("#fefce8", "#422006"),
                               corner_radius=R_XS, border_width=1,
                               border_color=("#fde68a", "#854d0e"))
            nf.pack(fill="x", padx=10, pady=(8, 4))
            ctk.CTkLabel(nf, text=exp["notes"], font=(FONT, 10), text_color=T2,
                          wraplength=240, justify="left", anchor="w").pack(
                padx=8, pady=4, anchor="w")

        # Section header with block count + "+ Add Block"
        sh = ctk.CTkFrame(tl, fg_color="transparent")
        sh.pack(fill="x", padx=10, pady=(6, 4))
        n = len(blocks)
        done_n = sum(1 for b in blocks if b.get("status") == "done")
        ctk.CTkLabel(sh,
                      text=(f"Timeline  ·  {n} block{'s' if n != 1 else ''}"
                            + (f"  ·  {done_n} done" if done_n else "")),
                      font=(FONT, 11, "bold"), text_color=T2, anchor="w"
                      ).pack(side="left")
        ctk.CTkButton(
            sh, text="＋ Add",
            command=lambda: self._show_add_block_menu(exp_id),
            fg_color=("#dbeafe", "#1e3a5f"), hover_color=("#bfdbfe", "#1e40af"),
            text_color=ACC, font=(FONT, 10), height=24, corner_radius=R_XS, width=66,
        ).pack(side="right")

        if not blocks:
            ctk.CTkLabel(tl, text="No blocks yet. Click ＋ Add to create one.",
                          font=(FONT, 11), text_color=T3).pack(
                padx=10, pady=12, anchor="w")
            return

        for idx, blk in enumerate(blocks):
            self._build_block_card(tl, exp_id, idx, blk, n)

        # Bottom spacer
        ctk.CTkFrame(tl, fg_color="transparent", height=8).pack()

    def _build_block_card(self, parent, exp_id, idx, blk, total):
        """Render one editable timeline block card."""
        btype  = blk.get("blockType", "protocol_step")
        status = blk.get("status", "planned")
        light, accent, badge_bg = BLOCK_TYPE_COLORS.get(
            btype, ("#f8fafc", "#64748b", "#d1d5db"))
        is_inactive = status in ("skipped", "canceled")
        # Grey out inactive blocks
        card_bg  = ("#f1f5f9", "#1e293b") if is_inactive else light
        card_bdr = ("#cbd5e1", "#334155") if is_inactive else badge_bg

        card = ctk.CTkFrame(parent, fg_color=card_bg, corner_radius=R_SM,
                              border_width=1, border_color=card_bdr)
        card.pack(fill="x", padx=10, pady=(0, 5))

        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="x", padx=8, pady=(6, 5))

        # ── Row 1: type badge + title + ↑↓ ──────────────────────────────────
        row1 = ctk.CTkFrame(inner, fg_color="transparent")
        row1.pack(fill="x")

        badge_lbl = BLOCK_TYPE_LABELS.get(btype, btype)
        bf = ctk.CTkFrame(row1, fg_color=badge_bg if not is_inactive else ("#94a3b8","#475569"),
                           corner_radius=R_XS)
        bf.pack(side="left", padx=(0, 5))
        ctk.CTkLabel(bf, text=badge_lbl, font=(FONT, 9),
                      text_color=("#1e293b", "#f8fafc")).pack(padx=5, pady=1)

        title_color = T3 if is_inactive else T1
        ctk.CTkLabel(row1, text=blk.get("title", "Block"),
                      font=(FONT, 11, "bold"), text_color=title_color,
                      anchor="w").pack(side="left", fill="x", expand=True, padx=(0, 2))

        # ↑↓ reorder buttons
        if idx < total - 1:
            ctk.CTkButton(
                row1, text="↓", width=20, height=18,
                fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
                text_color=T3, font=(FONT, 9), corner_radius=R_XS,
                command=lambda i=idx: self._move_block(exp_id, i, 1),
            ).pack(side="right", padx=(1, 0))
        if idx > 0:
            ctk.CTkButton(
                row1, text="↑", width=20, height=18,
                fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
                text_color=T3, font=(FONT, 9), corner_radius=R_XS,
                command=lambda i=idx: self._move_block(exp_id, i, -1),
            ).pack(side="right", padx=(1, 0))

        # ── Row 2: time range + duration ─────────────────────────────────────
        row2 = ctk.CTkFrame(inner, fg_color="transparent")
        row2.pack(fill="x", pady=(2, 0))
        try:
            s_dt = datetime.fromtimestamp(blk["startTime"] / 1000)
            e_dt = datetime.fromtimestamp(blk["endTime"]   / 1000)
            if blk["startTime"] == blk["endTime"]:
                t_str = "—  (no time)"
            else:
                t_str = (f"{s_dt.strftime('%I:%M').lstrip('0') or '12'}"
                         f"–{e_dt.strftime('%I:%M %p').lstrip('0') or '12'}")
        except Exception:
            t_str = "—"
        ctk.CTkLabel(row2, text=t_str, font=(FONT, 10),
                      text_color=T3 if is_inactive else T2).pack(side="left")
        ctk.CTkLabel(row2, text=fmt_mins(blk.get("durationMinutes", 0)),
                      font=(FONT, 10), text_color=T3).pack(side="right")

        # ── Row 3: status indicator (if not planned) ─────────────────────────
        if status != "planned":
            row3 = ctk.CTkFrame(inner, fg_color="transparent")
            row3.pack(fill="x", pady=(2, 0))
            st_icons = {"done": "✓ Done", "skipped": "⊘ Skipped",
                        "canceled": "✕ Canceled", "modified": "✎ Modified"}
            st_colors = {"done": "#22c55e", "skipped": "#94a3b8",
                         "canceled": "#ef4444", "modified": "#f59e0b"}
            ctk.CTkLabel(row3, text=st_icons.get(status, status.title()),
                          font=(FONT, 9, "bold"),
                          text_color=st_colors.get(status, T3)).pack(side="left")

        # ── Row 4: hands-on / wait (protocol step only) ───────────────────────
        if btype == "protocol_step":
            ho = blk.get("handsOnMinutes", 0)
            wt = blk.get("waitMinutes",    0)
            if ho or wt:
                row4 = ctk.CTkFrame(inner, fg_color="transparent")
                row4.pack(fill="x", pady=(2, 0))
                if ho:
                    ctk.CTkLabel(row4, text=f"✋ {fmt_mins(ho)}",
                                  font=(FONT, 9), text_color="#3b82f6").pack(
                        side="left", padx=(0, 8))
                if wt:
                    ctk.CTkLabel(row4, text=f"⏳ {fmt_mins(wt)}",
                                  font=(FONT, 9), text_color="#f97316").pack(side="left")

        # ── Row 5: notes ─────────────────────────────────────────────────────
        if blk.get("notes"):
            ctk.CTkLabel(inner, text=blk["notes"], font=(FONT, 9), text_color=T3,
                          wraplength=220, justify="left", anchor="w"
                          ).pack(anchor="w", pady=(3, 0))

        # ── Action row: Edit · Done/Skip/Restore · ⋯ (more) · ✕ ─────────────
        act_row = ctk.CTkFrame(inner, fg_color="transparent")
        act_row.pack(fill="x", pady=(5, 0))

        ctk.CTkButton(
            act_row, text="Edit", width=42, height=22,
            fg_color=("#dbeafe", "#1e3a5f"), hover_color=("#bfdbfe", "#1e40af"),
            text_color=ACC, font=(FONT, 10), corner_radius=R_XS,
            command=lambda i=idx: self._open_block_editor(exp_id, i),
        ).pack(side="left", padx=(0, 3))

        if status == "planned":
            ctk.CTkButton(
                act_row, text="✓ Done", width=54, height=22,
                fg_color=("#dcfce7", "#052e16"), hover_color=("#bbf7d0", "#064e3b"),
                text_color=GREEN, font=(FONT, 10), corner_radius=R_XS,
                command=lambda i=idx: self._mark_block_status(exp_id, i, "done"),
            ).pack(side="left", padx=(0, 3))
            ctk.CTkButton(
                act_row, text="⊘ Skip", width=50, height=22,
                fg_color=("#f1f5f9", "#1e293b"), hover_color=("#e2e8f0", "#334155"),
                text_color=T2, font=(FONT, 10), corner_radius=R_XS,
                command=lambda i=idx: self._ask_skip_block(exp_id, i),
            ).pack(side="left", padx=(0, 3))
        elif is_inactive:
            ctk.CTkButton(
                act_row, text="↩ Restore", width=66, height=22,
                fg_color=("#e2e8f0", "#334155"), hover_color=("#cbd5e1", "#475569"),
                text_color=T2, font=(FONT, 10), corner_radius=R_XS,
                command=lambda i=idx: self._mark_block_status(exp_id, i, "planned"),
            ).pack(side="left", padx=(0, 3))

        # ⋯ more menu and × delete
        ctk.CTkButton(
            act_row, text="✕", width=24, height=22,
            fg_color="transparent", hover_color=("#fecaca", "#7f1d1d"),
            text_color=DANGER, font=(FONT, 10), corner_radius=R_XS,
            command=lambda i=idx: self._confirm_delete_block(exp_id, i),
        ).pack(side="right", padx=(0, 0))
        ctk.CTkButton(
            act_row, text="⋯", width=26, height=22,
            fg_color="transparent", hover_color=("#e2e8f0", "#334155"),
            text_color=T3, font=(FONT, 12), corner_radius=R_XS,
            command=lambda i=idx, b=blk: self._show_block_ctx_menu(exp_id, i, b),
        ).pack(side="right", padx=(0, 2))

    # ── Block actions ─────────────────────────────────────────────────────────
    def _open_block_editor(self, exp_id, block_idx):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        blocks = exp.get("timelineBlocks", [])
        if block_idx >= len(blocks):
            return
        BlockEditDialog(
            self, blocks[block_idx],
            lambda updated: self._on_block_edited(exp_id, block_idx, updated))

    def _on_block_edited(self, exp_id, block_idx, updated_block):
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            if block_idx < len(blocks):
                blocks[block_idx] = updated_block
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _mark_block_status(self, exp_id, block_idx, new_status):
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            if block_idx < len(blocks):
                blocks[block_idx]["status"] = new_status
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _ask_skip_block(self, exp_id, block_idx):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        blocks = exp.get("timelineBlocks", [])
        if block_idx >= len(blocks):
            return
        title = blocks[block_idx].get("title", "Block")
        SkipBlockDialog(self, title,
                        lambda keep: self._apply_skip_block(exp_id, block_idx, keep))

    def _apply_skip_block(self, exp_id, block_idx, keep_time):
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            if block_idx < len(blocks):
                blocks[block_idx]["status"]   = "skipped"
                blocks[block_idx]["keepTime"] = keep_time
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _confirm_delete_block(self, exp_id, block_idx):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        blocks = exp.get("timelineBlocks", [])
        if block_idx >= len(blocks):
            return
        title = blocks[block_idx].get("title", "Block")
        if messagebox.askyesno("Delete Block",
                                f'Delete "{title}"?\nThis cannot be undone.'):
            self._delete_block(exp_id, block_idx)

    def _delete_block(self, exp_id, block_idx):
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            if block_idx < len(blocks):
                blocks.pop(block_idx)
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _move_block(self, exp_id, block_idx, direction):
        """Move block up (-1) or down (+1), recalculate, and rebuild."""
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            new_idx = block_idx + direction
            if 0 <= new_idx < len(blocks):
                blocks[block_idx], blocks[new_idx] = blocks[new_idx], blocks[block_idx]
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _show_add_block_menu(self, exp_id):
        """Open ScheduleBlockDialog to add a new timeline block."""
        ScheduleBlockDialog(self, "break",
                            lambda blk: self._add_block(exp_id, blk))

    def _add_block(self, exp_id, new_block):
        """Append a new block to the experiment timeline."""
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.setdefault("timelineBlocks", [])
            blocks.append(new_block)
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _insert_block_at(self, exp_id, after_idx, new_block):
        """Insert a new block after after_idx (or at end if -1)."""
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.setdefault("timelineBlocks", [])
            insert_pos = after_idx + 1 if after_idx >= 0 else len(blocks)
            blocks.insert(insert_pos, new_block)
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _show_block_ctx_menu(self, exp_id, block_idx, blk):
        """Show right-click style context menu for a block."""
        btype  = blk.get("blockType", "protocol_step")
        status = blk.get("status", "planned")
        actions = [
            ("✎  Edit",    lambda: self._open_block_editor(exp_id, block_idx)),
            ("⧉  Duplicate",
             lambda: self._duplicate_block(exp_id, block_idx)),
            None,
            ("＋  Insert block before",
             lambda: self._open_add_before(exp_id, block_idx)),
            ("＋  Insert block after",
             lambda: self._open_add_after(exp_id, block_idx)),
            None,
        ]
        if status == "planned":
            actions += [
                ("✓  Mark as done",
                 lambda: self._mark_block_status(exp_id, block_idx, "done")),
                ("⊘  Skip (remove time)",
                 lambda: self._apply_skip_block(exp_id, block_idx, False)),
                ("⊘  Skip (keep time)",
                 lambda: self._apply_skip_block(exp_id, block_idx, True)),
            ]
        elif status in ("skipped", "canceled", "done"):
            actions += [
                ("↩  Restore to planned",
                 lambda: self._mark_block_status(exp_id, block_idx, "planned")),
            ]
        if btype == "protocol_step":
            actions += [
                None,
                ("→  Convert to Break",
                 lambda: self._convert_block_type(exp_id, block_idx, "break")),
                ("→  Convert to Task",
                 lambda: self._convert_block_type(exp_id, block_idx, "task")),
                ("→  Convert to Note",
                 lambda: self._convert_block_type(exp_id, block_idx, "note")),
            ]
        actions += [
            None,
            ("Delete block",
             lambda: self._confirm_delete_block(exp_id, block_idx)),
        ]
        # Position menu near the card (use mouse pointer position)
        x = self.winfo_pointerx()
        y = self.winfo_pointery()
        BlockContextMenu(self, x, y, actions)

    def _duplicate_block(self, exp_id, block_idx):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if not exp:
            return
        blocks = exp.get("timelineBlocks", [])
        if block_idx >= len(blocks):
            return
        import copy
        dup = copy.deepcopy(blocks[block_idx])
        dup["id"] = new_id()
        dup["title"] = dup.get("title", "Block") + " (copy)"
        self._insert_block_at(exp_id, block_idx, dup)

    def _open_add_before(self, exp_id, block_idx):
        ScheduleBlockDialog(self, "task",
                            lambda blk: self._insert_block_at(exp_id, block_idx - 1, blk))

    def _open_add_after(self, exp_id, block_idx):
        ScheduleBlockDialog(self, "task",
                            lambda blk: self._insert_block_at(exp_id, block_idx, blk))

    def _convert_block_type(self, exp_id, block_idx, new_type):
        for exp in self.app.schedule:
            if exp["id"] != exp_id:
                continue
            blocks = exp.get("timelineBlocks", [])
            if block_idx < len(blocks):
                blocks[block_idx]["blockType"] = new_type
            _recalc_timeline(exp)
            break
        save_schedule(self.app.schedule)
        self._render_calendar()
        if self._detail_visible and self._selected_id == exp_id:
            self._rebuild_detail_timeline(exp_id)

    def _hide_detail(self):
        self._detail_visible = False
        self._selected_id    = None
        try:
            self._detail_frame.grid_remove()
        except Exception:
            pass
        self._draw_grid(self._get_days())

    # ── Actions ───────────────────────────────────────────────────────────────
    def _new_experiment(self):
        ScheduleExperimentDialog(self, self.app, None, self._on_exp_saved)

    def _edit_exp(self, exp_id):
        exp = next((e for e in self.app.schedule if e["id"] == exp_id), None)
        if exp:
            ScheduleExperimentDialog(self, self.app, exp, self._on_exp_saved)

    def _on_exp_saved(self, exp):
        sched = self.app.schedule
        idx = next((i for i, e in enumerate(sched) if e["id"] == exp["id"]), None)
        if idx is not None:
            sched[idx] = exp
        else:
            sched.insert(0, exp)
        save_schedule(sched)
        self.app.schedule = sched
        self._selected_id = exp["id"]
        self._render_calendar()
        self._show_detail(exp["id"])

    def _cancel_exp(self, exp_id):
        if not messagebox.askyesno("Cancel Experiment",
                                    "Mark this experiment as canceled?\n"
                                    "It will be hidden from the calendar."):
            return
        for e in self.app.schedule:
            if e["id"] == exp_id:
                e["status"] = "canceled"
                break
        save_schedule(self.app.schedule)
        self._hide_detail()
        self._render_calendar()


# ══════════════════════════════════════════════════════════════════════════════
# Main App
# ══════════════════════════════════════════════════════════════════════════════
NAV_ITEMS = [
    ("dashboard",  "⬛  Dashboard"),
    ("library",    "📋  Library"),
    ("schedule",   "🗓  Schedule"),
    ("flowchart",  "⎇  Flowchart"),
    ("run",        "▶  Run Mode"),
    ("history",    "📓  Lab Notebook"),
    ("settings",   "⚙  Settings"),
]

class BenchFlowApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("BenchFlow")
        self.geometry("1280x800")
        self.minsize(900, 600)
        self.configure(fg_color=("#e8eef5", "#0f172a"))

        self.protocols = load_protocols()
        self.runs = load_runs()
        self.schedule = load_schedule()
        self.categories = load_categories()
        self.tags = load_tags()
        self.templates = load_templates()

        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self._build_sidebar()
        self._build_pages()
        self.navigate("dashboard")

    def _build_sidebar(self):
        sb = ctk.CTkFrame(self, fg_color=SB, corner_radius=R_XL, width=190,
                          border_width=1, border_color=("#334155", "#1e293b"))
        sb.grid(row=0, column=0, sticky="nsew", padx=(12, 0), pady=12)
        sb.grid_propagate(False)
        sb.grid_rowconfigure(10, weight=1)

        # Logo
        logo = ctk.CTkFrame(sb, fg_color="transparent")
        logo.grid(row=0, column=0, sticky="ew", padx=16, pady=(20,16))
        ctk.CTkLabel(logo, text="BenchFlow", font=(FONT, 18, "bold"),
                     text_color=("#60a5fa","#60a5fa")).pack(anchor="w")
        ctk.CTkLabel(logo, text="Wet Lab Manager", font=(FONT, 11),
                     text_color=SB_TXT).pack(anchor="w")

        separator(sb).grid(row=1, column=0, sticky="ew", padx=12, pady=(0,8))

        self._nav_buttons = {}
        for i, (page_id, label_txt) in enumerate(NAV_ITEMS):
            b = ctk.CTkButton(
                sb, text=label_txt, anchor="w",
                font=(FONT, 13), height=40, corner_radius=R_LG,
                fg_color="transparent", hover_color=SB_HOV,
                text_color=SB_TXT,
                command=lambda p=page_id: self.navigate(p)
            )
            b.grid(row=i+2, column=0, sticky="ew", padx=10, pady=3)
            self._nav_buttons[page_id] = b

        separator(sb).grid(row=11, column=0, sticky="ew", padx=12, pady=8)
        ctk.CTkLabel(sb, text="Local · No cloud · No account",
                     font=(FONT, 10), text_color=SB_TXT).grid(row=12, column=0, padx=16, pady=(0,16))

    def _build_pages(self):
        container = ctk.CTkFrame(self, fg_color=BG, corner_radius=R_XL,
                                 border_width=1, border_color=("#dbe3ee", "#263244"))
        container.grid(row=0, column=1, sticky="nsew", padx=12, pady=12)
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)

        self.pages = {
            "dashboard": DashboardPage(container, self),
            "library":   LibraryPage(container, self),
            "editor":    EditorPage(container, self),
            "import":    ImportPage(container, self),
            "schedule":  SchedulePage(container, self),
            "flowchart": FlowchartPage(container, self),
            "run":       RunPage(container, self),
            "history":   HistoryPage(container, self),
            "settings":  SettingsPage(container, self),
        }
        for page in self.pages.values():
            page.grid(row=0, column=0, sticky="nsew")

    def navigate(self, page_id):
        # Editor is accessed from Library — keep Library highlighted in sidebar
        nav_highlight = "library" if page_id in ("editor", "import") else page_id
        for pid, b in self._nav_buttons.items():
            if pid == nav_highlight:
                b.configure(fg_color=SB_ACT, text_color=SB_ATXT)
            else:
                b.configure(fg_color="transparent", text_color=SB_TXT)

        for pid, page in self.pages.items():
            if pid == page_id:
                page.tkraise()
                page.refresh()

    def open_editor(self, protocol_id):
        if protocol_id is None:
            p = new_protocol()
        else:
            p = next((x for x in self.protocols if x["id"] == protocol_id), None)
            if p is None: return
        self.pages["editor"].load(p)
        self.navigate("editor")

    def open_flowchart(self, protocol_id):
        p = next((x for x in self.protocols if x["id"] == protocol_id), None)
        if p is None: return
        self.pages["flowchart"].load(p)
        self.navigate("flowchart")

    def open_create_protocol(self, mode="full"):
        ProtocolCreateDialog(self, self, mode=mode)

    def open_import(self, tab="text"):
        self.pages["import"].start(tab)
        self.navigate("import")

    def start_run(self, protocol_id):
        p = next((x for x in self.protocols if x["id"] == protocol_id), None)
        if p is None: return
        if not p.get("steps"):
            messagebox.showwarning("No Steps", "Add at least one step before running.")
            return
        self.pages["run"].start(p)
        self.navigate("run")


if __name__ == "__main__":
    app = BenchFlowApp()
    app.mainloop()
